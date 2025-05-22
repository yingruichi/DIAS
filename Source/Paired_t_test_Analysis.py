import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 指定中文字体，SimHei 是黑体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "配对 t 检验分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "配对 t 检验": "用于比较两个相关样本的均值是否有显著差异。",
            "样本量": "每个样本中的观测值数量。",
            "均值": "样本数据的平均值。",
            "t 统计量": "用于衡量两组样本均值差异的统计量。",
            "自由度": "在统计计算中能够自由取值的变量个数。",
            "p 值": "用于判断两组样本均值是否有显著差异的概率值。",
            "置信区间": "均值差异可能存在的区间范围。"
        },
        'interpretation': {
            "t 统计量": "t 统计量的绝对值越大，说明两组样本均值差异越显著。",
            "p 值": "p 值小于显著性水平（通常为 0.05）时，拒绝原假设，认为两组样本均值存在显著差异；否则，接受原假设，认为两组样本均值无显著差异。",
            "自由度": "自由度影响 t 分布的形状，自由度越大，t 分布越接近正态分布。",
            "置信区间": "如果置信区间不包含 0，说明两组样本均值存在显著差异。"
        }
    },
    'en': {
        'title': "Paired t-test Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "配对 t 检验": "Used to compare whether the means of two related samples are significantly different.",
            "样本量": "The number of observations in each sample.",
            "均值": "The average value of the sample data.",
            "t 统计量": "A statistic used to measure the difference between the means of two groups of samples.",
            "自由度": "The number of variables that can take on independent values in a statistical calculation.",
            "p 值": "A probability value used to determine whether there is a significant difference between the means of two groups of samples.",
            "置信区间": "The range within which the difference in means may lie."
        },
        'interpretation': {
            "t 统计量": "The larger the absolute value of the t statistic, the more significant the difference between the means of the two groups of samples.",
            "p 值": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between the means of the two groups of samples; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "自由度": "The degrees of freedom affect the shape of the t-distribution. The larger the degrees of freedom, the closer the t-distribution is to the normal distribution.",
            "置信区间": "If the confidence interval does not contain 0, it indicates a significant difference between the means of the two groups of samples."
        }
    }
}


class PairedTTestAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行配对 t 检验。")
            if numerical_df.shape[1] != 2:
                raise ValueError("数据必须包含且仅包含两列数值数据，才能进行配对 t 检验。")

            # 进行配对 t 检验
            t_stat, p_value = stats.ttest_rel(*numerical_df.T.values)
            df_value = len(numerical_df) - 1  # 自由度
            mean_diff = numerical_df.iloc[:, 0].mean() - numerical_df.iloc[:, 1].mean()
            std_err = stats.sem(numerical_df.iloc[:, 0] - numerical_df.iloc[:, 1])
            conf_int = stats.t.interval(0.95, df_value, loc=mean_diff, scale=std_err)

            # 计算样本量和均值
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()

            # 整理数据
            data = [
                ["配对 t 检验", t_stat, df_value, p_value, conf_int],
                ["样本量", sample_sizes.to_dict(), "", "", ""],
                ["均值", means.to_dict(), "", "", ""]
            ]
            headers = ["统计量", "t 统计量", "自由度", "p 值", "置信区间"]
            result_df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["配对 t 检验", "样本量", "均值", "t 统计量", "自由度", "p 值", "置信区间"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["t 统计量", "p 值", "自由度", "置信区间"])
            interpretation_df.insert(0, "统计量_结果解读",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                doc.add_heading('分析结果', level=1)
                table = doc.add_table(rows=result_df.shape[0] + 1, cols=result_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(result_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in result_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明表格
                doc.add_heading('解释说明', level=1)
                table = doc.add_table(rows=explanation_df.shape[0] + 1, cols=explanation_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(explanation_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in explanation_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加结果解读表格
                doc.add_heading('结果解读', level=1)
                table = doc.add_table(rows=interpretation_df.shape[0] + 1, cols=interpretation_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(interpretation_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in interpretation_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 绘制图表
                plot_path = self.plot_results(numerical_df, save_path)
                if plot_path:
                    doc.add_heading('图表', level=1)
                    doc.add_picture(plot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def plot_results(self, numerical_df, save_path):
        # 柱状图
        plt.figure(figsize=(12, 8))
        plt.subplot(2, 2, 1)
        means = numerical_df.mean()
        bars = plt.bar(means.index, means)
        for bar in bars:
            height = bar.get_height()
            plt.annotate(f'{height:.2f}',
                         xy=(bar.get_x() + bar.get_width() / 2, height),
                         xytext=(0, 3),  # 3 points vertical offset
                         textcoords="offset points",
                         ha='center', va='bottom')
        plt.title('柱状图' if self.current_language == 'zh' else 'Bar Chart')
        plt.xlabel('样本' if self.current_language == 'zh' else 'Sample')
        plt.ylabel('均值' if self.current_language == 'zh' else 'Mean')

        # 误差线图
        plt.subplot(2, 2, 2)
        means = numerical_df.mean()
        stds = numerical_df.std()
        plt.errorbar(means.index, means, yerr=stds, fmt='o')
        plt.title('误差线图' if self.current_language == 'zh' else 'Error Bar Chart')
        plt.xlabel('样本' if self.current_language == 'zh' else 'Sample')
        plt.ylabel('均值' if self.current_language == 'zh' else 'Mean')

        # 箱线图
        plt.subplot(2, 2, 3)
        numerical_df.boxplot()
        plt.title('箱线图' if self.current_language == 'zh' else 'Box Plot')
        plt.xlabel('样本' if self.current_language == 'zh' else 'Sample')
        plt.ylabel('数值' if self.current_language == 'zh' else 'Value')

        # 折线图
        plt.subplot(2, 2, 4)
        plt.plot(numerical_df)
        plt.title('折线图' if self.current_language == 'zh' else 'Line Chart')
        plt.xlabel('观测值' if self.current_language == 'zh' else 'Observation')
        plt.ylabel('数值' if self.current_language == 'zh' else 'Value')
        plt.legend(numerical_df.columns)

        plt.tight_layout()
        plot_path = save_path.replace('.docx', '_plots.png')
        plt.savefig(plot_path)
        plt.close()
        return plot_path

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = PairedTTestAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()