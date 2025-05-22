import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "模糊层次分析法 FAHP 分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "模糊特征向量": "反映各因素相对重要性的模糊向量",
            "一致性指标 CI": "衡量模糊判断矩阵一致性的指标",
            "随机一致性指标 RI": "根据矩阵阶数确定的随机一致性指标",
            "一致性比率 CR": "CI 与 RI 的比值，判断矩阵是否具有满意一致性"
        },
        'interpretation': {
            "模糊特征向量": "模糊特征向量值越大，对应因素越重要",
            "一致性指标 CI": "CI 值越小，矩阵一致性越好",
            "随机一致性指标 RI": "不同阶数矩阵有对应标准值",
            "一致性比率 CR": "CR < 0.1 时，矩阵具有满意一致性，结果可信"
        }
    },
    'en': {
        'title': "Fuzzy Analytic Hierarchy Process (FAHP) Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "模糊特征向量": "A fuzzy vector reflecting the relative importance of each factor",
            "一致性指标 CI": "An indicator to measure the consistency of the fuzzy judgment matrix",
            "随机一致性指标 RI": "A random consistency indicator determined by the order of the matrix",
            "一致性比率 CR": "The ratio of CI to RI to determine if the matrix has satisfactory consistency"
        },
        'interpretation': {
            "模糊特征向量": "The larger the value in the fuzzy eigenvector, the more important the corresponding factor",
            "一致性指标 CI": "The smaller the CI value, the better the consistency of the matrix",
            "随机一致性指标 RI": "There are corresponding standard values for matrices of different orders",
            "一致性比率 CR": "When CR < 0.1, the matrix has satisfactory consistency and the results are reliable"
        }
    }
}

# 随机一致性指标 RI 表
RI_TABLE = {
    1: 0, 2: 0, 3: 0.58, 4: 0.90, 5: 1.12, 6: 1.24, 7: 1.32, 8: 1.41, 9: 1.45
}

class FuzzyAnalyticHierarchyProcessFAHPApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.configure(style="TEntry")

    def on_entry_click(self, event):
        """当用户点击输入框时，清除提示信息"""
        if self.file_entry.get() == LANGUAGES[self.current_language]['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.configure(style="TEntry")

    def on_focusout(self, event):
        """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def fahp_analysis(self, data):
        """
        进行模糊层次分析法 FAHP 分析
        :param data: 模糊判断矩阵数据
        :return: 模糊特征向量、一致性指标 CI、一致性比率 CR
        """
        # 计算模糊特征向量
        row_sums = np.sum(data, axis=1)
        fuzzy_eigenvector = row_sums / np.sum(row_sums)

        # 计算模糊判断矩阵的最大特征值
        weighted_sum = np.dot(data, fuzzy_eigenvector)
        max_eigenvalue = np.sum(weighted_sum / (fuzzy_eigenvector * len(data)))

        # 计算一致性指标 CI
        n = data.shape[0]
        CI = (max_eigenvalue - n) / (n - 1)

        # 计算随机一致性指标 RI
        RI = RI_TABLE.get(n, None)
        if RI is None:
            raise ValueError("判断矩阵阶数超出支持范围")

        # 计算一致性比率 CR
        CR = CI / RI

        return fuzzy_eigenvector, CI, RI, CR

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]['file_entry_placeholder']:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 进行 FAHP 分析
            fuzzy_eigenvector, CI, RI, CR = self.fahp_analysis(data)

            # 整理数据
            data = [
                ["模糊特征向量", fuzzy_eigenvector.tolist(), ""],
                ["一致性指标 CI", CI, ""],
                ["随机一致性指标 RI", RI, ""],
                ["一致性比率 CR", CR, ""]
            ]
            headers = ["统计量", "统计量值", "p值"]

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('FAHP 分析结果', 0)

                # 添加分析结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for col, header in enumerate(headers):
                    hdr_cells[col].text = header
                for row_data in data:
                    row_cells = table.add_row().cells
                    for col, value in enumerate(row_data):
                        row_cells[col].text = str(value)

                # 添加解释说明
                doc.add_heading('解释说明', level=1)
                for key, value in explanations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 添加结果解读
                doc.add_heading('结果解读', level=1)
                for key, value in interpretations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 生成模糊特征向量柱状图
                fig, ax = plt.subplots()
                ax.bar(range(len(fuzzy_eigenvector)), fuzzy_eigenvector)
                ax.set_title('模糊特征向量柱状图' if self.current_language == 'zh' else 'Bar Chart of Fuzzy Eigenvector')
                ax.set_xlabel('因素' if self.current_language == 'zh' else 'Factors')
                ax.set_ylabel('权重' if self.current_language == 'zh' else 'Weights')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_fuzzy_eigenvector.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading('模糊特征向量柱状图', level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.configure(style="Gray.TEntry")

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建文件选择按钮
        self.select_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['select_button'], command=self.select_file,
                                        bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(self.root, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack(pady=5)
        self.file_entry.bind("<FocusIn>", self.on_entry_click)
        self.file_entry.bind("<FocusOut>", self.on_focusout)

        # 创建分析按钮
        self.analyze_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['analyze_button'], command=self.analyze_file,
                                         bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建语言切换标签
        self.switch_language_label = ttk.Label(self.root, text=LANGUAGES[self.current_language]['switch_language'], cursor="hand2")
        self.switch_language_label.pack(pady=10)
        self.switch_language_label.bind("<Button-1>", self.switch_language)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = FuzzyAnalyticHierarchyProcessFAHPApp()
    app.run()

if __name__ == "__main__":
    run_app()