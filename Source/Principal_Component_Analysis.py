import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from sklearn.decomposition import PCA
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "主成分分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "switch_language_button_text": "切换语言",
        "explanation": {
            "主成分载荷矩阵": "显示每个变量在各个主成分上的载荷，反映变量与主成分的相关性",
            "主成分得分": "每个样本在各个主成分上的得分",
            "特征值和方差贡献率": "特征值表示每个主成分解释的总方差，方差贡献率表示每个主成分解释的方差占总方差的比例",
            "碎石图": "展示特征值随主成分数量的变化情况，帮助确定主成分的数量"
        },
        "interpretation": {
            "主成分载荷矩阵": "绝对值越大，说明变量与主成分的相关性越强",
            "主成分得分": "得分越高，说明样本在该主成分上的特征越明显",
            "特征值和方差贡献率": "特征值大于1的主成分通常被保留，方差贡献率越高，说明该主成分越重要",
            "碎石图": "曲线的拐点处通常表示合适的主成分数量"
        }
    },
    "en": {
        "title": "Principal Component Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "switch_language_button_text": "Switch Language",
        "explanation": {
            "主成分载荷矩阵": "Shows the loadings of each variable on each principal component, reflecting the correlation between variables and principal components",
            "主成分得分": "The scores of each sample on each principal component",
            "特征值和方差贡献率": "The eigenvalue represents the total variance explained by each principal component, and the variance contribution rate represents the proportion of variance explained by each principal component to the total variance",
            "碎石图": "Shows the change of eigenvalues with the number of principal components, helping to determine the number of principal components"
        },
        "interpretation": {
            "主成分载荷矩阵": "The larger the absolute value, the stronger the correlation between the variable and the principal component",
            "主成分得分": "The higher the score, the more obvious the characteristics of the sample on this principal component",
            "特征值和方差贡献率": "Principal components with eigenvalues greater than 1 are usually retained. The higher the variance contribution rate, the more important the principal component",
            "碎石图": "The inflection point of the curve usually indicates the appropriate number of principal components"
        }
    }
}


class PrincipalComponentAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def pca_analysis(self, data):
        """
        进行主成分分析
        :param data: 输入数据
        :return: 主成分载荷矩阵、主成分得分、特征值、方差贡献率
        """
        # 创建主成分分析对象
        pca = PCA()
        pca.fit(data)

        # 计算特征值和方差贡献率
        ev = pca.explained_variance_
        v = pca.explained_variance_ratio_

        # 确定主成分数量
        num_components = sum(ev > 1)
        num_components = max(1, num_components)  # 确保至少有一个主成分

        # 重新进行主成分分析
        pca = PCA(n_components=num_components)
        scores = pca.fit_transform(data)
        loadings = pca.components_.T

        return loadings, scores, ev, v

    def plot_scree_plot(self, ev, save_path):
        """
        绘制碎石图
        :param ev: 特征值
        :param save_path: 图片保存路径
        """
        plt.figure(figsize=(10, 5))
        plt.plot(range(1, len(ev) + 1), ev, marker='o')
        plt.title('碎石图' if self.current_language == 'zh' else 'Scree Plot')
        plt.xlabel('主成分数量' if self.current_language == 'zh' else 'Number of Principal Components')
        plt.ylabel('特征值' if self.current_language == 'zh' else 'Eigenvalues')
        img_path = os.path.splitext(save_path)[0] + '_scree_plot.png'
        plt.savefig(img_path)
        plt.close()
        return img_path

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 进行主成分分析
            loadings, scores, ev, v = self.pca_analysis(df)

            # 整理数据
            component_names = [f'主成分{i + 1}' for i in range(len(loadings[0]))]
            loadings_df = pd.DataFrame(loadings, index=df.columns, columns=component_names)
            scores_df = pd.DataFrame(scores, columns=component_names)
            ev_df = pd.DataFrame(ev, columns=['特征值'])
            v_df = pd.DataFrame(v, columns=['方差贡献率'])

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加主成分载荷矩阵
                doc.add_heading('主成分载荷矩阵', level=1)
                table = doc.add_table(rows=loadings_df.shape[0] + 1, cols=loadings_df.shape[1] + 1)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '变量'
                for col_num, col_name in enumerate(loadings_df.columns):
                    hdr_cells[col_num + 1].text = col_name
                for row_num, (index, row) in enumerate(loadings_df.iterrows()):
                    row_cells = table.rows[row_num + 1].cells
                    row_cells[0].text = index
                    for col_num, value in enumerate(row):
                        row_cells[col_num + 1].text = str(value)

                # 添加主成分得分
                doc.add_heading('主成分得分', level=1)
                table = doc.add_table(rows=scores_df.shape[0] + 1, cols=scores_df.shape[1] + 1)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '样本'
                for col_num, col_name in enumerate(scores_df.columns):
                    hdr_cells[col_num + 1].text = col_name
                for row_num, (index, row) in enumerate(scores_df.iterrows()):
                    row_cells = table.rows[row_num + 1].cells
                    row_cells[0].text = str(index)
                    for col_num, value in enumerate(row):
                        row_cells[col_num + 1].text = str(value)

                # 添加特征值和方差贡献率
                doc.add_heading('特征值和方差贡献率', level=1)
                table = doc.add_table(rows=ev_df.shape[0] + 1, cols=ev_df.shape[1] + v_df.shape[1])
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '主成分'
                hdr_cells[1].text = '特征值'
                hdr_cells[2].text = '方差贡献率'
                for row_num, ((_, ev_row), (_, v_row)) in enumerate(zip(ev_df.iterrows(), v_df.iterrows())):
                    row_cells = table.rows[row_num + 1].cells
                    row_cells[0].text = str(row_num + 1)
                    row_cells[1].text = str(ev_row[0])
                    row_cells[2].text = str(v_row[0])

                # 添加解释说明
                doc.add_heading('解释说明', level=1)
                for key, value in languages[self.current_language]["explanation"].items():
                    doc.add_paragraph(f'{key}: {value}')

                # 添加结果解读
                doc.add_heading('结果解读', level=1)
                for key, value in languages[self.current_language]["interpretation"].items():
                    doc.add_paragraph(f'{key}: {value}')

                # 生成碎石图
                img_path = self.plot_scree_plot(ev, save_path)

                # 添加碎石图到文档
                doc.add_heading('碎石图', level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = PrincipalComponentAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()