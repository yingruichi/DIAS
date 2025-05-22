import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from sklearn.manifold import MDS
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "多维尺度分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'explanation': {
            "多维尺度分析": "多维尺度分析（MDS）是一种将多维空间中的对象之间的相似性或距离信息可视化的技术。",
        },
        'interpretation': {
            "多维尺度分析": "在多维尺度分析图中，距离较近的点表示对象之间的相似性较高，距离较远的点表示对象之间的相似性较低。",
        }
    },
    'en': {
        'title': "Multidimensional Scaling Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'explanation': {
            "Multidimensional Scaling Analysis": "Multidimensional Scaling (MDS) is a technique for visualizing the similarity or distance information between objects in a multi-dimensional space.",
        },
        'interpretation': {
            "Multidimensional Scaling Analysis": "In the MDS plot, points that are closer together indicate higher similarity between objects, while points that are farther apart indicate lower similarity between objects.",
        }
    }
}

class MultidimensionalScalingAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()
        
    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def mds_analysis(self, data):
        try:
            # 进行多维尺度分析
            mds = MDS(n_components=2, random_state=42)
            mds_result = mds.fit_transform(data)
            return mds_result
        except Exception as e:
            print(f"多维尺度分析出错: {e}")
            return None

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)
            # 将特征名称转换为字符串类型
            df.columns = df.columns.astype(str)

            # 进行多维尺度分析
            mds_result = self.mds_analysis(df)
            if mds_result is None:
                raise ValueError("多维尺度分析失败")

            # 整理数据
            all_data = []
            for j, point in enumerate(mds_result):
                all_data.append([f"对象{j + 1}", point[0], point[1]])
            headers = ["对象", "维度1", "维度2"]
            df_result = pd.DataFrame(all_data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["多维尺度分析" if self.current_language == 'zh' else "Multidimensional Scaling Analysis"])
            explanation_df.insert(0, "指标_解释说明",
                              "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["多维尺度分析" if self.current_language == 'zh' else "Multidimensional Scaling Analysis"])
            interpretation_df.insert(0, "指标_结果解读",
                                 "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 生成多维尺度分析图
            plt.figure(figsize=(10, 8))
            plt.scatter(mds_result[:, 0], mds_result[:, 1])
            for j, point in enumerate(mds_result):
                plt.annotate(f"对象{j + 1}", (point[0], point[1]))
            plt.title('多维尺度分析图' if self.current_language == 'zh' else 'Multidimensional Scaling Plot')
            plt.xlabel('维度1' if self.current_language == 'zh' else 'Dimension 1')
            plt.ylabel('维度2' if self.current_language == 'zh' else 'Dimension 2')

            # 保存图片
            image_path = os.path.splitext(file_path)[0] + '_mds_plot.png'
            plt.savefig(image_path)
            plt.close()

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('多维尺度分析结果' if self.current_language == 'zh' else 'Multidimensional Scaling Analysis Results', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(df_result.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df_result.columns):
                    hdr_cells[i].text = col

                # 添加数据到表格
                for index, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明表格
                doc.add_paragraph()
                table = doc.add_table(rows=1, cols=len(explanation_df.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(explanation_df.columns):
                    hdr_cells[i].text = col
                for index, row in explanation_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加结果解读表格
                doc.add_paragraph()
                table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(interpretation_df.columns):
                    hdr_cells[i].text = col
                for index, row in interpretation_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加图片
                doc.add_picture(image_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 设置 wraplength 属性让文本自动换行
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        
    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"], 
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = MultidimensionalScalingAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()