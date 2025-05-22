import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import tkinter.simpledialog  # 新增导入
import matplotlib.pyplot as plt
from statsmodels.formula.api import ols
# 修改导入部分
from statsmodels.stats import contrast
from docx import Document  # 导入 python-docx 库

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "联合分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'explanation': {
            "属性效应": "各属性对消费者偏好的影响程度。",
            "属性水平效应": "各属性不同水平对消费者偏好的影响程度。",
            "R-squared": "模型的拟合优度，值越接近1表示模型拟合效果越好。"
        },
        'interpretation': {
            "属性效应": "效应值越大，说明该属性对消费者偏好的影响越大。",
            "属性水平效应": "效应值越大，说明该属性水平越受消费者偏好。",
            "R-squared": "R-squared值接近1，说明模型能很好地解释消费者的偏好。"
        }
    },
    'en': {
        'title': "Conjoint Analysis",
        'select_button': "Select Files",
        'analyze_button': "Analyze Files",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'explanation': {
            "属性效应": "The influence of each attribute on consumer preferences.",
            "属性水平效应": "The influence of different levels of each attribute on consumer preferences.",
            "R-squared": "The goodness of fit of the model. A value closer to 1 indicates a better fit."
        },
        'interpretation': {
            "属性效应": "A larger effect value indicates a greater influence of the attribute on consumer preferences.",
            "属性水平效应": "A larger effect value indicates that the attribute level is more preferred by consumers.",
            "R-squared": "An R-squared value close to 1 indicates that the model can well explain consumer preferences."
        }
    }
}

# 当前语言
current_language = 'en'


def select_file():
    file_paths = filedialog.askopenfilenames(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_paths:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, ", ".join(file_paths))
        file_entry.configure(style="TEntry")  # 恢复默认样式


def conjoint_analysis(data, attribute_columns, preference_column):
    formula = f'{preference_column} ~ ' + ' + '.join(attribute_columns)
    model = ols(formula, data=data).fit()

    # 属性效应
    attribute_effects = model.params.drop('Intercept')
    # R-squared
    r_squared = model.rsquared

    all_results = {
        "属性效应": attribute_effects,
        "R-squared": r_squared
    }

    # 属性水平效应
    attribute_level_effects = {}
    for attr in attribute_columns:
        # 修改为直接使用导入的 contrast 模块
        contrast_results = contrast.ContrastResults(model.t_test_pairwise(attr))
        attribute_level_effects[attr] = contrast_results.effect

    all_results["属性水平效应"] = attribute_level_effects

    return all_results


def analyze_file():
    global current_language
    file_paths = file_entry.get().split(", ")
    if not file_paths or file_paths[0] == "请输入待分析 Excel 文件的完整路径" or file_paths[0] == "Please enter the full path of the Excel file to be analyzed":
        file_paths = []
    for file_path in file_paths:
        if not os.path.exists(file_path):
            result_label.config(text=LANGUAGES[current_language]['file_not_found'])
            return
    try:
        preference_column = tkinter.simpledialog.askstring("输入信息", "请输入偏好列的列名")
        if not preference_column:
            result_label.config(text="未输入有效的偏好列名，分析取消。")
            return
        attribute_columns = []
        while True:
            attribute_column = tkinter.simpledialog.askstring("输入信息", "请输入属性列的列名（点击取消结束输入）")
            if attribute_column is None:
                break
            if attribute_column.strip():
                attribute_columns.append(attribute_column.strip())
            else:
                print("输入的列名不能为空，请重新输入。")

        if not attribute_columns:
            result_label.config(text="未输入有效的属性列名，分析取消。")
            return

        all_results = []
        file_names = []
        for file_path in file_paths:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 进行联合分析
            conjoint_results = conjoint_analysis(df, attribute_columns, preference_column)
            all_results.append(conjoint_results)
            file_names.append(os.path.basename(file_path))

        # 整理数据
        all_data = []
        for i, results in enumerate(all_results):
            # 属性效应
            for attr, effect in results["属性效应"].items():
                all_data.append([f"{file_names[i]}_{attr}_属性效应", effect])
            # R-squared
            all_data.append([f"{file_names[i]}_R-squared", results["R-squared"]])
            # 属性水平效应
            for attr, levels in results["属性水平效应"].items():
                for level, effect in levels.items():
                    all_data.append([f"{file_names[i]}_{attr}_{level}_属性水平效应", effect])

        headers = ["指标", "数值"]
        df_result = pd.DataFrame(all_data, columns=headers)

        # 添加解释说明
        explanations = LANGUAGES[current_language]['explanation']
        interpretations = LANGUAGES[current_language]['interpretation']
        explanation_df = pd.DataFrame([explanations])
        explanation_df = explanation_df.reindex(
            columns=["属性效应", "属性水平效应", "R-squared"])
        explanation_df.insert(0, "指标_解释说明",
                              "解释说明" if current_language == 'zh' else "Explanation")

        # 添加分析结果解读
        interpretation_df = pd.DataFrame([interpretations])
        interpretation_df = interpretation_df.reindex(
            columns=["属性效应", "属性水平效应", "R-squared"])
        interpretation_df.insert(0, "指标_结果解读",
                                 "结果解读" if current_language == 'zh' else "Interpretation")

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建 Word 文档
            doc = Document()

            # 添加分析结果表格
            table = doc.add_table(rows=1, cols=len(headers))
            hdr_cells = table.rows[0].cells
            for col_idx, header in enumerate(headers):
                hdr_cells[col_idx].text = header
            for row in df_result.values.tolist():
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 添加解释说明表格
            doc.add_paragraph()
            doc.add_heading("解释说明" if current_language == 'zh' else "Explanation", level=2)
            exp_table = doc.add_table(rows=1, cols=len(explanation_df.columns))
            exp_hdr_cells = exp_table.rows[0].cells
            for col_idx, header in enumerate(explanation_df.columns):
                exp_hdr_cells[col_idx].text = header
            for row in explanation_df.values.tolist():
                exp_row_cells = exp_table.add_row().cells
                for col_idx, value in enumerate(row):
                    exp_row_cells[col_idx].text = str(value)

            # 添加结果解读表格
            doc.add_paragraph()
            doc.add_heading("结果解读" if current_language == 'zh' else "Interpretation", level=2)
            interp_table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
            interp_hdr_cells = interp_table.rows[0].cells
            for col_idx, header in enumerate(interpretation_df.columns):
                interp_hdr_cells[col_idx].text = header
            for row in interpretation_df.values.tolist():
                interp_row_cells = interp_table.add_row().cells
                for col_idx, value in enumerate(row):
                    interp_row_cells[col_idx].text = str(value)

            # 保存 Word 文档
            doc.save(save_path)

            result_msg = LANGUAGES[current_language]['analysis_success'].format(
                save_path)
            result_label.config(text=result_msg, wraplength=400)

            # 生成属性效应柱状图
            attribute_effects = [result["属性效应"] for result in all_results]
            all_attributes = []
            all_effects = []
            for effects in attribute_effects:
                for attr, effect in effects.items():
                    all_attributes.append(attr)
                    all_effects.append(effect)

            fig, ax = plt.subplots()
            ax.bar(all_attributes, all_effects)
            ax.set_title('属性效应' if current_language == 'zh' else 'Attribute Effects')
            ax.set_ylabel('效应值' if current_language == 'zh' else 'Effect Value')
            ax.set_xlabel('属性' if current_language == 'zh' else 'Attribute')
            plt.xticks(rotation=45)

            # 保存图片
            img_path = os.path.splitext(save_path)[0] + '_attribute_effects.png'
            plt.savefig(img_path)
            plt.close()

        else:
            result_label.config(text=LANGUAGES[current_language]['no_save_path'])

    except Exception as e:
        result_label.config(text=LANGUAGES[current_language]['analysis_error'].format(str(e)))


def switch_language():
    global current_language
    current_language = 'en' if current_language == 'zh' else 'zh'
    root.title(LANGUAGES[current_language]['title'])
    select_button.config(text=LANGUAGES[current_language]['select_button'])
    analyze_button.config(text=LANGUAGES[current_language]['analyze_button'])
    language_label.config(text=LANGUAGES[current_language]['switch_language'])
    # 切换语言时更新提示信息
    file_entry.delete(0, tk.END)
    if current_language == 'zh':
        file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
        file_entry.configure(style="Gray.TEntry")
    else:
        file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
        file_entry.configure(style="Gray.TEntry")


def on_entry_click(event):
    """当用户点击输入框时，清除提示信息"""
    if file_entry.get() == "请输入待分析 Excel 文件的完整路径" or file_entry.get() == "Please enter the full path of the Excel file to be analyzed":
        file_entry.delete(0, tk.END)
        file_entry.configure(style="TEntry")  # 恢复默认样式


def on_focusout(event):
    """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
    if file_entry.get() == "":
        if current_language == 'zh':
            file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
            file_entry.configure(style="Gray.TEntry")
        else:
            file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
            file_entry.configure(style="Gray.TEntry")


# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(LANGUAGES[current_language]['title'])

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 设置窗口的宽度和高度
window_width = 500
window_height = 300

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建自定义样式
style = ttk.Style()
style.configure("Gray.TEntry", foreground="gray")

# 创建文件选择按钮
select_button = ttk.Button(root, text=LANGUAGES[current_language]['select_button'], command=select_file,
                           bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建文件路径输入框
file_entry = ttk.Entry(root, width=50, style="Gray.TEntry")
if current_language == 'zh':
    file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
else:
    file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
file_entry.pack(pady=5)
file_entry.bind("<FocusIn>", on_entry_click)
file_entry.bind("<FocusOut>", on_focusout)

# 创建分析按钮
analyze_button = ttk.Button(root, text=LANGUAGES[current_language]['analyze_button'], command=analyze_file,
                            bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建语言切换标签
language_label = ttk.Label(root, text=LANGUAGES[current_language]['switch_language'], cursor="hand2")
language_label.pack(pady=10)
language_label.bind("<Button-1>", lambda event: switch_language())

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
root.mainloop()