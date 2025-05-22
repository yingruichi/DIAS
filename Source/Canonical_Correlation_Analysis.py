import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from sklearn.cross_decomposition import CCA
from docx import Document

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    'zh': {
        'title': "典型相关分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}",
        'images_saved': "结果图片已保存到 {}",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'explanation': {
            "典型相关分析": "研究两组变量之间的相关性，找到两组变量的线性组合，使得它们之间的相关性最大。",
        },
        'interpretation': {
            "典型相关系数": "反映两组变量的线性组合之间的相关性，取值范围为 -1 到 1，绝对值越接近 1 表示相关性越强。",
            "典型变量": "两组变量的线性组合，用于揭示两组变量之间的潜在关系。",
        }
    },
    'en': {
        'title': "Canonical Correlation Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}",
        'images_saved': "The result image has been saved to {}",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'explanation': {
            "Canonical Correlation Analysis": "Study the correlation between two sets of variables and find the linear combinations of the two sets of variables that maximize the correlation between them.",
        },
        'interpretation': {
            "Canonical correlation coefficient": "Reflects the correlation between the linear combinations of two sets of variables, ranging from -1 to 1. The closer the absolute value is to 1, the stronger the correlation.",
            "Canonical variables": "Linear combinations of two sets of variables used to reveal the potential relationship between the two sets of variables.",
        }
    }
}

class CanonicalCorrelationAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()
        
    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设前半部分列是第一组变量，后半部分列是第二组变量
            mid = len(df.columns) // 2
            X = df.iloc[:, :mid]
            Y = df.iloc[:, mid:]

            # 进行典型相关分析
            cca = CCA()
            cca.fit(X, Y)
            X_c, Y_c = cca.transform(X, Y)

            # 计算典型相关系数
            canonical_correlations = []
            for i in range(min(X_c.shape[1], Y_c.shape[1])):
                corr = np.corrcoef(X_c[:, i], Y_c[:, i])[0, 1]
                canonical_correlations.append(corr)

            # 整理结果
            canonical_corr_df = pd.DataFrame({
                '典型相关系数' if self.current_language == 'zh' else 'Canonical correlation coefficient': canonical_correlations
            })

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["典型相关分析" if self.current_language == 'zh' else "Canonical Correlation Analysis"])
            explanation_df.insert(0, "统计量", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=[
                "典型相关系数" if self.current_language == 'zh' else "Canonical correlation coefficient",
                "典型变量" if self.current_language == 'zh' else "Canonical variables"
            ])
            interpretation_df.insert(0, "统计量", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加典型相关系数表格
                table = doc.add_table(rows=1, cols=len(canonical_corr_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(canonical_corr_df.columns):
                    hdr_cells[col_idx].text = col_name
                for _, row in canonical_corr_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明表格
                doc.add_paragraph()
                table = doc.add_table(rows=1, cols=len(explanation_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(explanation_df.columns):
                    hdr_cells[col_idx].text = col_name
                for _, row in explanation_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加结果解读表格
                doc.add_paragraph()
                table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(interpretation_df.columns):
                    hdr_cells[col_idx].text = col_name
                for _, row in interpretation_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 保存 Word 文档
                doc.save(save_path)

                # 生成结果图片
                desktop_path = pathlib.Path.home() / 'Desktop'
                plot_path = desktop_path / 'canonical_corr_plot.png'
                plt.figure()
                plt.scatter(X_c[:, 0], Y_c[:, 0])
                plt.xlabel(
                    '第一组典型变量第一维' if self.current_language == 'zh' else 'First dimension of the first set of canonical variables')
                plt.ylabel(
                    '第二组典型变量第一维' if self.current_language == 'zh' else 'First dimension of the second set of canonical variables')
                plt.title('典型相关分析结果' if self.current_language == 'zh' else 'Canonical Correlation Analysis Results')
                plt.savefig(plot_path)
                plt.close()

                result_msg = languages[self.current_language]['analysis_complete'].format(save_path)
                result_msg += "\n" + languages[self.current_language]['images_saved'].format(plot_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        
    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"], 
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = CanonicalCorrelationAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()