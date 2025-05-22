import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "灰色关联分析法分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "参考序列": "用于比较的基准序列",
            "比较序列": "待与参考序列进行比较的序列",
            "关联系数矩阵": "反映各比较序列与参考序列在各个时刻的关联程度的矩阵",
            "关联度": "各比较序列与参考序列的整体关联程度",
            "关联度排序结果": "根据关联度对各比较序列进行排序的结果"
        },
        'interpretation': {
            "参考序列": "作为衡量其他序列关联程度的标准",
            "比较序列": "需要分析与参考序列关联程度的序列",
            "关联系数矩阵": "数值越大，该时刻比较序列与参考序列的关联程度越高",
            "关联度": "值越大，说明比较序列与参考序列的整体关联程度越高",
            "关联度排序结果": "排名越靠前，与参考序列的关联程度越高"
        }
    },
    'en': {
        'title': "Grey Relational Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "参考序列": "The reference sequence for comparison",
            "比较序列": "The sequences to be compared with the reference sequence",
            "关联系数矩阵": "A matrix reflecting the degree of association between each comparison sequence and the reference sequence at each time point",
            "关联度": "The overall degree of association between each comparison sequence and the reference sequence",
            "关联度排序结果": "The result of ranking each comparison sequence according to the degree of association"
        },
        'interpretation': {
            "参考序列": "As a standard for measuring the degree of association of other sequences",
            "比较序列": "Sequences whose degree of association with the reference sequence needs to be analyzed",
            "关联系数矩阵": "The larger the value, the higher the degree of association between the comparison sequence and the reference sequence at that time point",
            "关联度": "The larger the value, the higher the overall degree of association between the comparison sequence and the reference sequence",
            "关联度排序结果": "The higher the ranking, the higher the degree of association with the reference sequence"
        }
    }
}

class GreyRelationalAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.configure(style="TEntry")

    def on_entry_click(self, event):
        """当用户点击输入框时，清除提示信息"""
        if self.file_entry.get() == LANGUAGES[self.current_language]['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.configure(style="TEntry")

    def on_focusout(self, event):
        """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def grey_relational_analysis(self, reference_sequence, comparison_sequences):
        """
        实现灰色关联分析法
        :param reference_sequence: 参考序列
        :param comparison_sequences: 比较序列
        :return: 关联系数矩阵, 关联度, 关联度排序结果
        """
        # 数据预处理，这里采用初值化处理
        reference_sequence = reference_sequence / reference_sequence[0]
        comparison_sequences = comparison_sequences / comparison_sequences[:, 0].reshape(-1, 1)

        # 计算差序列
        diff_matrix = np.abs(comparison_sequences - reference_sequence)

        # 计算两级最小差和两级最大差
        min_min_diff = np.min(np.min(diff_matrix))
        max_max_diff = np.max(np.max(diff_matrix))

        # 分辨系数
        rho = 0.5

        # 计算关联系数矩阵
        relational_coefficient_matrix = (min_min_diff + rho * max_max_diff) / (diff_matrix + rho * max_max_diff)

        # 计算关联度
        relational_degree = np.mean(relational_coefficient_matrix, axis=1)

        # 对关联度进行排序
        ranking = np.argsort(-relational_degree) + 1

        return relational_coefficient_matrix, relational_degree, ranking

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]['file_entry_placeholder']:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 将数据转换为浮点类型
            data = data.astype(float)

            # 假设第一行为参考序列，其余行为比较序列
            reference_sequence = data[0]
            comparison_sequences = data[1:]

            # 进行灰色关联分析
            relational_coefficient_matrix, relational_degree, ranking = self.grey_relational_analysis(reference_sequence,
                                                                                                      comparison_sequences)

            # 整理数据
            data = [
                ["参考序列", reference_sequence.tolist(), ""],
                ["比较序列", comparison_sequences.tolist(), ""],
                ["关联系数矩阵", relational_coefficient_matrix.tolist(), ""],
                ["关联度", relational_degree.tolist(), ""],
                ["关联度排序结果", ranking.tolist(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["参考序列", "比较序列", "关联系数矩阵", "关联度", "关联度排序结果"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["参考序列", "比较序列", "关联系数矩阵", "关联度", "关联度排序结果"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                document = Document()

                # 添加标题
                document.add_heading('灰色关联分析结果', 0)

                # 添加分析结果表格
                table = document.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header

                for row in combined_df.values.tolist():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 生成关联度柱状图
                fig, ax = plt.subplots()
                ax.bar(range(len(relational_degree)), relational_degree)
                ax.set_title(
                    '关联度柱状图' if self.current_language == 'zh' else 'Bar Chart of Relational Degree')
                ax.set_xlabel('比较序列编号' if self.current_language == 'zh' else 'Comparison Sequence Number')
                ax.set_ylabel('关联度' if self.current_language == 'zh' else 'Relational Degree')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_relational_degree.png'
                plt.savefig(img_path)
                plt.close()

                # 在 Word 文档中插入图片
                document.add_heading('关联度柱状图', level=1)
                document.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                document.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.configure(style="Gray.TEntry")

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建文件选择按钮
        self.select_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file,
                                        bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(self.root, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack(pady=5)
        self.file_entry.bind("<FocusIn>", self.on_entry_click)
        self.file_entry.bind("<FocusOut>", self.on_focusout)

        # 创建分析按钮
        self.analyze_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file,
                                         bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建语言切换标签
        self.switch_language_label = ttk.Label(self.root, text=LANGUAGES[self.current_language]['switch_language'],
                                               cursor="hand2")
        self.switch_language_label.pack(pady=10)
        self.switch_language_label.bind("<Button-1>", self.switch_language)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = GreyRelationalAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()