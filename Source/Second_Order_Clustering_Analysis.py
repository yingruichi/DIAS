import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from sklearn.cluster import AgglomerativeClustering
from scipy.cluster.hierarchy import dendrogram

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "二阶聚类分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}\n",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "switch_language_button_text": "切换语言",
        "explanation": {
            "聚类结果": "每个样本所属的聚类类别",
            "聚类树状图": "展示样本之间的二阶聚类层次关系"
        },
        "interpretation": {
            "聚类结果": "可用于区分不同样本所属的类别",
            "聚类树状图": "直观展示样本之间的二阶聚类层次结构"
        }
    },
    "en": {
        "title": "Second-Order Clustering Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}\n",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "switch_language_button_text": "Switch Language",
        "explanation": {
            "聚类结果": "The cluster label to which each sample belongs",
            "聚类树状图": "Show the second-order hierarchical clustering relationship between samples"
        },
        "interpretation": {
            "聚类结果": "Can be used to distinguish the categories to which different samples belong",
            "聚类树状图": "Visually show the second-order hierarchical clustering structure of samples"
        }
    }
}

class SecondOrderClusteringAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def second_order_clustering(self, data, n_clusters=3):
        """
        进行二阶聚类分析
        :param data: 输入数据
        :param n_clusters: 聚类的数量
        :return: 聚类标签
        """
        model = AgglomerativeClustering(n_clusters=n_clusters, linkage='ward')
        labels = model.fit_predict(data)
        return labels, model

    def plot_dendrogram(self, model, **kwargs):
        # 生成树状图所需的链接矩阵
        counts = np.zeros(model.children_.shape[0])
        n_samples = len(model.labels_)
        for i, merge in enumerate(model.children_):
            current_count = 0
            for child_idx in merge:
                if child_idx < n_samples:
                    current_count += 1
                else:
                    current_count += counts[child_idx - n_samples]
            counts[i] = current_count

        linkage_matrix = np.column_stack([model.children_, model.distances_,
                                          counts]).astype(float)

        # 绘制树状图
        dendrogram(linkage_matrix, **kwargs)

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 进行二阶聚类分析
            labels, model = self.second_order_clustering(data)

            # 整理数据
            data = [
                ["聚类结果", labels.tolist(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["聚类结果", "聚类树状图"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["聚类结果", "聚类树状图"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加分析结果表格
                doc.add_heading('Analysis Results', level=1)
                table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明表格
                doc.add_heading('Explanation', level=1)
                table = doc.add_table(rows=explanation_df.shape[0] + 1, cols=explanation_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(explanation_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in explanation_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加结果解读表格
                doc.add_heading('Interpretation', level=1)
                table = doc.add_table(rows=interpretation_df.shape[0] + 1, cols=interpretation_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(interpretation_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in interpretation_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 生成聚类树状图
                plt.figure(figsize=(10, 5))
                self.plot_dendrogram(model, labels=range(data.shape[0]))
                plt.title('聚类树状图' if self.current_language == 'zh' else 'Second-Order Clustering Dendrogram')
                plt.xlabel('样本编号' if self.current_language == 'zh' else 'Sample Index')
                plt.ylabel('距离' if self.current_language == 'zh' else 'Distance')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_second_order_clustering_dendrogram.png'
                plt.savefig(img_path)
                plt.close()

                # 在 Word 文档中插入树状图
                doc.add_heading('Second-Order Clustering Dendrogram', level=1)
                doc.add_picture(img_path)

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_complete'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = SecondOrderClusteringAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()