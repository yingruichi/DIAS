import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "单样本Wilcoxon检验分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "单样本Wilcoxon检验": "用于检验样本中位数是否与给定的假设中位数存在显著差异。",
            "样本量": "样本中的观测值数量。",
            "中位数": "样本数据的中间值，将数据分为上下两部分。",
            "t统计量": "单样本Wilcoxon检验的统计量值，用于衡量样本与假设中位数的差异程度。",
            "自由度": "在统计计算中能够自由变化的变量个数。",
            "p值": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为样本中位数与假设中位数存在显著差异；否则，接受原假设，认为无显著差异。",
            "均值差异的置信区间": "包含真实均值差异的一个区间，反映了估计的不确定性。"
        },
        'interpretation': {
            "t统计量": "t统计量的绝对值越大，说明样本与假设中位数的差异越显著。",
            "自由度": "自由度影响t分布的形状，进而影响p值的计算。",
            "p值": "用于判断样本与假设中位数之间是否存在显著差异的依据。",
            "样本量": "样本量的大小会影响统计检验的功效，较大的样本量通常能提供更准确的结果。",
            "中位数": "中位数反映了数据的中心位置，可用于比较样本与假设中位数的差异。",
            "均值差异的置信区间": "如果置信区间不包含0，说明样本与假设中位数存在显著差异。"
        }
    },
    'en': {
        'title': "One-Sample Wilcoxon Test Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "单样本Wilcoxon检验": "Used to test whether the median of a sample is significantly different from a given hypothesized median.",
            "样本量": "The number of observations in the sample.",
            "中位数": "The middle value of the sample data, dividing the data into two parts.",
            "t统计量": "The test statistic value of the one-sample Wilcoxon test, used to measure the degree of difference between the sample and the hypothesized median.",
            "自由度": "The number of independent variables in a statistical calculation.",
            "p值": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between the sample median and the hypothesized median; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "均值差异的置信区间": "An interval that contains the true mean difference, reflecting the uncertainty of the estimate."
        },
        'interpretation': {
            "t统计量": "The larger the absolute value of the t-statistic, the more significant the difference between the sample and the hypothesized median.",
            "自由度": "The degrees of freedom affect the shape of the t-distribution, which in turn affects the calculation of the p-value.",
            "p值": "The basis for determining whether there is a significant difference between the sample and the hypothesized median.",
            "样本量": "The sample size affects the power of the statistical test. A larger sample size usually provides more accurate results.",
            "中位数": "The median reflects the central position of the data and can be used to compare the difference between the sample and the hypothesized median.",
            "均值差异的置信区间": "If the confidence interval does not contain 0, it indicates a significant difference between the sample and the hypothesized median."
        }
    }
}

class OneSampleWilcoxonTestAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行单样本Wilcoxon检验。")

            # 假设中位数为 0
            hypothesized_median = 0

            # 进行单样本Wilcoxon检验
            t_stat, p_value = stats.wilcoxon(numerical_df.squeeze() - hypothesized_median)

            # 计算样本量、中位数
            sample_size = numerical_df.count().values[0]
            median = numerical_df.median().values[0]

            # 计算自由度
            degrees_of_freedom = sample_size - 1

            # 计算均值差异的置信区间
            mean_diff = numerical_df.squeeze().mean() - hypothesized_median
            std_err = numerical_df.squeeze().std() / np.sqrt(sample_size)
            confidence_interval = stats.t.interval(0.95, degrees_of_freedom, loc=mean_diff, scale=std_err)

            # 整理数据
            data = [
                ["单样本Wilcoxon检验", t_stat, degrees_of_freedom, p_value, confidence_interval],
                ["样本量", sample_size, "", "", ""],
                ["中位数", median, "", "", ""]
            ]
            headers = ["统计量", "t统计量", "自由度", "p值", "均值差异的置信区间"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["单样本Wilcoxon检验", "样本量", "中位数", "t统计量", "自由度", "p值", "均值差异的置信区间"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["t统计量", "自由度", "p值", "样本量", "中位数", "均值差异的置信区间"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表格数据
                table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header
                for row_idx in range(df.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx in range(df.shape[1]):
                        row_cells[col_idx].text = str(df.iloc[row_idx, col_idx])

                # 添加解释说明
                doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", level=2)
                for stat, explanation in explanations.items():
                    doc.add_paragraph(f"{stat}: {explanation}")

                # 添加分析结果解读
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Interpretation", level=2)
                for stat, interpretation in interpretations.items():
                    doc.add_paragraph(f"{stat}: {interpretation}")

                # 绘制柱状图
                plt.figure(figsize=(8, 6))
                plt.bar(['样本'], [numerical_df.squeeze().mean()])
                plt.title('柱状图' if self.current_language == 'zh' else 'Bar Chart')
                plt.ylabel('均值' if self.current_language == 'zh' else 'Mean')
                bar_chart_path = save_path.replace('.docx', '_bar_chart.png')
                plt.savefig(bar_chart_path)
                plt.close()

                # 绘制误差线图
                plt.figure(figsize=(8, 6))
                plt.errorbar(['样本'], [numerical_df.squeeze().mean()], yerr=[std_err], fmt='o')
                plt.title('误差线图' if self.current_language == 'zh' else 'Error Bar Chart')
                plt.ylabel('均值' if self.current_language == 'zh' else 'Mean')
                error_bar_chart_path = save_path.replace('.docx', '_error_bar_chart.png')
                plt.savefig(error_bar_chart_path)
                plt.close()

                # 绘制箱线图
                plt.figure(figsize=(8, 6))
                numerical_df.squeeze().plot(kind='box')
                plt.title('箱线图' if self.current_language == 'zh' else 'Box Plot')
                plt.ylabel('数值' if self.current_language == 'zh' else 'Value')
                box_plot_path = save_path.replace('.docx', '_box_plot.png')
                plt.savefig(box_plot_path)
                plt.close()

                # 绘制折线图
                plt.figure(figsize=(8, 6))
                numerical_df.squeeze().plot(kind='line')
                plt.title('折线图' if self.current_language == 'zh' else 'Line Chart')
                plt.ylabel('数值' if self.current_language == 'zh' else 'Value')
                line_chart_path = save_path.replace('.docx', '_line_chart.png')
                plt.savefig(line_chart_path)
                plt.close()

                # 将图表插入 Word 文档
                doc.add_heading("图表" if self.current_language == 'zh' else "Charts", level=2)
                doc.add_picture(bar_chart_path, width=Inches(6))
                doc.add_picture(error_bar_chart_path, width=Inches(6))
                doc.add_picture(box_plot_path, width=Inches(6))
                doc.add_picture(line_chart_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = OneSampleWilcoxonTestAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()