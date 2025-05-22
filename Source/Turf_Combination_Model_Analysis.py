import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    'zh': {
        'title': "Turf组合模型分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'explanation': {
            "Turf组合得分": "Turf组合模型中每个组合的得分，反映了该组合的吸引力。",
            "最优组合": "得分最高的组合，代表了最有吸引力的产品或服务组合。",
            "样本量": "每个样本中的观测值数量。",
            "均值": "样本数据的平均值。"
        },
        'interpretation': {
            "Turf组合得分": "Turf组合得分越高，说明该组合越受用户欢迎。",
            "最优组合": "最优组合是最能满足用户需求的组合，可作为产品或服务的推荐组合。",
            "样本量": "样本量的大小会影响统计检验的稳定性，较大的样本量通常能提供更可靠的结果。",
            "均值": "均值反映了数据的平均水平，可用于比较不同变量的集中趋势。"
        }
    },
    'en': {
        'title': "Turf Combination Model Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'explanation': {
            "Turf组合得分": "The score of each combination in the Turf combination model, reflecting the attractiveness of the combination.",
            "最优组合": "The combination with the highest score, representing the most attractive product or service combination.",
            "样本量": "The number of observations in each sample.",
            "均值": "The average value of the sample data."
        },
        'interpretation': {
            "Turf组合得分": "The higher the Turf combination score, the more popular the combination is among users.",
            "最优组合": "The optimal combination is the one that best meets the needs of users and can be recommended as a product or service combination.",
            "样本量": "The sample size affects the stability of the statistical test. A larger sample size usually provides more reliable results.",
            "均值": "The mean reflects the average level of the data and can be used to compare the central tendencies of different variables."
        }
    }
}

class TurfCombinationModelAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def turf_analysis(self, data):
        # 这里简单模拟Turf组合模型分析，实际应用中需要根据具体需求实现
        # 假设数据的每一行代表一个样本，每一列代表一个产品或服务
        # 计算每个组合的得分
        combination_scores = data.sum(axis=0)
        # 找到最优组合
        optimal_combination = combination_scores.idxmax()

        return combination_scores, optimal_combination

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行Turf组合模型分析。")

            # 进行Turf组合模型分析
            combination_scores, optimal_combination = self.turf_analysis(numerical_df)

            # 计算样本量和均值
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()

            # 整理数据
            data = [
                ["Turf组合得分", combination_scores.to_dict(), ""],
                ["最优组合", optimal_combination, ""],
                ["样本量", sample_sizes.to_dict(), ""],
                ["均值", means.to_dict(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["Turf组合得分", "最优组合", "样本量", "均值"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["Turf组合得分", "最优组合", "样本量", "均值"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df_result, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表格
                table = doc.add_table(rows=combined_df.shape[0] + 1, cols=combined_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(combined_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx in range(combined_df.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(combined_df.iloc[row_idx]):
                        row_cells[col_idx].text = str(value)

                # 生成图片（Turf组合得分柱状图）
                fig, ax = plt.subplots()
                combination_scores.plot(kind='bar', ax=ax)
                ax.set_title('Turf组合得分柱状图' if self.current_language == 'zh' else 'Bar Chart of Turf Combination Scores')
                ax.set_xlabel('组合' if self.current_language == 'zh' else 'Combinations')
                ax.set_ylabel('得分' if self.current_language == 'zh' else 'Scores')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = TurfCombinationModelAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()