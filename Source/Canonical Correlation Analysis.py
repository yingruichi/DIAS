import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from sklearn.cross_decomposition import CCA
from docx import Document

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "典型相关分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'explanation': {
            "典型相关分析": "研究两组变量之间的相关性，找到两组变量的线性组合，使得它们之间的相关性最大。",
        },
        'interpretation': {
            "典型相关系数": "反映两组变量的线性组合之间的相关性，取值范围为 -1 到 1，绝对值越接近 1 表示相关性越强。",
            "典型变量": "两组变量的线性组合，用于揭示两组变量之间的潜在关系。",
        }
    },
    'en': {
        'title': "Canonical Correlation Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'explanation': {
            "Canonical Correlation Analysis": "Study the correlation between two sets of variables and find the linear combinations of the two sets of variables that maximize the correlation between them.",
        },
        'interpretation': {
            "Canonical correlation coefficient": "Reflects the correlation between the linear combinations of two sets of variables, ranging from -1 to 1. The closer the absolute value is to 1, the stronger the correlation.",
            "Canonical variables": "Linear combinations of two sets of variables used to reveal the potential relationship between the two sets of variables.",
        }
    }
}

# 当前语言
current_language = 'en'


def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        file_entry.configure(style="TEntry")


def analyze_file():
    global current_language
    file_path = file_entry.get()
    if file_path == "请输入待分析 Excel 文件的完整路径" or file_path == "Please enter the full path of the Excel file to be analyzed":
        file_path = ""
    if not os.path.exists(file_path):
        result_label.config(text=LANGUAGES[current_language]['file_not_found'])
        return
    try:
        # 打开 Excel 文件
        df = pd.read_excel(file_path)

        # 假设前半部分列是第一组变量，后半部分列是第二组变量
        mid = len(df.columns) // 2
        X = df.iloc[:, :mid]
        Y = df.iloc[:, mid:]

        # 进行典型相关分析
        cca = CCA()
        cca.fit(X, Y)
        X_c, Y_c = cca.transform(X, Y)

        # 计算典型相关系数
        canonical_correlations = []
        for i in range(min(X_c.shape[1], Y_c.shape[1])):
            corr = np.corrcoef(X_c[:, i], Y_c[:, i])[0, 1]
            canonical_correlations.append(corr)

        # 整理结果
        canonical_corr_df = pd.DataFrame({
            '典型相关系数' if current_language == 'zh' else 'Canonical correlation coefficient': canonical_correlations
        })

        # 添加解释说明
        explanations = LANGUAGES[current_language]['explanation']
        interpretations = LANGUAGES[current_language]['interpretation']
        explanation_df = pd.DataFrame([explanations])
        explanation_df = explanation_df.reindex(
            columns=["典型相关分析" if current_language == 'zh' else "Canonical Correlation Analysis"])
        explanation_df.insert(0, "统计量", "解释说明" if current_language == 'zh' else "Explanation")

        # 添加分析结果解读
        interpretation_df = pd.DataFrame([interpretations])
        interpretation_df = interpretation_df.reindex(columns=[
            "典型相关系数" if current_language == 'zh' else "Canonical correlation coefficient",
            "典型变量" if current_language == 'zh' else "Canonical variables"
        ])
        interpretation_df.insert(0, "统计量", "结果解读" if current_language == 'zh' else "Interpretation")

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建一个新的 Word 文档
            doc = Document()

            # 添加典型相关系数表格
            table = doc.add_table(rows=1, cols=len(canonical_corr_df.columns))
            hdr_cells = table.rows[0].cells
            for col_idx, col_name in enumerate(canonical_corr_df.columns):
                hdr_cells[col_idx].text = col_name
            for _, row in canonical_corr_df.iterrows():
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 添加解释说明表格
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=len(explanation_df.columns))
            hdr_cells = table.rows[0].cells
            for col_idx, col_name in enumerate(explanation_df.columns):
                hdr_cells[col_idx].text = col_name
            for _, row in explanation_df.iterrows():
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 添加结果解读表格
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
            hdr_cells = table.rows[0].cells
            for col_idx, col_name in enumerate(interpretation_df.columns):
                hdr_cells[col_idx].text = col_name
            for _, row in interpretation_df.iterrows():
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 保存 Word 文档
            doc.save(save_path)

            # 生成结果图片
            desktop_path = pathlib.Path.home() / 'Desktop'
            plot_path = desktop_path / 'canonical_corr_plot.png'
            plt.figure()
            plt.scatter(X_c[:, 0], Y_c[:, 0])
            plt.xlabel(
                '第一组典型变量第一维' if current_language == 'zh' else 'First dimension of the first set of canonical variables')
            plt.ylabel(
                '第二组典型变量第一维' if current_language == 'zh' else 'First dimension of the second set of canonical variables')
            plt.title('典型相关分析结果' if current_language == 'zh' else 'Canonical Correlation Analysis Results')
            plt.savefig(plot_path)
            plt.close()

            result_msg = LANGUAGES[current_language]['analysis_success'].format(
                save_path) + f"结果图片已保存到 {plot_path}" if current_language == 'zh' else f"The result image has been saved to {plot_path}"
            result_label.config(text=result_msg, wraplength=400)
        else:
            result_label.config(text=LANGUAGES[current_language]['no_save_path'])

    except Exception as e:
        result_label.config(text=LANGUAGES[current_language]['analysis_error'].format(str(e)))


def switch_language():
    global current_language
    current_language = 'en' if current_language == 'zh' else 'zh'
    root.title(LANGUAGES[current_language]['title'])
    select_button.config(text=LANGUAGES[current_language]['select_button'])
    analyze_button.config(text=LANGUAGES[current_language]['analyze_button'])
    language_label.config(text=LANGUAGES[current_language]['switch_language'])
    # 切换语言时更新提示信息
    file_entry.delete(0, tk.END)
    if current_language == 'zh':
        file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
        file_entry.configure(style="Gray.TEntry")
    else:
        file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
        file_entry.configure(style="Gray.TEntry")


def on_entry_click(event):
    """当用户点击输入框时，清除提示信息"""
    if file_entry.get() == "请输入待分析 Excel 文件的完整路径" or file_entry.get() == "Please enter the full path of the Excel file to be analyzed":
        file_entry.delete(0, tk.END)
        file_entry.configure(style="TEntry")


def on_focusout(event):
    """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
    if file_entry.get() == "":
        if current_language == 'zh':
            file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
            file_entry.configure(style="Gray.TEntry")
        else:
            file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
            file_entry.configure(style="Gray.TEntry")


# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(LANGUAGES[current_language]['title'])

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 设置窗口的宽度和高度
window_width = 500
window_height = 300

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建自定义样式
style = ttk.Style()
style.configure("Gray.TEntry", foreground="gray")

# 创建文件选择按钮
select_button = ttk.Button(root, text=LANGUAGES[current_language]['select_button'], command=select_file,
                           bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建文件路径输入框
file_entry = ttk.Entry(root, width=50, style="Gray.TEntry")
if current_language == 'zh':
    file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
else:
    file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
file_entry.pack(pady=5)
file_entry.bind("<FocusIn>", on_entry_click)
file_entry.bind("<FocusOut>", on_focusout)

# 创建分析按钮
analyze_button = ttk.Button(root, text=LANGUAGES[current_language]['analyze_button'], command=analyze_file,
                            bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建语言切换标签
language_label = ttk.Label(root, text=LANGUAGES[current_language]['switch_language'], cursor="hand2")
language_label.pack(pady=10)
language_label.bind("<Button-1>", lambda event: switch_language())

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
try:
    root.mainloop()
except KeyboardInterrupt:
    print("程序被手动中断。")
    root.destroy()  # 销毁主窗口