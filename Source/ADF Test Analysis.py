import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
from statsmodels.tsa.stattools import adfuller
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    "zh": {
        "title": "ADF检验分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "columns_stats": ["变量名", "ADF检验统计量", "p值", "滞后阶数", "结果解读"],
        "interpretation_stationary": "p值小于 0.05，表明该时间序列是平稳的。",
        "interpretation_non_stationary": "p值大于等于 0.05，表明该时间序列是非平稳的。",
        "switch_language_button_text": "切换语言"
    },
    "en": {
        "title": "ADF Test Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "columns_stats": ["Variable Name", "ADF Test Statistic", "p-value", "Lags", "Result Interpretation"],
        "interpretation_stationary": "The p-value is less than 0.05, indicating that the time series is stationary.",
        "interpretation_non_stationary": "The p-value is greater than or equal to 0.05, indicating that the time series is non-stationary.",
        "switch_language_button_text": "Switch Language"
    }
}

# 当前语言，默认为英文
current_language = "en"


def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        file_entry.config(foreground='black')


def on_entry_click(event):
    if file_entry.get() == languages[current_language]["file_entry_placeholder"]:
        file_entry.delete(0, tk.END)
        file_entry.config(foreground='black')


def on_focusout(event):
    if file_entry.get() == "":
        file_entry.insert(0, languages[current_language]["file_entry_placeholder"])
        file_entry.config(foreground='gray')


# 计算 ADF 检验的函数
def calculate_adf(X):
    adf_data = pd.DataFrame()
    adf_data["Variable Name"] = X.columns
    adf_statistics = []
    p_values = []
    lags = []
    interpretations = []

    for col in X.columns:
        result = adfuller(X[col])
        adf_statistics.append(result[0])
        p_values.append(result[1])
        lags.append(result[2])
        if result[1] < 0.05:
            interpretations.append(languages[current_language]["interpretation_stationary"])
        else:
            interpretations.append(languages[current_language]["interpretation_non_stationary"])

    adf_data["ADF Test Statistic"] = adf_statistics
    adf_data["p-value"] = p_values
    adf_data["Lags"] = lags
    adf_data["Result Interpretation"] = interpretations
    return adf_data


def analyze_file():
    file_path = file_entry.get()
    if file_path == languages[current_language]["file_entry_placeholder"]:
        result_label.config(text=languages[current_language]["no_file_selected"])
        return
    if not os.path.exists(file_path):
        result_label.config(text=languages[current_language]["file_not_exists"])
        return
    try:
        # 打开 Excel 文件
        df = pd.read_excel(file_path)

        # 假设所有列都是时间序列数据
        X = df

        # 计算 ADF 检验
        adf_data = calculate_adf(X)

        # 绘制 p 值的柱状图
        plt.figure()
        plt.bar(adf_data["Variable Name"], adf_data["p-value"])
        plt.xlabel('Variable Name')
        plt.ylabel('p-value')
        plt.title('ADF Test p-values for Each Variable')
        plt.xticks(rotation=45)

        # 保存图片
        image_path = os.path.splitext(file_path)[0] + '_adf_plot.png'
        plt.savefig(image_path)
        plt.close()

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建一个新的 Word 文档
            doc = Document()

            # 添加标题
            doc.add_heading('ADF Test Results', 0)

            # 添加表格
            table = doc.add_table(rows=1, cols=len(adf_data.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(adf_data.columns):
                hdr_cells[i].text = col

            # 添加数据到表格
            for index, row in adf_data.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # 添加图片
            doc.add_picture(image_path, width=Inches(6))

            # 保存 Word 文档
            doc.save(save_path)

            # 设置 wraplength 属性让文本自动换行
            result_label.config(text=languages[current_language]["analysis_complete"].format(save_path), wraplength=400)
        else:
            result_label.config(text=languages[current_language]["no_save_path_selected"])

    except Exception as e:
        result_label.config(text=languages[current_language]["analysis_error"].format(str(e)))


def switch_language(event):
    global current_language
    if current_language == "zh":
        current_language = "en"
    else:
        current_language = "zh"

    # 更新界面文字
    root.title(languages[current_language]["title"])
    select_button.config(text=languages[current_language]["select_button_text"])
    file_entry.delete(0, tk.END)
    file_entry.insert(0, languages[current_language]["file_entry_placeholder"])
    file_entry.config(foreground='gray')
    analyze_button.config(text=languages[current_language]["analyze_button_text"])
    switch_language_label.config(text=languages[current_language]["switch_language_button_text"])


# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(languages[current_language]["title"])

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 设置窗口的宽度和高度
window_width = 500
window_height = 300

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建一个框架来包含按钮和输入框
frame = ttk.Frame(root)
frame.pack(expand=True)

# 创建文件选择按钮
select_button = ttk.Button(frame, text=languages[current_language]["select_button_text"], command=select_file,
                           bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建文件路径输入框
file_entry = ttk.Entry(frame, width=50)
file_entry.insert(0, languages[current_language]["file_entry_placeholder"])
file_entry.config(foreground='gray')
file_entry.bind('<FocusIn>', on_entry_click)
file_entry.bind('<FocusOut>', on_focusout)
file_entry.pack(pady=5)

# 创建分析按钮
analyze_button = ttk.Button(frame, text=languages[current_language]["analyze_button_text"], command=analyze_file,
                            bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建切换语言标签
switch_language_label = ttk.Label(frame, text=languages[current_language]["switch_language_button_text"],
                                  foreground="gray", cursor="hand2")
switch_language_label.bind("<Button-1>", switch_language)
switch_language_label.pack(pady=10)

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
root.mainloop()