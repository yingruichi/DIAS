import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "耦合协调度模型分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'explanation': {
            "耦合度": "反映多个系统之间相互作用的强度",
            "耦合协调度": "综合考虑系统发展水平和耦合程度，衡量系统之间的协调发展状况",
            "耦合度分布直方图": "展示耦合度值分布情况的直方图",
            "耦合协调度分布直方图": "展示耦合协调度值分布情况的直方图"
        },
        'interpretation': {
            "耦合度": "值越接近 1，系统间相互作用越强",
            "耦合协调度": "值越接近 1，系统间协调发展程度越高",
            "耦合度分布直方图": "直观观察耦合度值的分布特征",
            "耦合协调度分布直方图": "直观观察耦合协调度值的分布特征"
        }
    },
    'en': {
        'title': "Coupling Coordination Degree Model Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'explanation': {
            "耦合度": "Reflects the intensity of interaction between multiple systems",
            "耦合协调度": "Comprehensively considers the development level and coupling degree of systems to measure the coordinated development status between systems",
            "耦合度分布直方图": "A histogram showing the distribution of coupling degree values",
            "耦合协调度分布直方图": "A histogram showing the distribution of coupling coordination degree values"
        },
        'interpretation': {
            "耦合度": "The closer the value is to 1, the stronger the interaction between systems",
            "耦合协调度": "The closer the value is to 1, the higher the coordinated development degree between systems",
            "耦合度分布直方图": "Visually observe the distribution characteristics of coupling degree values",
            "耦合协调度分布直方图": "Visually observe the distribution characteristics of coupling coordination degree values"
        }
    }
}

# 当前语言
current_language = 'en'


def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        file_entry.configure(style="TEntry")


def coupling_coordination_degree(data):
    """
    计算耦合度和耦合协调度
    :param data: 原始数据矩阵，每列代表一个系统
    :return: 耦合度, 耦合协调度
    """
    n = data.shape[1]  # 系统数量
    # 计算各系统的综合发展水平
    u = np.mean(data, axis=0)

    # 计算耦合度
    C = n * np.power(np.prod(u), 1 / n) / np.sum(u)

    # 计算综合发展指数
    T = np.sum(u)

    # 计算耦合协调度
    D = np.sqrt(C * T)

    return C, D


def analyze_file():
    global current_language
    file_path = file_entry.get()
    if file_path == "请输入待分析 Excel 文件的完整路径" or file_path == "Please enter the full path of the Excel file to be analyzed":
        file_path = ""
    if not os.path.exists(file_path):
        result_label.config(text=LANGUAGES[current_language]['file_not_found'])
        return
    try:
        # 打开 Excel 文件
        df = pd.read_excel(file_path, header=None)
        data = df.values

        # 将数据转换为浮点类型
        data = data.astype(float)

        # 进行耦合协调度分析
        C, D = coupling_coordination_degree(data)

        # 整理数据
        data = [
            ["耦合度", C, ""],
            ["耦合协调度", D, ""]
        ]
        headers = ["统计量", "统计量值", "p值"]
        df = pd.DataFrame(data, columns=headers)

        # 添加解释说明
        explanations = LANGUAGES[current_language]['explanation']
        interpretations = LANGUAGES[current_language]['interpretation']
        explanation_df = pd.DataFrame([explanations])
        explanation_df = explanation_df.reindex(
            columns=["耦合度", "耦合协调度", "耦合度分布直方图", "耦合协调度分布直方图"])
        explanation_df.insert(0, "统计量_解释说明", "解释说明" if current_language == 'zh' else "Explanation")

        # 添加分析结果解读
        interpretation_df = pd.DataFrame([interpretations])
        interpretation_df = interpretation_df.reindex(
            columns=["耦合度", "耦合协调度", "耦合度分布直方图", "耦合协调度分布直方图"])
        interpretation_df.insert(0, "统计量_结果解读", "结果解读" if current_language == 'zh' else "Interpretation")

        # 合并数据、解释说明和结果解读
        combined_df = pd.concat([df, explanation_df, interpretation_df], ignore_index=True)

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建 Word 文档
            doc = Document()

            # 添加表格数据
            table = doc.add_table(rows=1, cols=len(combined_df.columns))
            hdr_cells = table.rows[0].cells
            for col_idx, header in enumerate(combined_df.columns):
                hdr_cells[col_idx].text = header
            for row in combined_df.values:
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 生成耦合度和耦合协调度分布直方图
            # 这里假设数据是多组的，若只有一组数据则需要调整逻辑
            C_list = [C]
            D_list = [D]
            fig, axes = plt.subplots(2, 1, figsize=(6, 8))

            axes[0].hist(C_list, bins=10)
            axes[0].set_title(
                '耦合度分布直方图' if current_language == 'zh' else 'Histogram of Coupling Degree Distribution')
            axes[0].set_xlabel('耦合度' if current_language == 'zh' else 'Coupling Degree')
            axes[0].set_ylabel('频数' if current_language == 'zh' else 'Frequency')

            axes[1].hist(D_list, bins=10)
            axes[1].set_title(
                '耦合协调度分布直方图' if current_language == 'zh' else 'Histogram of Coupling Coordination Degree Distribution')
            axes[1].set_xlabel('耦合协调度' if current_language == 'zh' else 'Coupling Coordination Degree')
            axes[1].set_ylabel('频数' if current_language == 'zh' else 'Frequency')

            # 保存图片
            img_path = os.path.splitext(save_path)[0] + '_histograms.png'
            plt.tight_layout()
            plt.savefig(img_path)
            plt.close()

            # 将图片插入到 Word 文档中
            doc.add_picture(img_path, width=Inches(6))

            # 保存 Word 文档
            doc.save(save_path)

            result_msg = LANGUAGES[current_language]['analysis_success'].format(save_path)
            result_label.config(text=result_msg, wraplength=400)

        else:
            result_label.config(text=LANGUAGES[current_language]['no_save_path'])

    except Exception as e:
        result_label.config(text=LANGUAGES[current_language]['analysis_error'].format(str(e)))


def switch_language():
    global current_language
    current_language = 'en' if current_language == 'zh' else 'zh'
    root.title(LANGUAGES[current_language]['title'])
    select_button.config(text=LANGUAGES[current_language]['select_button'])
    analyze_button.config(text=LANGUAGES[current_language]['analyze_button'])
    language_label.config(text=LANGUAGES[current_language]['switch_language'])
    # 切换语言时更新提示信息
    file_entry.delete(0, tk.END)
    if current_language == 'zh':
        file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
        file_entry.configure(style="Gray.TEntry")
    else:
        file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
        file_entry.configure(style="Gray.TEntry")


def on_entry_click(event):
    """当用户点击输入框时，清除提示信息"""
    if file_entry.get() == "请输入待分析 Excel 文件的完整路径" or file_entry.get() == "Please enter the full path of the Excel file to be analyzed":
        file_entry.delete(0, tk.END)
        file_entry.configure(style="TEntry")


def on_focusout(event):
    """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
    if file_entry.get() == "":
        if current_language == 'zh':
            file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
            file_entry.configure(style="Gray.TEntry")
        else:
            file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
            file_entry.configure(style="Gray.TEntry")


# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(LANGUAGES[current_language]['title'])

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 设置窗口的宽度和高度
window_width = 500
window_height = 300

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建自定义样式
style = ttk.Style()
style.configure("Gray.TEntry", foreground="gray")

# 创建文件选择按钮
select_button = ttk.Button(root, text=LANGUAGES[current_language]['select_button'], command=select_file,
                           bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建文件路径输入框
file_entry = ttk.Entry(root, width=50, style="Gray.TEntry")
if current_language == 'zh':
    file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
else:
    file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
file_entry.pack(pady=5)
file_entry.bind("<FocusIn>", on_entry_click)
file_entry.bind("<FocusOut>", on_focusout)

# 创建分析按钮
analyze_button = ttk.Button(root, text=LANGUAGES[current_language]['analyze_button'], command=analyze_file,
                            bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建语言切换标签
language_label = ttk.Label(root, text=LANGUAGES[current_language]['switch_language'], cursor="hand2")
language_label.pack(pady=10)
language_label.bind("<Button-1>", lambda event: switch_language())

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
root.mainloop()