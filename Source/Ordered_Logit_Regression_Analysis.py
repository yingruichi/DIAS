import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from sklearn.metrics import accuracy_score
import matplotlib.pyplot as plt
import statsmodels.api as sm
from statsmodels.miscmodels.ordinal_model import OrderedModel  # 修正导入路径
from docx import Document
from docx.shared import Inches

# 设置中文字体，确保中文正常显示
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]
plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "有序Logit回归分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'images_saved': "图片已保存到 {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "Coefficients": "回归系数，表示每个自变量对因变量的影响程度。",
            "Intercept": "截距，是当所有自变量为 0 时因变量的预测值。",
            "Accuracy": "准确率，衡量模型预测正确的比例。",
            "z-value": "z 统计量，用于检验每个自变量的显著性。",
            "p-value": "p 值，用于判断自变量的显著性，p 值越小，自变量越显著。",
            "Thresholds": "阈值参数，用于确定有序分类的边界。"
        },
        'interpretation': {
            "Coefficients": "正值表示该自变量增加时，因变量取值有增大的趋势；负值表示该自变量增加时，因变量取值有减小的趋势。",
            "Intercept": "反映了基础水平下的响应倾向。",
            "Accuracy": "模型预测准确率越高，说明模型对数据的拟合效果越好。",
            "z-value": "绝对值越大，表示该自变量对因变量的影响越显著。",
            "p-value": "通常以0.05为阈值，小于0.05表示该自变量对因变量有显著影响。",
            "Thresholds": "决定了有序分类的临界点，用于将线性预测转换为分类结果。"
        }
    },
    'en': {
        'title': "Ordered Logit Regression Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'images_saved': "Images have been saved to {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "Coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "Intercept": "Intercept, which is the predicted value of the dependent variable when all independent variables are 0.",
            "Accuracy": "Accuracy, measuring the proportion of correct predictions of the model.",
            "z-value": "z statistic, used to test the significance of each independent variable.",
            "p-value": "p value, used to determine the significance of the independent variable. The smaller the p value, the more significant the independent variable.",
            "Thresholds": "Threshold parameters that determine the boundaries between ordered categories."
        },
        'interpretation': {
            "Coefficients": "A positive value indicates that as the independent variable increases, the dependent variable tends to increase; a negative value indicates the opposite trend.",
            "Intercept": "Reflects the baseline response tendency.",
            "Accuracy": "The higher the accuracy, the better the model fits the data.",
            "z-value": "The larger the absolute value, the more significant the impact of the independent variable on the dependent variable.",
            "p-value": "Typically using a threshold of 0.05, values smaller than 0.05 indicate a significant impact of the independent variable on the dependent variable.",
            "Thresholds": "Determine the critical points for ordered classification, converting linear predictions into categorical results."
        }
    }
}


class OrderedLogitRegressionAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否包含足够的列
            if df.shape[1] < 2:
                raise ValueError("数据至少需要两列（自变量和因变量）。")

            # 假设最后一列是因变量，其余列是自变量
            X = df.iloc[:, :-1]
            y = df.iloc[:, -1]

            # 检查因变量是否为有序分类变量
            if not isinstance(y.dtype, pd.CategoricalDtype) or not (hasattr(y, 'ordered') and y.ordered):
                # 尝试将因变量转换为有序分类
                y = pd.Categorical(y, ordered=True)
                if not y.ordered:
                    # 如果无法自动转换，提示用户
                    unique_values = y.unique()
                    if len(unique_values) < 2:
                        raise ValueError("因变量必须至少有两个不同的值。")
                    # 尝试按值的顺序排序
                    y = pd.Categorical(y, categories=sorted(unique_values), ordered=True)
                    if not y.ordered:
                        raise ValueError("因变量需要是有序分类变量。请确保数据中的因变量列具有有序分类性质。")

            # 移除添加常数项的步骤
            # X_with_const = sm.add_constant(X)  # 注释掉这行代码
            logit_model = OrderedModel(y, X, distr='logit')  # 直接使用 X
            result = logit_model.fit()

            # 预测类别
            y_pred = result.predict().argmax(axis=1)
            # 将预测结果转换为与原始数据相同的类别
            # 修正此处，直接使用 y.categories
            y_pred_categorical = pd.Categorical.from_codes(y_pred, categories=y.categories)

            # 计算指标
            coefficients = result.params[:len(X.columns)]  # 前n个是系数
            z_values = result.tvalues[:len(X.columns)]
            p_values = result.pvalues[:len(X.columns)]
            accuracy = accuracy_score(y, y_pred_categorical)
            # 修改此处，使用 result.params[len(X.columns):] 获取阈值参数
            thresholds = result.params[len(X.columns):]

            # 准备数据
            columns_stats = ["Coefficients", "z-value", "p-value"]
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']

            # 准备系数数据
            coef_data = []
            for i, col in enumerate(X.columns):
                coef_data.append([
                    col,
                    coefficients[i],  # 不需要 +1
                    z_values[i],
                    p_values[i],
                    interpretations["Coefficients"]
                ])

            # 移除添加常数项的部分
            # 添加常数项
            # coef_data.append([
            #     "const",
            #     coefficients[0],
            #     z_values[0],
            #     p_values[0],
            #     interpretations["Intercept"]
            # ])

            # 创建系数表格
            coef_df = pd.DataFrame(coef_data, columns=[
                "Variable",
                "Coefficients",
                "z-value",
                "p-value",
                "Interpretation"
            ])

            # 创建阈值表格
            threshold_data = []
            for i, threshold in enumerate(thresholds):
                threshold_data.append([
                    f"Threshold {i + 1}",
                    threshold,
                    "",
                    "",
                    interpretations["Thresholds"]
                ])

            threshold_df = pd.DataFrame(threshold_data, columns=[
                "Variable",
                "Coefficients",
                "z-value",
                "p-value",
                "Interpretation"
            ])

            # 创建模型汇总表格
            model_summary = pd.DataFrame({
                "Metric": ["Accuracy", "No. Observations", "Log-Likelihood", "LL-Null", "LLR p-value"],
                "Value": [
                    accuracy,
                    len(y),
                    result.llf,
                    result.llnull,
                    # 修改此处，使用 result.llr_pvalue（不带括号）
                    result.llr_pvalue
                ],
                "Interpretation": [
                    interpretations["Accuracy"],
                    "",
                    "",
                    "",
                    "如果p值小于0.05，表示模型显著优于空模型"
                ]
            })

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的Word文档
                doc = Document()

                # 添加标题
                doc.add_heading('Ordered Logit Regression Analysis Results', 0)

                # 添加模型摘要
                doc.add_heading('Model Summary', 1)
                self._add_dataframe_to_doc(doc, model_summary)

                # 添加系数表格
                doc.add_heading('Coefficients', 1)
                self._add_dataframe_to_doc(doc, coef_df)

                # 添加阈值表格
                doc.add_heading('Thresholds', 1)
                self._add_dataframe_to_doc(doc, threshold_df)

                # 添加解释说明
                doc.add_heading('Explanations', 1)
                for stat, explanation in explanations.items():
                    doc.add_paragraph(f"{stat}: {explanation}")

                # 保存Word文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def _add_dataframe_to_doc(self, doc, df):
        """将DataFrame添加到Word文档中"""
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        hdr_cells = table.rows[0].cells
        for col_idx, header in enumerate(df.columns):
            hdr_cells[col_idx].text = header

        for row_idx, row in enumerate(df.values):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, value in enumerate(row):
                row_cells[col_idx].text = str(value)

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 250

        # 计算窗口的 x 和 y 坐标，使其居中
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)  # 使用 expand 选项使框架在上下方向上居中

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = OrderedLogitRegressionAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()