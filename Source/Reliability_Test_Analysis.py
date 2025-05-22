import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "信度检验分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "Cronbach's Alpha系数": "用于衡量量表或测试的内部一致性信度。",
            "样本量": "每个样本中的观测值数量。",
            "均值": "样本数据的平均值。"
        },
        'interpretation': {
            "Cronbach's Alpha系数": "Cronbach's Alpha系数越接近1，表示量表的内部一致性越好；越接近0，表示内部一致性越差。",
            "样本量": "样本量的大小会影响统计检验的稳定性，较大的样本量通常能提供更可靠的结果。",
            "均值": "均值反映了数据的平均水平，可用于比较不同变量的集中趋势。"
        }
    },
    'en': {
        'title': "Reliability Test Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "Cronbach's Alpha系数": "Used to measure the internal consistency reliability of a scale or test.",
            "样本量": "The number of observations in each sample.",
            "均值": "The average value of the sample data."
        },
        'interpretation': {
            "Cronbach's Alpha系数": "The closer the Cronbach's Alpha coefficient is to 1, the better the internal consistency of the scale; the closer it is to 0, the worse the internal consistency.",
            "样本量": "The sample size affects the stability of the statistical test. A larger sample size usually provides more reliable results.",
            "均值": "The mean reflects the average level of the data and can be used to compare the central tendencies of different variables."
        }
    }
}


class ReliabilityTestAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def cronbach_alpha(self, data):
        # 计算Cronbach's Alpha系数
        k = data.shape[1]
        variances = data.var(axis=0, ddof=1)
        total_variance = data.sum(axis=1).var(ddof=1)
        alpha = (k / (k - 1)) * (1 - variances.sum() / total_variance)
        return alpha

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行信度检验。")

            # 进行信度检验（Cronbach's Alpha）
            alpha = self.cronbach_alpha(numerical_df)

            # 计算样本量和均值
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()

            # 整理数据
            data = [
                ["Cronbach's Alpha系数", alpha, ""],
                ["样本量", sample_sizes.to_dict(), ""],
                ["均值", means.to_dict(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            result_df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["Cronbach's Alpha系数", "样本量", "均值"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["Cronbach's Alpha系数", "样本量", "均值"])
            interpretation_df.insert(0, "统计量_结果解读",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([result_df, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(
                    '信度检验分析结果' if self.current_language == 'zh' else 'Reliability Test Analysis Results', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(combined_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(combined_df.columns):
                    hdr_cells[col_idx].text = header

                for _, row in combined_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 生成图片（均值柱状图）
                fig, ax = plt.subplots()
                means.plot(kind='bar', ax=ax)
                ax.set_title('变量均值柱状图' if self.current_language == 'zh' else 'Bar Chart of Variable Means')
                ax.set_xlabel('变量' if self.current_language == 'zh' else 'Variables')
                ax.set_ylabel('均值' if self.current_language == 'zh' else 'Mean')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ReliabilityTestAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()