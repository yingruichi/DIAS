import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import tkinter.simpledialog
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "调节作用分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'input_info': "输入信息",
        'input_ind_var': "请输入自变量的列名",
        'input_mod_var': "请输入调节变量的列名",
        'input_dep_var': "请输入因变量的列名",
        'input_incomplete': "未输入完整的变量名，分析取消。",
        'explanation': {
            "自变量对因变量的主效应": "不考虑调节变量时，自变量对因变量的影响。",
            "调节变量对因变量的主效应": "不考虑自变量时，调节变量对因变量的影响。",
            "调节效应": "调节变量对自变量和因变量关系的影响。",
            "样本量": "参与分析的样本数量。"
        },
        'interpretation': {
            "自变量对因变量的主效应": "主效应显著表示自变量对因变量有直接影响。",
            "调节变量对因变量的主效应": "主效应显著表示调节变量对因变量有直接影响。",
            "调节效应": "调节效应显著表示调节变量改变了自变量和因变量之间的关系。",
            "样本量": "样本量的大小会影响统计结果的可靠性，较大的样本量通常能提供更可靠的结果。"
        }
    },
    'en': {
        'title': "Moderation Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'input_info': "Input Information",
        'input_ind_var': "Please enter the column name of the independent variable",
        'input_mod_var': "Please enter the column name of the moderator variable",
        'input_dep_var': "Please enter the column name of the dependent variable",
        'input_incomplete': "Incomplete variable names entered, analysis canceled.",
        'explanation': {
            "自变量对因变量的主效应": "The direct effect of the independent variable on the dependent variable without considering the moderator.",
            "调节变量对因变量的主效应": "The direct effect of the moderator on the dependent variable without considering the independent variable.",
            "调节效应": "The effect of the moderator on the relationship between the independent variable and the dependent variable.",
            "样本量": "The number of samples involved in the analysis."
        },
        'interpretation': {
            "自变量对因变量的主效应": "A significant main effect indicates that the independent variable has a direct impact on the dependent variable.",
            "调节变量对因变量的主效应": "A significant main effect indicates that the moderator has a direct impact on the dependent variable.",
            "调节效应": "A significant moderation effect indicates that the moderator changes the relationship between the independent variable and the dependent variable.",
            "样本量": "The sample size affects the reliability of the statistical results. A larger sample size usually provides more reliable results."
        }
    }
}

class ModerationAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def moderation_analysis(self, data, ind_var, mod_var, dep_var):
        # 第一步：自变量对因变量的主效应
        X1 = data[ind_var]
        X1 = sm.add_constant(X1)
        model1 = sm.OLS(data[dep_var], X1).fit()
        main_effect_ind = model1.params[ind_var]
        p_value_ind = model1.pvalues[ind_var]

        # 第二步：调节变量对因变量的主效应
        X2 = data[mod_var]
        X2 = sm.add_constant(X2)
        model2 = sm.OLS(data[dep_var], X2).fit()
        main_effect_mod = model2.params[mod_var]
        p_value_mod = model2.pvalues[mod_var]

        # 第三步：调节效应
        data['interaction'] = data[ind_var] * data[mod_var]
        X3 = data[[ind_var, mod_var, 'interaction']]
        X3 = sm.add_constant(X3)
        model3 = sm.OLS(data[dep_var], X3).fit()
        moderation_effect = model3.params['interaction']
        p_value_moderation = model3.pvalues['interaction']

        sample_size = len(data)

        return main_effect_ind, p_value_ind, main_effect_mod, p_value_mod, moderation_effect, p_value_moderation, sample_size

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 让用户输入自变量、调节变量和因变量的列名
            ind_var = tkinter.simpledialog.askstring(languages[self.current_language]['input_info'],
                                                     languages[self.current_language]['input_ind_var'])
            mod_var = tkinter.simpledialog.askstring(languages[self.current_language]['input_info'],
                                                     languages[self.current_language]['input_mod_var'])
            dep_var = tkinter.simpledialog.askstring(languages[self.current_language]['input_info'],
                                                     languages[self.current_language]['input_dep_var'])

            if not ind_var or not mod_var or not dep_var:
                self.result_label.config(text=languages[self.current_language]['input_incomplete'])
                return

            # 进行调节作用分析
            main_effect_ind, p_value_ind, main_effect_mod, p_value_mod, moderation_effect, p_value_moderation, sample_size = self.moderation_analysis(
                df, ind_var, mod_var, dep_var)

            # 整理数据
            data = [
                ["自变量对因变量的主效应", main_effect_ind, p_value_ind],
                ["调节变量对因变量的主效应", main_effect_mod, p_value_mod],
                ["调节效应", moderation_effect, p_value_moderation],
                ["样本量", sample_size, ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["自变量对因变量的主效应", "调节变量对因变量的主效应", "调节效应", "样本量"])
            explanation_df.insert(0, "统计量_解释说明",
                                  "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["自变量对因变量的主效应", "调节变量对因变量的主效应", "调节效应", "样本量"])
            interpretation_df.insert(0, "统计量_结果解读",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                table = doc.add_table(rows=df_result.shape[0] + 1, cols=df_result.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header
                for row_idx in range(df_result.shape[0]):
                    for col_idx in range(df_result.shape[1]):
                        table.cell(row_idx + 1, col_idx).text = str(df_result.iloc[row_idx, col_idx])

                # 添加解释说明表格
                doc.add_paragraph()
                doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", level=2)
                table = doc.add_table(rows=explanation_df.shape[0] + 1, cols=explanation_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(explanation_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx in range(explanation_df.shape[0]):
                    for col_idx in range(explanation_df.shape[1]):
                        table.cell(row_idx + 1, col_idx).text = str(explanation_df.iloc[row_idx, col_idx])

                # 添加结果解读表格
                doc.add_paragraph()
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Interpretation", level=2)
                table = doc.add_table(rows=interpretation_df.shape[0] + 1, cols=interpretation_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(interpretation_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx in range(interpretation_df.shape[0]):
                    for col_idx in range(interpretation_df.shape[1]):
                        table.cell(row_idx + 1, col_idx).text = str(interpretation_df.iloc[row_idx, col_idx])

                # 生成图片（调节效应柱状图）
                fig, ax = plt.subplots()
                effects = [main_effect_ind, main_effect_mod, moderation_effect]
                labels = ["自变量主效应", "调节变量主效应", "调节效应"] if self.current_language == 'zh' else [
                    "Independent Variable Main Effect", "Moderator Variable Main Effect", "Moderation Effect"]
                ax.bar(labels, effects)
                ax.set_title('调节作用分析结果' if self.current_language == 'zh' else 'Moderation Analysis Results')
                ax.set_ylabel('效应值' if self.current_language == 'zh' else 'Effect Value')

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_paragraph()
                doc.add_heading("调节作用分析结果图" if self.current_language == 'zh' else "Moderation Analysis Results Chart", level=2)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]['analysis_complete'].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ModerationAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()