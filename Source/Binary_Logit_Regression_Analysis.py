import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, roc_auc_score
import matplotlib.pyplot as plt
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    'zh': {
        'title': "二元Logit回归分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}",
        'images_saved': "图片已保存到 {}",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'explanation': {
            "Coefficients": "回归系数，表示每个自变量对因变量的影响程度。",
            "Intercept": "截距，是当所有自变量为 0 时因变量的预测值。",
            "Accuracy": "准确率，衡量模型预测正确的比例。",
            "ROC-AUC": "ROC曲线下面积，衡量模型的分类能力。",
            "z-value": "z 统计量，用于检验每个自变量的显著性。",
            "p-value": "p 值，用于判断自变量的显著性，p 值越小，自变量越显著。"
        }
    },
    'en': {
        'title': "Binary Logit Regression Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}",
        'images_saved': "Images have been saved to {}",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'explanation': {
            "Coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "Intercept": "Intercept, which is the predicted value of the dependent variable when all independent variables are 0.",
            "Accuracy": "Accuracy, measuring the proportion of correct predictions of the model.",
            "ROC-AUC": "Area under the ROC curve, measuring the classification ability of the model.",
            "z-value": "z statistic, used to test the significance of each independent variable.",
            "p-value": "p value, used to determine the significance of the independent variable. The smaller the p value, the more significant the independent variable."
        }
    }
}

class BinaryLogitRegressionAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()
    
    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设最后一列是因变量，其余列是自变量
            X = df.iloc[:, :-1].values
            y = df.iloc[:, -1].values

            # 检查因变量的值是否在 [0, 1] 区间内
            if not all(np.logical_and(y >= 0, y <= 1)):
                error_msg = "因变量的值必须在 [0, 1] 区间内，请检查数据。" if self.current_language == 'zh' else "The values of the dependent variable must be in the interval [0, 1]. Please check the data."
                self.result_label.config(text=error_msg)
                return

            # 进行二元Logit回归分析
            logit = LogisticRegression()
            logit.fit(X, y)
            y_pred = logit.predict(X)
            y_pred_proba = logit.predict_proba(X)[:, 1]

            # 计算指标
            coefficients = logit.coef_[0]
            intercept = logit.intercept_[0]
            accuracy = accuracy_score(y, y_pred)
            roc_auc = roc_auc_score(y, y_pred_proba)

            # 计算 z 值和 p 值
            X_with_const = sm.add_constant(X)
            logit_model = sm.Logit(y, X_with_const).fit()
            z_values = logit_model.tvalues
            p_values = logit_model.pvalues

            # 准备数据
            columns_stats = ["Coefficients", "z-value", "p-value", "Accuracy", "ROC-AUC"]
            explanations = languages[self.current_language]['explanation']
            data = [["Binary Logit Regression"] + list(coefficients) + [intercept, accuracy, roc_auc] + list(z_values) + list(p_values)]
            headers = ["Model"] + [f"Coefficient_{i+1}" for i in range(len(coefficients))] + ["Intercept"] + columns_stats[3:] + [f"z-value_{i+1}" for i in range(len(z_values))] + [f"p-value_{i+1}" for i in range(len(p_values))]

            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=columns_stats)
            explanation_df.insert(0, "Model", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 合并数据和解释说明
            combined_df = pd.concat([df, explanation_df], ignore_index=True)

            # 转置数据框
            transposed_df = combined_df.set_index('Model').T.reset_index().rename(columns={'index': 'Model'})

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表格标题
                doc.add_heading('Binary Logit Regression Analysis Results', level=1)

                # 添加表格
                table = doc.add_table(rows=transposed_df.shape[0]+1, cols=transposed_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(transposed_df.columns):
                    hdr_cells[col_idx].text = col_name

                for row_idx in range(transposed_df.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(transposed_df.iloc[row_idx]):
                        row_cells[col_idx].text = str(value)

                # 获取保存路径的目录
                save_dir = os.path.dirname(save_path)

                # 生成ROC曲线
                from sklearn.metrics import roc_curve
                fpr, tpr, thresholds = roc_curve(y, y_pred_proba)
                plt.figure(figsize=(10, 6))
                plt.plot(fpr, tpr, label=f'ROC curve (area = {roc_auc:.2f})')
                plt.plot([0, 1], [0, 1], 'k--')
                plt.xlim([0.0, 1.0])
                plt.ylim([0.0, 1.05])
                plt.xlabel('False Positive Rate')
                plt.ylabel('True Positive Rate')
                plt.title('Receiver Operating Characteristic')
                plt.legend(loc="lower right")
                img_name = "logit_regression_roc.png"
                img_path = os.path.join(save_dir, img_name)
                plt.savefig(img_path)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_heading('ROC Curve', level=2)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_complete'].format(save_path)
                result_msg += "\n" + languages[self.current_language]['images_saved'].format(save_dir)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
    
    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"], 
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
    
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = BinaryLogitRegressionAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()