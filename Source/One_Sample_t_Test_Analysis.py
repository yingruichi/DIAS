import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置字体为支持中文的字体，如 SimHei
plt.rcParams['font.family'] = 'SimHei'
# 解决负号显示问题
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "单样本 t 检验分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'explanation': {
            "单样本 t 检验": "用于检验样本均值与已知总体均值之间是否存在显著差异。",
            "t 统计量": "衡量样本均值与总体均值之间差异的统计量。",
            "自由度": "在统计计算中可以自由变动的数值的个数。",
            "p 值": "用于判断是否拒绝原假设的概率值。",
            "均值差异的置信区间": "包含总体均值差异的一个区间范围。",
            "样本均值": "样本数据的平均值。",
            "样本标准差": "样本数据的离散程度度量。"
        },
        'interpretation': {
            "t 统计量": "t 统计量的绝对值越大，说明样本均值与总体均值之间的差异越显著。",
            "p 值": "当 p 值小于显著性水平（通常为 0.05）时，拒绝原假设，认为样本均值与总体均值存在显著差异；否则，接受原假设。",
            "自由度": "自由度影响 t 分布的形状，自由度越大，t 分布越接近正态分布。",
            "均值差异的置信区间": "如果置信区间不包含 0，则说明样本均值与总体均值存在显著差异。",
            "样本均值": "反映了样本数据的平均水平。",
            "样本标准差": "反映了样本数据的离散程度。"
        }
    },
    'en': {
        'title': "One-Sample t-Test Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'explanation': {
            "One-Sample t-Test": "Used to test whether there is a significant difference between the sample mean and a known population mean.",
            "t-statistic": "A statistic that measures the difference between the sample mean and the population mean.",
            "Degrees of Freedom": "The number of independent values that can vary in a statistical calculation.",
            "p-value": "A probability value used to determine whether to reject the null hypothesis.",
            "Confidence Interval": "An interval range that contains the difference in population means.",
            "Sample Mean": "The average value of the sample data.",
            "Sample Standard Deviation": "A measure of the dispersion of the sample data."
        },
        'interpretation': {
            "t-statistic": "The larger the absolute value of the t-statistic, the more significant the difference between the sample mean and the population mean.",
            "p-value": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between the sample mean and the population mean; otherwise, the null hypothesis is accepted.",
            "Degrees of Freedom": "The degrees of freedom affect the shape of the t-distribution. The larger the degrees of freedom, the closer the t-distribution is to the normal distribution.",
            "Confidence Interval": "If the confidence interval does not contain 0, it indicates a significant difference between the sample mean and the population mean.",
            "Sample Mean": "Reflects the average level of the sample data.",
            "Sample Standard Deviation": "Reflects the dispersion of the sample data."
        }
    }
}

class OneSampleTTestAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 假设的总体均值
        self.population_mean = 0
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()
        
    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行单样本 t 检验。")

            # 进行单样本 t 检验
            results = []
            for column in numerical_df.columns:
                data = numerical_df[column].dropna()
                t_stat, p_value = stats.ttest_1samp(data, self.population_mean)
                degrees_of_freedom = len(data) - 1
                sample_mean = data.mean()
                sample_std = data.std()
                confidence_interval = stats.t.interval(0.95, degrees_of_freedom, loc=sample_mean - self.population_mean,
                                                   scale=stats.sem(data))

                results.append([column, t_stat, degrees_of_freedom, p_value, confidence_interval, sample_mean, sample_std])

            # 整理数据
            headers = ["变量" if self.current_language == 'zh' else "Variable", 
                      "t 统计量" if self.current_language == 'zh' else "t-statistic", 
                      "自由度" if self.current_language == 'zh' else "Degrees of Freedom", 
                      "p 值" if self.current_language == 'zh' else "p-value", 
                      "均值差异的置信区间" if self.current_language == 'zh' else "Confidence Interval", 
                      "样本均值" if self.current_language == 'zh' else "Sample Mean", 
                      "样本标准差" if self.current_language == 'zh' else "Sample Standard Deviation"]
            df_result = pd.DataFrame(results, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["单样本 t 检验" if self.current_language == 'zh' else "One-Sample t-Test", 
                         "t 统计量" if self.current_language == 'zh' else "t-statistic", 
                         "自由度" if self.current_language == 'zh' else "Degrees of Freedom", 
                         "p 值" if self.current_language == 'zh' else "p-value", 
                         "均值差异的置信区间" if self.current_language == 'zh' else "Confidence Interval", 
                         "样本均值" if self.current_language == 'zh' else "Sample Mean", 
                         "样本标准差" if self.current_language == 'zh' else "Sample Standard Deviation"])
            explanation_df.insert(0, "统计量_解释说明", 
                                "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["t 统计量" if self.current_language == 'zh' else "t-statistic", 
                         "p 值" if self.current_language == 'zh' else "p-value", 
                         "自由度" if self.current_language == 'zh' else "Degrees of Freedom", 
                         "均值差异的置信区间" if self.current_language == 'zh' else "Confidence Interval", 
                         "样本均值" if self.current_language == 'zh' else "Sample Mean", 
                         "样本标准差" if self.current_language == 'zh' else "Sample Standard Deviation"])
            interpretation_df.insert(0, "统计量_结果解读", 
                                   "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df_result, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('单样本 t 检验分析结果' if self.current_language == 'zh' else 'One-Sample t-Test Analysis Results', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(combined_df.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(combined_df.columns):
                    hdr_cells[i].text = col

                # 添加数据到表格
                for index, row in combined_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 获取桌面路径
                desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")

                # 绘制图形
                for column in numerical_df.columns:
                    data = numerical_df[column].dropna()
                    plt.figure(figsize=(12, 8))

                    # 柱状图
                    plt.subplot(2, 2, 1)
                    plt.bar([column], [data.mean()], yerr=data.std(), capsize=5)
                    plt.title('柱状图' if self.current_language == 'zh' else 'Bar Chart')
                    plt.ylabel('值' if self.current_language == 'zh' else 'Value')

                    # 误差线图
                    plt.subplot(2, 2, 2)
                    plt.errorbar(range(len(data)), data, yerr=data.std(), fmt='o', capsize=5)
                    plt.title('误差线图' if self.current_language == 'zh' else 'Error Bar Chart')
                    plt.ylabel('值' if self.current_language == 'zh' else 'Value')

                    # 箱线图
                    plt.subplot(2, 2, 3)
                    plt.boxplot(data)
                    plt.title('箱线图' if self.current_language == 'zh' else 'Box Plot')
                    plt.ylabel('值' if self.current_language == 'zh' else 'Value')

                    # 折线图
                    plt.subplot(2, 2, 4)
                    plt.plot(range(len(data)), data)
                    plt.title('折线图' if self.current_language == 'zh' else 'Line Chart')
                    plt.ylabel('值' if self.current_language == 'zh' else 'Value')

                    plt.tight_layout()

                    # 保存图片
                    chart_path = os.path.splitext(save_path)[0] + f"_{column}_charts.png"
                    plt.savefig(chart_path)
                    plt.close()
                    
                    # 添加图片到文档
                    doc.add_picture(chart_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 设置 wraplength 属性让文本自动换行
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        
    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"], 
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = OneSampleTTestAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()