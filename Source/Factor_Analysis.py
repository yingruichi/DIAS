import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from factor_analyzer import FactorAnalyzer
from factor_analyzer.factor_analyzer import calculate_bartlett_sphericity
from factor_analyzer.factor_analyzer import calculate_kmo
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "因子分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "因子载荷矩阵": "显示每个变量在各个因子上的载荷，反映变量与因子的相关性",
            "共同度": "表示每个变量被因子所解释的方差比例",
            "特征值和方差贡献率": "特征值表示每个因子解释的总方差，方差贡献率表示每个因子解释的方差占总方差的比例",
            "Bartlett球形检验": "检验变量之间是否存在相关性",
            "KMO检验": "衡量变量之间的偏相关性，判断数据是否适合进行因子分析",
            "碎石图": "展示特征值随因子数量的变化情况，帮助确定因子的数量"
        },
        'interpretation': {
            "因子载荷矩阵": "绝对值越大，说明变量与因子的相关性越强",
            "共同度": "值越接近1，说明变量被因子解释的程度越高",
            "特征值和方差贡献率": "特征值大于1的因子通常被保留，方差贡献率越高，说明该因子越重要",
            "Bartlett球形检验": "p值小于0.05时，拒绝原假设，表明变量之间存在相关性，适合进行因子分析",
            "KMO检验": "KMO值大于0.6时，适合进行因子分析",
            "碎石图": "曲线的拐点处通常表示合适的因子数量"
        }
    },
    'en': {
        'title': "Factor Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "因子载荷矩阵": "Shows the loadings of each variable on each factor, reflecting the correlation between variables and factors",
            "共同度": "Represents the proportion of variance of each variable explained by the factors",
            "特征值和方差贡献率": "The eigenvalue represents the total variance explained by each factor, and the variance contribution rate represents the proportion of variance explained by each factor to the total variance",
            "Bartlett球形检验": "Tests whether there is a correlation between variables",
            "KMO检验": "Measures the partial correlation between variables to determine whether the data is suitable for factor analysis",
            "碎石图": "Shows the change of eigenvalues with the number of factors, helping to determine the number of factors"
        },
        'interpretation': {
            "因子载荷矩阵": "The larger the absolute value, the stronger the correlation between the variable and the factor",
            "共同度": "The closer the value is to 1, the higher the degree to which the variable is explained by the factors",
            "特征值和方差贡献率": "Factors with eigenvalues greater than 1 are usually retained. The higher the variance contribution rate, the more important the factor",
            "Bartlett球形检验": "When the p-value is less than 0.05, the null hypothesis is rejected, indicating that there is a correlation between variables and factor analysis is suitable",
            "KMO检验": "When the KMO value is greater than 0.6, factor analysis is suitable",
            "碎石图": "The inflection point of the curve usually indicates the appropriate number of factors"
        }
    }
}

class FactorAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def factor_analysis(self, data):
        """
        进行因子分析
        :param data: 输入数据
        :return: 因子载荷矩阵、共同度、特征值、方差贡献率、Bartlett球形检验结果、KMO检验结果
        """
        # Bartlett球形检验
        chi_square_value, p_value = calculate_bartlett_sphericity(data)

        # KMO检验
        kmo_all, kmo_model = calculate_kmo(data)

        # 创建因子分析对象
        fa = FactorAnalyzer()
        fa.fit(data)

        # 计算特征值和方差贡献率
        ev, v = fa.get_eigenvalues()

        # 确定因子数量
        num_factors = sum(ev > 1)

        # 重新进行因子分析
        fa = FactorAnalyzer(n_factors=num_factors, rotation='varimax')
        fa.fit(data)

        # 获取因子载荷矩阵
        loadings = fa.loadings_

        # 获取共同度
        communalities = fa.get_communalities()

        return loadings, communalities, ev, v, (chi_square_value, p_value), kmo_model

    def plot_scree_plot(self, ev, save_path):
        """
        绘制碎石图
        :param ev: 特征值
        :param save_path: 图片保存路径
        """
        plt.figure(figsize=(10, 5))
        plt.plot(range(1, len(ev) + 1), ev, marker='o')
        plt.title('碎石图' if self.current_language == 'zh' else 'Scree Plot')
        plt.xlabel('因子数量' if self.current_language == 'zh' else 'Number of Factors')
        plt.ylabel('特征值' if self.current_language == 'zh' else 'Eigenvalues')
        img_path = os.path.splitext(save_path)[0] + '_scree_plot.png'
        plt.savefig(img_path)
        plt.close()
        return img_path

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)
            original_data = df.values

            # 进行因子分析
            loadings, communalities, ev, v, bartlett_result, kmo_result = self.factor_analysis(df)

            # 整理数据
            factor_names = [f'因子{i + 1}' for i in range(len(loadings[0]))]
            loadings_df = pd.DataFrame(loadings, index=df.columns, columns=factor_names)
            communalities_df = pd.DataFrame(communalities, index=df.columns, columns=['共同度'])
            ev_df = pd.DataFrame(ev, columns=['特征值'])
            v_df = pd.DataFrame(v, columns=['方差贡献率'])
            bartlett_df = pd.DataFrame([bartlett_result], columns=['卡方值', 'p值'], index=['Bartlett球形检验'])
            kmo_df = pd.DataFrame([kmo_result], columns=['KMO值'], index=['KMO检验'])

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["因子载荷矩阵", "共同度", "特征值和方差贡献率", "Bartlett球形检验", "KMO检验", "碎石图"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["因子载荷矩阵", "共同度", "特征值和方差贡献率", "Bartlett球形检验", "KMO检验", "碎石图"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 创建 Word 文档
            doc = Document()

            # 添加标题
            doc.add_heading('因子分析结果', 0)

            # 添加因子载荷矩阵
            doc.add_heading('因子载荷矩阵', 1)
            table = doc.add_table(rows=1, cols=len(loadings_df.columns) + 1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '变量'
            for col_idx, col_name in enumerate(loadings_df.columns):
                hdr_cells[col_idx + 1].text = col_name
            for row_idx, row in loadings_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = row_idx
                for col_idx, value in enumerate(row):
                    row_cells[col_idx + 1].text = str(value)

            # 添加共同度
            doc.add_heading('共同度', 1)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '变量'
            hdr_cells[1].text = '共同度'
            for row_idx, row in communalities_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = row_idx
                row_cells[1].text = str(row[0])

            # 添加特征值和方差贡献率
            doc.add_heading('特征值和方差贡献率', 1)
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '因子'
            hdr_cells[1].text = '特征值'
            hdr_cells[2].text = '方差贡献率'
            for i in range(len(ev_df)):
                row_cells = table.add_row().cells
                row_cells[0].text = f'因子{i + 1}'
                row_cells[1].text = str(ev_df.iloc[i, 0])
                row_cells[2].text = str(v_df.iloc[i, 0])

            # 添加 Bartlett 球形检验
            doc.add_heading('Bartlett球形检验', 1)
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '检验名称'
            hdr_cells[1].text = '卡方值'
            hdr_cells[2].text = 'p值'
            row_cells = table.add_row().cells
            row_cells[0].text = 'Bartlett球形检验'
            row_cells[1].text = str(bartlett_df.iloc[0, 0])
            row_cells[2].text = str(bartlett_df.iloc[0, 1])

            # 添加 KMO 检验
            doc.add_heading('KMO检验', 1)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '检验名称'
            hdr_cells[1].text = 'KMO值'
            row_cells = table.add_row().cells
            row_cells[0].text = 'KMO检验'
            row_cells[1].text = str(kmo_df.iloc[0, 0])

            # 添加解释说明
            doc.add_heading('解释说明', 1)
            table = doc.add_table(rows=1, cols=len(explanation_df.columns))
            hdr_cells = table.rows[0].cells
            for col_idx, col_name in enumerate(explanation_df.columns):
                hdr_cells[col_idx].text = col_name
            for row_idx, row in explanation_df.iterrows():
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 添加分析结果解读
            doc.add_heading('结果解读', 1)
            table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
            hdr_cells = table.rows[0].cells
            for col_idx, col_name in enumerate(interpretation_df.columns):
                hdr_cells[col_idx].text = col_name
            for row_idx, row in interpretation_df.iterrows():
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 生成碎石图
                img_path = self.plot_scree_plot(ev, save_path)

                # 添加碎石图到 Word 文档
                doc.add_heading('碎石图', 1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = FactorAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()