import tkinter as tk
from tkinter import filedialog
import openpyxl
import os
import pandas as pd
import numpy as np
from statsmodels.stats.outliers_influence import variance_inflation_factor
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import matplotlib
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
matplotlib.rcParams['font.family'] = 'SimHei'
# 解决负号显示问题
matplotlib.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "共线性分析 (VIF)",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "columns_stats": ["变量名", "方差膨胀因子（VIF）", "", "结果解读"],
        "interpretation_low_vif": "方差膨胀因子（VIF）小于 5，表明该变量与其他变量之间不存在严重的共线性。",
        "interpretation_medium_vif": "方差膨胀因子（VIF）在 5 到 10 之间，表明该变量与其他变量之间可能存在一定的共线性。",
        "interpretation_high_vif": "方差膨胀因子（VIF）大于 10，表明该变量与其他变量之间存在严重的共线性。",
        "switch_language_button_text": "切换语言"
    },
    "en": {
        "title": "Collinearity Analysis (VIF)",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "columns_stats": ["Variable Name", "Variance Inflation Factor (VIF)", "", "Result Interpretation"],
        "interpretation_low_vif": "The Variance Inflation Factor (VIF) is less than 5, indicating that there is no severe collinearity between this variable and other variables.",
        "interpretation_medium_vif": "The Variance Inflation Factor (VIF) is between 5 and 10, indicating that there may be some collinearity between this variable and other variables.",
        "interpretation_high_vif": "The Variance Inflation Factor (VIF) is greater than 10, indicating that there is severe collinearity between this variable and other variables.",
        "switch_language_button_text": "Switch Language"
    }
}

# 当前语言，默认为英文
current_language = "en"


def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        file_entry.config(foreground='black')


def on_entry_click(event):
    if file_entry.get() == languages[current_language]["file_entry_placeholder"]:
        file_entry.delete(0, tk.END)
        file_entry.config(foreground='black')


def on_focusout(event):
    if file_entry.get() == "":
        file_entry.insert(0, languages[current_language]["file_entry_placeholder"])
        file_entry.config(foreground='gray')


# 计算方差膨胀因子（VIF）的函数
def calculate_vif(X):
    vif_data = pd.DataFrame()
    vif_data["Variable Name"] = X.columns
    vif_data["Variance Inflation Factor (VIF)"] = [variance_inflation_factor(X.values, i) for i in range(X.shape[1])]
    return vif_data


def analyze_file():
    file_path = file_entry.get()
    if file_path == languages[current_language]["file_entry_placeholder"]:
        result_label.config(text=languages[current_language]["no_file_selected"])
        return
    if not os.path.exists(file_path):
        result_label.config(text=languages[current_language]["file_not_exists"])
        return
    try:
        # 打开 Excel 文件
        df = pd.read_excel(file_path)

        # 假设最后一列为因变量，其余为自变量
        X = df.iloc[:, :-1]

        # 计算方差膨胀因子（VIF）
        vif_data = calculate_vif(X)

        # 根据 VIF 值进行结果解读
        interpretations = []
        for vif in vif_data["Variance Inflation Factor (VIF)"]:
            if vif < 5:
                interpretations.append(languages[current_language]["interpretation_low_vif"])
            elif 5 <= vif < 10:
                interpretations.append(languages[current_language]["interpretation_medium_vif"])
            else:
                interpretations.append(languages[current_language]["interpretation_high_vif"])

        vif_data["Result Interpretation"] = interpretations

        # 绘制 VIF 值的柱状图
        plt.figure()
        plt.bar(vif_data["Variable Name"], vif_data["Variance Inflation Factor (VIF)"])
        plt.xlabel('Variable Name')
        plt.ylabel('Variance Inflation Factor (VIF)')
        plt.title('Variance Inflation Factor (VIF) for Each Variable')
        plt.xticks(rotation=45)

        # 保存图片
        image_path = os.path.splitext(file_path)[0] + '_vif_plot.png'
        plt.savefig(image_path)
        plt.close()

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建一个新的 Word 文档
            doc = Document()

            # 添加标题
            doc.add_heading('共线性分析 (VIF) 结果', 0)

            # 添加表格
            table = doc.add_table(rows=1, cols=len(languages[current_language]["columns_stats"]))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(languages[current_language]["columns_stats"]):
                hdr_cells[i].text = col

            # 添加数据到表格
            for index, row in vif_data.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['Variable Name'])
                row_cells[1].text = str(row['Variance Inflation Factor (VIF)'])
                row_cells[3].text = str(row['Result Interpretation'])

            # 添加图片到文档
            doc.add_picture(image_path, width=Inches(6))

            # 保存 Word 文档
            doc.save(save_path)

            # 设置 wraplength 属性让文本自动换行
            result_label.config(text=languages[current_language]["analysis_complete"].format(save_path), wraplength=400)
        else:
            result_label.config(text=languages[current_language]["no_save_path_selected"])

    except Exception as e:
        result_label.config(text=languages[current_language]["analysis_error"].format(str(e)))


def switch_language(event):
    global current_language
    if current_language == "zh":
        current_language = "en"
    else:
        current_language = "zh"

    # 更新界面文字
    root.title(languages[current_language]["title"])
    select_button.config(text=languages[current_language]["select_button_text"])
    file_entry.delete(0, tk.END)
    file_entry.insert(0, languages[current_language]["file_entry_placeholder"])
    file_entry.config(foreground='gray')
    analyze_button.config(text=languages[current_language]["analyze_button_text"])
    switch_language_label.config(text=languages[current_language]["switch_language_button_text"])


# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(languages[current_language]["title"])

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 设置窗口的宽度和高度
window_width = 500
window_height = 300

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建一个框架来包含按钮和输入框
frame = ttk.Frame(root)
frame.pack(expand=True)

# 创建文件选择按钮
select_button = ttk.Button(frame, text=languages[current_language]["select_button_text"], command=select_file,
                           bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建文件路径输入框
file_entry = ttk.Entry(frame, width=50)
file_entry.insert(0, languages[current_language]["file_entry_placeholder"])
file_entry.config(foreground='gray')
file_entry.bind('<FocusIn>', on_entry_click)
file_entry.bind('<FocusOut>', on_focusout)
file_entry.pack(pady=5)

# 创建分析按钮
analyze_button = ttk.Button(frame, text=languages[current_language]["analyze_button_text"], command=analyze_file,
                            bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建切换语言标签
switch_language_label = ttk.Label(frame, text=languages[current_language]["switch_language_button_text"],
                                  foreground="gray", cursor="hand2")
switch_language_label.bind("<Button-1>", switch_language)
switch_language_label.pack(pady=10)

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
root.mainloop()