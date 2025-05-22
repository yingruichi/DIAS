import tkinter as tk
from tkinter import filedialog
import openpyxl
import os
import pandas as pd
import numpy as np
from scipy import stats
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import matplotlib
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
matplotlib.rcParams['font.family'] = 'SimHei'
# 解决负号显示问题
matplotlib.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "偏相关分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "columns_stats": ["变量对", "偏相关系数", "p 值", "结果解读"],
        "interpretation_low_p": "p 值小于 0.05，表明该变量对之间的偏相关性显著。",
        "interpretation_high_p": "p 值大于等于 0.05，表明该变量对之间的偏相关性不显著。",
        "switch_language_button_text": "切换语言"
    },
    "en": {
        "title": "Partial Correlation Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "columns_stats": ["Variable Pair", "Partial Correlation Coefficient", "p-value", "Result Interpretation"],
        "interpretation_low_p": "The p-value is less than 0.05, indicating that the partial correlation between this variable pair is significant.",
        "interpretation_high_p": "The p-value is greater than or equal to 0.05, indicating that the partial correlation between this variable pair is not significant.",
        "switch_language_button_text": "Switch Language"
    }
}


class PartialCorrelationAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    # 计算偏相关系数的函数
    def partial_corr(self, data, x, y, z):
        sub_data = data[[x, y] + z]
        sub_data = sub_data.dropna()
        X = sub_data[[x] + z]
        Y = sub_data[[y] + z]
        beta_x = np.linalg.lstsq(X, sub_data[x], rcond=None)[0]
        beta_y = np.linalg.lstsq(Y, sub_data[y], rcond=None)[0]
        resid_x = sub_data[x] - np.dot(X, beta_x)
        resid_y = sub_data[y] - np.dot(Y, beta_y)
        corr, p = stats.pearsonr(resid_x, resid_y)
        return corr, p

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 获取所有变量名
            variables = df.columns.tolist()

            # 存储结果的列表
            results = []

            # 进行偏相关分析
            for i in range(len(variables)):
                for j in range(i + 1, len(variables)):
                    x = variables[i]
                    y = variables[j]
                    other_vars = [var for var in variables if var not in [x, y]]
                    corr, p = self.partial_corr(df, x, y, other_vars)
                    pair = f"{x} - {y}"
                    if p < 0.05:
                        interpretation = languages[self.current_language]["interpretation_low_p"]
                    else:
                        interpretation = languages[self.current_language]["interpretation_high_p"]
                    results.append([pair, corr, p, interpretation])

            # 创建结果 DataFrame
            result_df = pd.DataFrame(results, columns=languages[self.current_language]["columns_stats"])

            # 绘制偏相关系数的柱状图
            plt.figure()
            plt.bar(result_df["变量对" if self.current_language == "zh" else "Variable Pair"],
                    result_df["偏相关系数" if self.current_language == "zh" else "Partial Correlation Coefficient"])
            plt.xlabel('变量对' if self.current_language == "zh" else 'Variable Pair')
            plt.ylabel('偏相关系数' if self.current_language == "zh" else 'Partial Correlation Coefficient')
            plt.title('偏相关系数分析结果' if self.current_language == "zh" else 'Partial Correlation Analysis Results')
            plt.xticks(rotation=45)

            # 保存图片
            image_path = os.path.splitext(file_path)[0] + '_partial_corr_plot.png'
            plt.savefig(image_path)
            plt.close()

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的Word文档
                doc = Document()

                # 添加标题
                doc.add_heading(
                    '偏相关分析结果' if self.current_language == "zh" else 'Partial Correlation Analysis Results', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(result_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(result_df.columns):
                    hdr_cells[col_idx].text = col_name

                for index, row in result_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加图片
                doc.add_picture(image_path, width=Inches(6))

                # 保存Word文档
                doc.save(save_path)

                # 设置 wraplength 属性让文本自动换行
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = PartialCorrelationAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()