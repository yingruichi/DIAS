import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import tkinter.simpledialog
import matplotlib.pyplot as plt
from statsmodels.formula.api import ols
from statsmodels.stats import contrast
from docx import Document

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    'zh': {
        'title': "联合分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "属性效应": "各属性对消费者偏好的影响程度。",
            "属性水平效应": "各属性不同水平对消费者偏好的影响程度。",
            "R-squared": "模型的拟合优度，值越接近1表示模型拟合效果越好。"
        },
        'interpretation': {
            "属性效应": "效应值越大，说明该属性对消费者偏好的影响越大。",
            "属性水平效应": "效应值越大，说明该属性水平越受消费者偏好。",
            "R-squared": "R-squared值接近1，说明模型能很好地解释消费者的偏好。"
        }
    },
    'en': {
        'title': "Conjoint Analysis",
        'select_button': "Select Files",
        'analyze_button': "Analyze Files",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "属性效应": "The influence of each attribute on consumer preferences.",
            "属性水平效应": "The influence of different levels of each attribute on consumer preferences.",
            "R-squared": "The goodness of fit of the model. A value closer to 1 indicates a better fit."
        },
        'interpretation': {
            "属性效应": "A larger effect value indicates a greater influence of the attribute on consumer preferences.",
            "属性水平效应": "A larger effect value indicates that the attribute level is more preferred by consumers.",
            "R-squared": "An R-squared value close to 1 indicates that the model can well explain consumer preferences."
        }
    }
}

class ConjointAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()
    
    def select_file(self):
        file_paths = filedialog.askopenfilenames(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_paths:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, ", ".join(file_paths))
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def conjoint_analysis(self, data, attribute_columns, preference_column):
        formula = f'{preference_column} ~ ' + ' + '.join(attribute_columns)
        model = ols(formula, data=data).fit()

        # 属性效应
        attribute_effects = model.params.drop('Intercept')
        # R-squared
        r_squared = model.rsquared

        all_results = {
            "属性效应": attribute_effects,
            "R-squared": r_squared
        }

        # 属性水平效应
        attribute_level_effects = {}
        for attr in attribute_columns:
            contrast_results = contrast.ContrastResults(model.t_test_pairwise(attr))
            attribute_level_effects[attr] = contrast_results.effect

        all_results["属性水平效应"] = attribute_level_effects

        return all_results

    def analyze_file(self):
        file_paths = self.file_entry.get().split(", ")
        if not file_paths or file_paths[0] == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["file_not_found"])
            return
        for file_path in file_paths:
            if not os.path.exists(file_path):
                self.result_label.config(text=languages[self.current_language]['file_not_found'])
                return
        try:
            preference_column = tkinter.simpledialog.askstring("输入信息", "请输入偏好列的列名")
            if not preference_column:
                self.result_label.config(text="未输入有效的偏好列名，分析取消。")
                return
            attribute_columns = []
            while True:
                attribute_column = tkinter.simpledialog.askstring("输入信息", "请输入属性列的列名（点击取消结束输入）")
                if attribute_column is None:
                    break
                if attribute_column.strip():
                    attribute_columns.append(attribute_column.strip())
                else:
                    print("输入的列名不能为空，请重新输入。")

            if not attribute_columns:
                self.result_label.config(text="未输入有效的属性列名，分析取消。")
                return

            all_results = []
            file_names = []
            for file_path in file_paths:
                # 打开 Excel 文件
                df = pd.read_excel(file_path)

                # 进行联合分析
                conjoint_results = self.conjoint_analysis(df, attribute_columns, preference_column)
                all_results.append(conjoint_results)
                file_names.append(os.path.basename(file_path))

            # 整理数据
            all_data = []
            for i, results in enumerate(all_results):
                # 属性效应
                for attr, effect in results["属性效应"].items():
                    all_data.append([f"{file_names[i]}_{attr}_属性效应", effect])
                # R-squared
                all_data.append([f"{file_names[i]}_R-squared", results["R-squared"]])
                # 属性水平效应
                for attr, levels in results["属性水平效应"].items():
                    for level, effect in levels.items():
                        all_data.append([f"{file_names[i]}_{attr}_{level}_属性水平效应", effect])

            headers = ["指标", "数值"]
            df_result = pd.DataFrame(all_data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["属性效应", "属性水平效应", "R-squared"])
            explanation_df.insert(0, "指标_解释说明",
                                "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["属性效应", "属性水平效应", "R-squared"])
            interpretation_df.insert(0, "指标_结果解读",
                                    "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header
                for row in df_result.values.tolist():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明表格
                doc.add_paragraph()
                doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", level=2)
                exp_table = doc.add_table(rows=1, cols=len(explanation_df.columns))
                exp_hdr_cells = exp_table.rows[0].cells
                for col_idx, header in enumerate(explanation_df.columns):
                    exp_hdr_cells[col_idx].text = header
                for row in explanation_df.values.tolist():
                    exp_row_cells = exp_table.add_row().cells
                    for col_idx, value in enumerate(row):
                        exp_row_cells[col_idx].text = str(value)

                # 添加结果解读表格
                doc.add_paragraph()
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Interpretation", level=2)
                interp_table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
                interp_hdr_cells = interp_table.rows[0].cells
                for col_idx, header in enumerate(interpretation_df.columns):
                    interp_hdr_cells[col_idx].text = header
                for row in interpretation_df.values.tolist():
                    interp_row_cells = interp_table.add_row().cells
                    for col_idx, value in enumerate(row):
                        interp_row_cells[col_idx].text = str(value)

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)

                # 生成属性效应柱状图
                attribute_effects = [result["属性效应"] for result in all_results]
                all_attributes = []
                all_effects = []
                for effects in attribute_effects:
                    for attr, effect in effects.items():
                        all_attributes.append(attr)
                        all_effects.append(effect)

                fig, ax = plt.subplots()
                ax.bar(all_attributes, all_effects)
                ax.set_title('属性效应' if self.current_language == 'zh' else 'Attribute Effects')
                ax.set_ylabel('效应值' if self.current_language == 'zh' else 'Effect Value')
                ax.set_xlabel('属性' if self.current_language == 'zh' else 'Attribute')
                plt.xticks(rotation=45)

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_attribute_effects.png'
                plt.savefig(img_path)
                plt.close()

            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language"])
        
    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button"], 
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ConjointAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()