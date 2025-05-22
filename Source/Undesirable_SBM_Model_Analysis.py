import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from pyDOE2 import *
from pyDEA.core.models.model_builder import ModelBuilder
from pyDEA.core.data_processing.read_data import read_data
from pyDEA.core.data_processing.write_data import write_data
from pyDEA.core.utils.dea_utils import clean_up_pickled_files

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "非期望SBM模型分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'explanation': {
            "效率值": "各决策单元的非期望SBM效率得分",
            "效率值柱状图": "展示各决策单元效率值的分布情况"
        },
        'interpretation': {
            "效率值": "效率值越接近1，决策单元的效率越高",
            "效率值柱状图": "可直观比较不同决策单元的效率高低"
        }
    },
    'en': {
        'title': "Undesirable SBM Model Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'explanation': {
            "效率值": "The efficiency scores of each decision-making unit based on the undesirable SBM model",
            "效率值柱状图": "Show the distribution of efficiency values of each decision-making unit"
        },
        'interpretation': {
            "效率值": "The closer the efficiency value is to 1, the higher the efficiency of the decision-making unit",
            "效率值柱状图": "It can visually compare the efficiency levels of different decision-making units"
        }
    }
}

class UndesirableSBMModelAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def undesirable_sbm_analysis(self, data):
        """
        进行非期望SBM模型分析
        :param data: 输入数据，假设格式为 Excel 文件，包含输入、期望产出和非期望产出列
        :return: 各决策单元的效率值
        """
        # 这里简单假设数据格式符合 pyDEA 要求
        # 读取数据
        categories, input_categories, output_categories, undesirable_output_categories = read_data(data)

        # 构建非期望SBM模型
        model = ModelBuilder('SBM_Input_Oriented', 'constant', 'non_increasing',
                             input_categories, output_categories,
                             undesirable_output_categories=undesirable_output_categories)

        # 运行模型
        results = model.run(categories)

        # 提取效率值
        efficiency_values = [results.get_efficiency_score(dmu) for dmu in results.dmu_list]

        return efficiency_values

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 进行非期望SBM模型分析
            efficiency_values = self.undesirable_sbm_analysis(file_path)

            # 整理数据
            dmu_names = [f"DMU{i + 1}" for i in range(len(efficiency_values))]
            data = [
                ["效率值", efficiency_values, ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["效率值", "效率值柱状图"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["效率值", "效率值柱状图"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df_result, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表格
                table = doc.add_table(rows=combined_df.shape[0], cols=combined_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(combined_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx in range(1, combined_df.shape[0]):
                    row_cells = table.rows[row_idx].cells
                    for col_idx, value in enumerate(combined_df.iloc[row_idx]):
                        row_cells[col_idx].text = str(value)

                # 生成效率值柱状图
                fig, ax = plt.subplots()
                ax.bar(dmu_names, efficiency_values)
                ax.set_title(
                    '效率值柱状图' if self.current_language == 'zh' else 'Bar Chart of Efficiency Values')
                ax.set_xlabel('决策单元' if self.current_language == 'zh' else 'Decision-Making Units')
                ax.set_ylabel('效率值' if self.current_language == 'zh' else 'Efficiency Values')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_efficiency_bar.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入 Word 文档
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 清理临时文件
                clean_up_pickled_files()

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = UndesirableSBMModelAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()