import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
import pingouin as pg
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    'zh': {
        'title': "协方差分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存到 {}",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'explanation': {
            "协方差分析": "在控制一个或多个协变量的影响下，分析不同组之间因变量的均值是否存在显著差异。",
        },
        'interpretation': {
            "F": "F 统计量，用于检验组间差异的显著性。",
            "p-unc": "未经校正的 p 值，小于显著性水平（通常为 0.05）时，认为组间差异显著。",
            "np2": "偏 eta 平方，反映了组间差异在总变异中所占的比例。"
        }
    },
    'en': {
        'title': "Analysis of Covariance (ANCOVA)",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the image has been saved to {}",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'explanation': {
            "Analysis of Covariance (ANCOVA)": "Analyze whether there are significant differences in the means of the dependent variable between different groups while controlling for the effects of one or more covariates.",
        },
        'interpretation': {
            "F": "F statistic, used to test the significance of the differences between groups.",
            "p-unc": "Uncorrected p-value. When it is less than the significance level (usually 0.05), the differences between groups are considered significant.",
            "np2": "Partial eta squared, reflecting the proportion of the variance explained by the group differences in the total variance."
        }
    }
}

class ANCOVAAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()
        
    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设第一列是分组变量，最后一列是因变量，其余列是协变量
            group_var = df.columns[0]
            dep_var = df.columns[-1]
            covar_vars = df.columns[1:-1]

            # 进行协方差分析
            ancova = pg.ancova(data=df, dv=dep_var, between=group_var, covar=covar_vars.tolist())

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["协方差分析" if self.current_language == 'zh' else "Analysis of Covariance (ANCOVA)"])
            explanation_df.insert(0, "统计量", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=[
                "F", "p-unc", "np2"
            ])
            interpretation_df.insert(0, "统计量", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([ancova, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的Word文档
                doc = Document()

                # 添加标题
                doc.add_heading('协方差分析结果' if self.current_language == 'zh' else 'Analysis of Covariance (ANCOVA) Results', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(combined_df.columns))
                hdr_cells = table.rows[0].cells
                for col_index, col_name in enumerate(combined_df.columns):
                    hdr_cells[col_index].text = col_name

                # 添加数据行
                for index, row in combined_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_index, value in enumerate(row):
                        row_cells[col_index].text = str(value)

                # 生成结果图片
                plot_path = os.path.splitext(save_path)[0] + '_ancova_plot.png'
                plt.figure()
                for group in df[group_var].unique():
                    group_data = df[df[group_var] == group]
                    plt.scatter(group_data[covar_vars[0]], group_data[dep_var], label=group)
                plt.xlabel(covar_vars[0])
                plt.ylabel(dep_var)
                plt.title('协方差分析结果' if self.current_language == 'zh' else 'Analysis of Covariance (ANCOVA) Results')
                plt.legend()
                plt.savefig(plot_path)
                plt.close()

                # 将图片插入到Word文档中
                doc.add_picture(plot_path, width=Inches(6))

                # 保存Word文档
                doc.save(save_path)

                # 设置 wraplength 属性让文本自动换行
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path, plot_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        
    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"], 
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ANCOVAAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()