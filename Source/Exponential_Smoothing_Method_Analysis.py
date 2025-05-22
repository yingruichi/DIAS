import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "指数平滑法分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "原始数据": "输入的待分析数据",
            "一次指数平滑值": "通过指数平滑法计算得到的一次平滑值序列",
            "预测值": "基于一次指数平滑值得到的预测值序列",
            "预测结果折线图": "展示原始数据和预测值的折线图"
        },
        'interpretation': {
            "原始数据": "作为分析的基础数据",
            "一次指数平滑值": "反映数据的平滑趋势",
            "预测值": "反映未来趋势的预测结果",
            "预测结果折线图": "直观展示原始数据和预测值的变化趋势"
        }
    },
    'en': {
        'title': "Exponential Smoothing Method Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "原始数据": "The input data to be analyzed",
            "一次指数平滑值": "The first-order exponentially smoothed value sequence calculated by the exponential smoothing method",
            "预测值": "The predicted value sequence based on the first-order exponentially smoothed values",
            "预测结果折线图": "A line chart showing the original data and predicted values"
        },
        'interpretation': {
            "原始数据": "As the basic data for analysis",
            "一次指数平滑值": "Reflects the smoothing trend of the data",
            "预测值": "The predicted results reflecting future trends",
            "预测结果折线图": "Visually display the changing trends of the original data and predicted values"
        }
    }
}


class ExponentialSmoothingMethodAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def exponential_smoothing(self, x, alpha):
        """
        一次指数平滑法
        :param x: 原始数据序列
        :param alpha: 平滑系数
        :return: 一次指数平滑值序列
        """
        s = [x[0]]
        for i in range(1, len(x)):
            s.append(alpha * x[i] + (1 - alpha) * s[i - 1])
        return s

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values.flatten()

            # 将数据转换为浮点类型
            data = data.astype(float)

            # 进行指数平滑分析，平滑系数取 0.3
            alpha = 0.3
            smoothed_values = self.exponential_smoothing(data, alpha)
            # 预测值为最后一个平滑值
            pred_values = smoothed_values + [smoothed_values[-1]]

            # 整理数据
            data_list = [
                ["原始数据", data.tolist(), ""],
                ["一次指数平滑值", smoothed_values, ""],
                ["预测值", pred_values, ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data_list, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["原始数据", "一次指数平滑值", "预测值", "预测结果折线图"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["原始数据", "一次指数平滑值", "预测值", "预测结果折线图"])
            interpretation_df.insert(0, "统计量_结果解读",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('指数平滑法分析结果', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(combined_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(combined_df.columns):
                    hdr_cells[col_idx].text = header

                for index, row in combined_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, cell_data in enumerate(row):
                        row_cells[col_idx].text = str(cell_data)

                # 生成预测结果折线图
                plt.figure()
                plt.plot(range(len(data)), data, label='原始数据' if self.current_language == 'zh' else 'Original Data')
                plt.plot(range(len(pred_values)), pred_values,
                         label='预测值' if self.current_language == 'zh' else 'Predicted Values', linestyle='--')
                plt.title('预测结果折线图' if self.current_language == 'zh' else 'Line Chart of Prediction Results')
                plt.xlabel('时间步' if self.current_language == 'zh' else 'Time Step')
                plt.ylabel('值' if self.current_language == 'zh' else 'Value')
                plt.legend()

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_prediction_chart.png'
                plt.savefig(img_path)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_text.delete(1.0, tk.END)
                self.result_text.insert(tk.END, result_msg)

            else:
                self.result_text.delete(1.0, tk.END)
                self.result_text.insert(tk.END, LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示文本框
        self.result_text = tk.Text(self.root, height=4, width=60, wrap=tk.WORD)
        self.result_text.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ExponentialSmoothingMethodAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()