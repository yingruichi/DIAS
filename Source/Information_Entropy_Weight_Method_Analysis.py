import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "信息量权重法分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "原始数据": "输入的待分析数据",
            "指标熵值": "反映各指标信息无序程度的统计量",
            "指标冗余度": "指标熵值的互补量，反映指标提供的有效信息",
            "信息量权重": "根据指标冗余度计算得到的各指标权重"
        },
        'interpretation': {
            "原始数据": "作为分析的基础数据",
            "指标熵值": "值越大，指标信息越无序，提供的有效信息越少",
            "指标冗余度": "值越大，指标提供的有效信息越多",
            "信息量权重": "权重越大，该指标在综合评价中越重要"
        }
    },
    'en': {
        'title': "Information Entropy Weight Method Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "原始数据": "The input data to be analyzed",
            "指标熵值": "A statistic reflecting the degree of disorder of information for each indicator",
            "指标冗余度": "The complementary quantity of the indicator entropy, reflecting the effective information provided by the indicator",
            "信息量权重": "The weight of each indicator calculated based on the indicator redundancy"
        },
        'interpretation': {
            "原始数据": "As the basic data for analysis",
            "指标熵值": "The larger the value, the more disordered the indicator information and the less effective information it provides",
            "指标冗余度": "The larger the value, the more effective information the indicator provides",
            "信息量权重": "The larger the weight, the more important the indicator is in the comprehensive evaluation"
        }
    }
}


class InformationEntropyWeightMethodAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def information_entropy_weight_method(self, data):
        """
        实现信息量权重法
        :param data: 原始数据矩阵
        :return: 指标熵值, 指标冗余度, 信息量权重
        """
        # 数据标准化
        standardized_data = data / data.sum(axis=0)

        # 计算指标熵值
        entropy = -np.sum(standardized_data * np.log(standardized_data + 1e-8), axis=0) / np.log(data.shape[0])

        # 计算指标冗余度
        redundancy = 1 - entropy

        # 计算信息量权重
        weights = redundancy / redundancy.sum()

        return entropy, redundancy, weights

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 将数据转换为浮点类型
            data = data.astype(float)

            # 进行信息量权重法分析
            entropy, redundancy, weights = self.information_entropy_weight_method(data)

            # 整理数据
            data = [
                ["原始数据", data.tolist(), ""],
                ["指标熵值", entropy.tolist(), ""],
                ["指标冗余度", redundancy.tolist(), ""],
                ["信息量权重", weights.tolist(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["原始数据", "指标熵值", "指标冗余度", "信息量权重"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["原始数据", "指标熵值", "指标冗余度", "信息量权重"])
            interpretation_df.insert(0, "统计量_结果解读",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表格数据
                table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header
                for row_idx in range(df.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx in range(df.shape[1]):
                        row_cells[col_idx].text = str(df.iloc[row_idx, col_idx])

                # 添加解释说明
                doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", level=2)
                for index, row in explanation_df.iterrows():
                    for col in explanation_df.columns[1:]:
                        doc.add_paragraph(f"{col}: {row[col]}")

                # 添加分析结果解读
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Interpretation", level=2)
                for index, row in interpretation_df.iterrows():
                    for col in interpretation_df.columns[1:]:
                        doc.add_paragraph(f"{col}: {row[col]}")

                # 生成信息量权重分布饼图
                fig, ax = plt.subplots()
                labels = [f'指标{i + 1}' for i in range(len(weights))] if self.current_language == 'zh' else [
                    f'Indicator {i + 1}' for i in range(len(weights))]
                ax.pie(weights, labels=labels, autopct='%1.1f%%')
                ax.set_title(
                    '信息量权重分布饼图' if self.current_language == 'zh' else 'Pie Chart of Information Entropy Weights')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_weights_pie_chart.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading(
                    "信息量权重分布饼图" if self.current_language == 'zh' else 'Pie Chart of Information Entropy Weights',
                    level=2)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]['switch_language'],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = InformationEntropyWeightMethodAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()