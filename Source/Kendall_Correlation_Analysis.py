import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
from scipy import stats
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import pathlib
import pandas.plotting as pd_plotting
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    'zh': {
        'title': "Kendall相关性分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_success': "分析完成，结果已保存到 {}",
        'no_save_path': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'explanation': {
            "Kendall相关系数": "用于衡量两个有序变量之间的相关性，考虑了变量的顺序信息。"
        },
        'interpretation': {
            "相关系数": "相关系数的绝对值越接近1，说明两个变量之间的相关性越强；接近0则表示相关性较弱。",
            "p值": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为两个变量之间存在显著相关性；否则，接受原假设，认为两个变量之间无显著相关性。"
        }
    },
    'en': {
        'title': "Kendall Correlation Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_success': "Analysis completed. The results have been saved to {}",
        'no_save_path': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'explanation': {
            "Kendall相关系数": "Used to measure the correlation between two ordinal variables, taking into account the order information of the variables."
        },
        'interpretation': {
            "相关系数": "The closer the absolute value of the correlation coefficient is to 1, the stronger the correlation between the two variables; close to 0 indicates a weak correlation.",
            "p值": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant correlation between the two variables; otherwise, the null hypothesis is accepted, indicating no significant correlation."
        }
    }
}

class KendallCorrelationAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行相关性分析。")

            # 计算Kendall相关性
            kendall_corr = numerical_df.corr(method='kendall')

            # 计算p值
            def calculate_pvalues(df):
                df = df.dropna()._get_numeric_data()
                dfcols = pd.DataFrame(columns=df.columns)
                pvalues = dfcols.transpose().join(dfcols, how='outer')
                for r in df.columns:
                    for c in df.columns:
                        _, p = stats.kendalltau(df[r], df[c])
                        pvalues.loc[r, c] = p
                return pvalues

            kendall_pvalues = calculate_pvalues(numerical_df)

            # 整理数据
            data = []
            correlation_types = ["Kendall相关系数"]
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            for i, (corr, pvalues) in enumerate(zip([kendall_corr], [kendall_pvalues])):
                for col1 in corr.columns:
                    for col2 in corr.columns:
                        if col1 != col2:
                            data.append([f"{correlation_types[i]} ({col1} vs {col2})", corr.loc[col1, col2], pvalues.loc[col1, col2]])

            headers = ["统计量", "相关系数", "p值"]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=correlation_types)
            explanation_df.insert(0, "统计量", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["相关系数", "p值"])
            interpretation_df.insert(0, "统计量", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df_result, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()
                doc.add_heading('Kendall相关性分析结果', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header

                for _, row in combined_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 生成相关性热力图
                desktop_path = pathlib.Path.home() / 'Desktop'
                plot_path = desktop_path / 'correlation_heatmap.png'
                plt.figure(figsize=(10, 8))
                plt.imshow(kendall_corr, cmap='coolwarm', interpolation='nearest')
                plt.colorbar()
                plt.xticks(range(len(kendall_corr.columns)), kendall_corr.columns, rotation=45)
                plt.yticks(range(len(kendall_corr.columns)), kendall_corr.columns)
                for i in range(len(kendall_corr.columns)):
                    for j in range(len(kendall_corr.columns)):
                        plt.text(j, i, f'{kendall_corr.iloc[i, j]:.2f}', ha='center', va='center', color='black')
                plt.title('Kendall Correlation Heatmap')
                plt.savefig(plot_path)
                plt.close()

                # 生成散点图矩阵
                scatter_matrix_path = desktop_path / 'scatter_matrix.png'
                pd_plotting.scatter_matrix(numerical_df, alpha=0.8, figsize=(10, 10), diagonal='hist')
                plt.suptitle('Scatter Matrix')
                plt.savefig(scatter_matrix_path)
                plt.close()

                # 生成相关性柱状图
                selected_variable = numerical_df.columns[0]
                correlation_column = kendall_corr[selected_variable]
                bar_plot_path = desktop_path / 'correlation_bar_plot.png'
                plt.figure(figsize=(10, 6))
                correlation_column.plot(kind='bar')
                plt.title(f'Correlation with {selected_variable}')
                plt.xlabel('Variables')
                plt.ylabel('Correlation Coefficient')
                plt.xticks(rotation=45)
                plt.tight_layout()
                plt.savefig(bar_plot_path)
                plt.close()

                # 将图片插入 Word 文档
                doc.add_heading('相关性热力图', level=1)
                doc.add_picture(str(plot_path), width=Inches(6))
                doc.add_heading('散点图矩阵', level=1)
                doc.add_picture(str(scatter_matrix_path), width=Inches(6))
                doc.add_heading('相关性柱状图', level=1)
                doc.add_picture(str(bar_plot_path), width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                result_msg += f"\n相关性热力图已保存到 {plot_path}"
                result_msg += f"\n散点图矩阵已保存到 {scatter_matrix_path}"
                result_msg += f"\n相关性柱状图已保存到 {bar_plot_path}"
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = KendallCorrelationAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()