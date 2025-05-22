import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "CRITIC 权重法分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "原始数据矩阵": "从 Excel 文件中读取的原始数据矩阵",
            "标准差矩阵": "各指标的标准差矩阵，反映指标的对比强度",
            "相关系数矩阵": "各指标之间的相关系数矩阵，反映指标之间的冲突性",
            "信息量矩阵": "结合标准差和相关系数计算得到的各指标信息量矩阵",
            "指标权重": "根据信息量矩阵计算得到的各指标权重",
        },
        'interpretation': {
            "原始数据矩阵": "用于后续分析的基础数据",
            "标准差矩阵": "标准差越大，该指标的对比强度越大，在综合评价中越重要",
            "相关系数矩阵": "相关系数越小，指标之间的冲突性越大，该指标在综合评价中越重要",
            "信息量矩阵": "反映各指标包含的信息量，信息量越大，该指标越重要",
            "指标权重": "各指标在综合评价中的相对重要程度，权重越大越重要"
        }
    },
    'en': {
        'title': "CRITIC Weighting Method Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "原始数据矩阵": "The original data matrix read from the Excel file",
            "标准差矩阵": "The standard deviation matrix of each indicator, reflecting the contrast intensity of the indicators",
            "相关系数矩阵": "The correlation coefficient matrix between each indicator, reflecting the conflict between the indicators",
            "信息量矩阵": "The information matrix of each indicator calculated by combining the standard deviation and correlation coefficient",
            "指标权重": "The weight of each indicator calculated based on the information matrix"
        },
        'interpretation': {
            "原始数据矩阵": "The basic data for subsequent analysis",
            "标准差矩阵": "The larger the standard deviation, the greater the contrast intensity of the indicator, and the more important it is in the comprehensive evaluation",
            "相关系数矩阵": "The smaller the correlation coefficient, the greater the conflict between the indicators, and the more important the indicator is in the comprehensive evaluation",
            "信息量矩阵": "Reflects the information contained in each indicator. The greater the information, the more important the indicator",
            "指标权重": "The relative importance of each indicator in the comprehensive evaluation. The larger the weight, the more important it is"
        }
    }
}

class CRITICWeightingMethodAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()
    
    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def critic_weight_method(self, data):
        """
        实现 CRITIC 权重法
        :param data: 原始数据矩阵
        :return: 标准差矩阵, 相关系数矩阵, 信息量矩阵, 指标权重
        """
        # 计算标准差矩阵
        std_matrix = np.std(data, axis=0)

        # 计算相关系数矩阵
        corr_matrix = np.corrcoef(data, rowvar=False)

        # 计算冲突性
        conflict = 1 - corr_matrix

        # 计算信息量矩阵
        info_matrix = std_matrix * np.sum(conflict, axis=0)

        # 计算指标权重
        weights = info_matrix / np.sum(info_matrix)

        return std_matrix, corr_matrix, info_matrix, weights

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["file_not_found"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 将数据转换为浮点类型
            data = data.astype(float)

            # 进行 CRITIC 权重法分析
            std_matrix, corr_matrix, info_matrix, weights = self.critic_weight_method(data)

            # 整理数据
            data = [
                ["原始数据矩阵", data.tolist(), ""],
                ["标准差矩阵", std_matrix.tolist(), ""],
                ["相关系数矩阵", corr_matrix.tolist(), ""],
                ["信息量矩阵", info_matrix.tolist(), ""],
                ["指标权重", weights.tolist(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["原始数据矩阵", "标准差矩阵", "相关系数矩阵", "信息量矩阵", "指标权重"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["原始数据矩阵", "标准差矩阵", "相关系数矩阵", "信息量矩阵", "指标权重"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                doc.add_heading('分析结果', level=1)
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明表格
                doc.add_heading('解释说明', level=1)
                table = doc.add_table(rows=1, cols=len(explanation_df.columns))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(explanation_df.columns):
                    hdr_cells[i].text = header
                for _, row in explanation_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加结果解读表格
                doc.add_heading('结果解读', level=1)
                table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(interpretation_df.columns):
                    hdr_cells[i].text = header
                for _, row in interpretation_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 生成指标权重柱状图
                fig, ax = plt.subplots()
                ax.bar(range(len(weights)), weights)
                ax.set_title(
                    '指标权重柱状图' if self.current_language == 'zh' else 'Bar Chart of Indicator Weights')
                ax.set_xlabel('指标编号' if self.current_language == 'zh' else 'Indicator Number')
                ax.set_ylabel('指标权重' if self.current_language == 'zh' else 'Indicator Weight')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_indicator_weights.png'
                plt.savefig(img_path)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_heading('指标权重柱状图', level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language"])
        
    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button"], 
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = CRITICWeightingMethodAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()