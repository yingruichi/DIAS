import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "灰色预测模型分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "原始数据": "输入的待分析数据",
            "累加生成序列": "对原始数据进行一次累加生成得到的序列",
            "预测值": "通过灰色预测模型得到的预测值",
            "预测结果折线图": "展示原始数据和预测值的折线图"
        },
        'interpretation': {
            "原始数据": "作为分析的基础数据",
            "累加生成序列": "用于构建灰色预测模型",
            "预测值": "反映未来趋势的预测结果",
            "预测结果折线图": "直观展示原始数据和预测值的变化趋势"
        }
    },
    'en': {
        'title': "Gray Prediction Model Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "原始数据": "The input data to be analyzed",
            "累加生成序列": "The sequence obtained by accumulating the original data once",
            "预测值": "The predicted values obtained through the gray prediction model",
            "预测结果折线图": "A line chart showing the original data and predicted values"
        },
        'interpretation': {
            "原始数据": "As the basic data for analysis",
            "累加生成序列": "Used to build the gray prediction model",
            "预测值": "The predicted results reflecting future trends",
            "预测结果折线图": "Visually display the changing trends of the original data and predicted values"
        }
    }
}


class GrayPredictionModelAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供 root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.configure(style="TEntry")

    def on_entry_click(self, event):
        """当用户点击输入框时，清除提示信息"""
        if self.file_entry.get() == LANGUAGES[self.current_language]['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.configure(style="TEntry")

    def on_focusout(self, event):
        """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def gm11(self, x0, n_pred):
        """
        GM(1,1) 灰色预测模型
        :param x0: 原始数据序列
        :param n_pred: 预测步数
        :return: 预测值序列
        """
        x1 = np.cumsum(x0)
        z1 = (x1[:-1] + x1[1:]) / 2
        B = np.vstack([-z1, np.ones_like(z1)]).T
        Y = x0[1:].reshape(-1, 1)
        # 最小二乘法求解参数
        a, b = np.linalg.lstsq(B, Y, rcond=None)[0].flatten()
        # 预测累加序列
        x1_pred = [(x0[0] - b / a) * np.exp(-a * k) + b / a for k in range(len(x0) + n_pred)]
        # 还原为原始序列
        x0_pred = [x1_pred[0]] + [x1_pred[k] - x1_pred[k - 1] for k in range(1, len(x1_pred))]
        return x0_pred

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]['file_entry_placeholder']:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values.flatten()

            # 将数据转换为浮点类型
            data = data.astype(float)

            # 进行灰色预测分析，预测未来 5 步
            n_pred = 5
            pred_values = self.gm11(data, n_pred)

            # 整理数据
            data_list = [
                ["原始数据", data.tolist(), ""],
                ["累加生成序列", np.cumsum(data).tolist(), ""],
                ["预测值", pred_values, ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data_list, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["原始数据", "累加生成序列", "预测值", "预测结果折线图"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["原始数据", "累加生成序列", "预测值", "预测结果折线图"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表格数据
                doc.add_heading('分析数据', level=1)
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for index, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明
                doc.add_heading('解释说明', level=1)
                table = doc.add_table(rows=1, cols=len(explanation_df.columns))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(explanation_df.columns):
                    hdr_cells[i].text = header
                for index, row in explanation_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加分析结果解读
                doc.add_heading('结果解读', level=1)
                table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(interpretation_df.columns):
                    hdr_cells[i].text = header
                for index, row in interpretation_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 生成预测结果折线图
                plt.figure()
                plt.plot(range(len(data)), data, label='原始数据' if self.current_language == 'zh' else 'Original Data')
                plt.plot(range(len(pred_values)), pred_values, label='预测值' if self.current_language == 'zh' else 'Predicted Values',
                         linestyle='--')
                plt.title('预测结果折线图' if self.current_language == 'zh' else 'Line Chart of Prediction Results')
                plt.xlabel('时间步' if self.current_language == 'zh' else 'Time Step')
                plt.ylabel('值' if self.current_language == 'zh' else 'Value')
                plt.legend()

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_prediction_chart.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading('预测结果折线图', level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.configure(style="Gray.TEntry")

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建文件选择按钮
        self.select_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file,
                                        bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(self.root, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack(pady=5)
        self.file_entry.bind("<FocusIn>", self.on_entry_click)
        self.file_entry.bind("<FocusOut>", self.on_focusout)

        # 创建分析按钮
        self.analyze_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file,
                                         bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建语言切换标签
        self.switch_language_label = ttk.Label(self.root, text=LANGUAGES[self.current_language]['switch_language'],
                                               cursor="hand2")
        self.switch_language_label.pack(pady=10)
        self.switch_language_label.bind("<Button-1>", self.switch_language)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = GrayPredictionModelAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()