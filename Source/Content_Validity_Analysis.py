import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "内容效度分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "平均内容效度比（CVR）": "平均内容效度比用于衡量测量工具中各个题项与测量内容的相关性，取值范围在 -1 到 1 之间，越接近 1 表示相关性越强。",
            "样本量": "每个样本中的观测值数量。",
            "均值": "样本数据的平均值。",
            "标准差": "样本数据的离散程度。",
            "中位数": "样本数据的中间值。",
            "偏度": "样本数据分布的偏斜程度。",
            "峰度": "样本数据分布的峰态程度。"
        },
        'interpretation': {
            "平均内容效度比（CVR）": "平均内容效度比越接近 1，说明测量工具的内容与所测量的概念或领域相关性越强，内容效度越高。",
            "样本量": "样本量的大小会影响统计检验的稳定性，较大的样本量通常能提供更可靠的结果。",
            "均值": "均值反映了数据的平均水平，可用于比较不同变量的集中趋势。",
            "标准差": "标准差越大，说明数据的离散程度越大。",
            "中位数": "中位数不受极端值的影响，能更好地反映数据的中间水平。",
            "偏度": "偏度为正表示数据右偏，偏度为负表示数据左偏。",
            "峰度": "峰度大于 3 表示数据分布比正态分布更尖峭，峰度小于 3 表示数据分布比正态分布更平坦。"
        }
    },
    'en': {
        'title': "Content Validity Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "平均内容效度比（CVR）": "The average content validity ratio (CVR) is used to measure the correlation between each item in the measurement tool and the measured content. The value ranges from -1 to 1, and the closer it is to 1, the stronger the correlation.",
            "样本量": "The number of observations in each sample.",
            "均值": "The average value of the sample data.",
            "标准差": "The degree of dispersion of the sample data.",
            "中位数": "The median value of the sample data.",
            "偏度": "The degree of skewness of the sample data distribution.",
            "峰度": "The degree of kurtosis of the sample data distribution."
        },
        'interpretation': {
            "平均内容效度比（CVR）": "The closer the average content validity ratio (CVR) is to 1, the stronger the correlation between the content of the measurement tool and the measured concept or domain, and the higher the content validity.",
            "样本量": "The sample size affects the stability of the statistical test. A larger sample size usually provides more reliable results.",
            "均值": "The mean reflects the average level of the data and can be used to compare the central tendencies of different variables.",
            "标准差": "A larger standard deviation indicates a greater degree of dispersion of the data.",
            "中位数": "The median is not affected by extreme values and can better reflect the middle level of the data.",
            "偏度": "A positive skewness indicates a right-skewed distribution, while a negative skewness indicates a left-skewed distribution.",
            "峰度": "A kurtosis greater than 3 indicates a more peaked distribution than the normal distribution, while a kurtosis less than 3 indicates a flatter distribution than the normal distribution."
        }
    }
}

class ContentValidityAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])
            
        self.create_ui()

    def content_validity_analysis(self, data):
        # 假设数据是专家对每个题项与测量内容相关性的评分（1 表示相关，0 表示不相关）
        # 计算每个题项的内容效度比（CVR）
        cvr_values = []
        for column in data.columns:
            n = len(data[column])
            ne = data[column].sum()
            cvr = (ne - n / 2) / (n / 2)
            cvr_values.append(cvr)

        # 计算平均内容效度比
        average_cvr = np.mean(cvr_values)

        return average_cvr
        
    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.configure(style="TEntry")  # 恢复默认样式

    def on_entry_click(self, event):
        """当用户点击输入框时，清除提示信息"""
        if self.file_entry.get() == LANGUAGES[self.current_language]['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.configure(style="TEntry")  # 恢复默认样式

    def on_focusout(self, event):
        """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]['file_entry_placeholder']:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行内容效度分析。")

            # 进行内容效度分析
            average_cvr = self.content_validity_analysis(numerical_df)

            # 计算更多的统计指标
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()
            stds = numerical_df.std()
            medians = numerical_df.median()
            skewnesses = numerical_df.skew()
            kurtoses = numerical_df.kurt()

            # 整理数据
            data = [
                ["平均内容效度比（CVR）", average_cvr, ""],
                ["样本量", sample_sizes.to_dict(), ""],
                ["均值", means.to_dict(), ""],
                ["标准差", stds.to_dict(), ""],
                ["中位数", medians.to_dict(), ""],
                ["偏度", skewnesses.to_dict(), ""],
                ["峰度", kurtoses.to_dict(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["平均内容效度比（CVR）", "样本量", "均值", "标准差", "中位数", "偏度", "峰度"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["平均内容效度比（CVR）", "样本量", "均值", "标准差", "中位数", "偏度", "峰度"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表格
                table = doc.add_table(rows=combined_df.shape[0] + 1, cols=combined_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(combined_df.columns):
                    hdr_cells[col_idx].text = header

                for row_idx, row in combined_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 生成图片（均值柱状图）
                fig, ax = plt.subplots()
                means.plot(kind='bar', ax=ax)
                ax.set_title('变量均值柱状图' if self.current_language == 'zh' else 'Bar Chart of Variable Means')
                ax.set_xlabel('变量' if self.current_language == 'zh' else 'Variables')
                ax.set_ylabel('均值' if self.current_language == 'zh' else 'Mean')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        if self.current_language == 'zh':
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")
        else:
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")
        
    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建文件选择按钮
        self.select_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['select_button'], 
                                       command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(self.root, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack(pady=5)
        self.file_entry.bind("<FocusIn>", self.on_entry_click)
        self.file_entry.bind("<FocusOut>", self.on_focusout)

        # 创建分析按钮
        self.analyze_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['analyze_button'], 
                                        command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建语言切换标签
        self.language_label = ttk.Label(self.root, text=LANGUAGES[self.current_language]['switch_language'], cursor="hand2")
        self.language_label.pack(pady=10)
        self.language_label.bind("<Button-1>", self.switch_language)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ContentValidityAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()