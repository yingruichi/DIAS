import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "德尔菲专家法分析",
        "select_button": "选择文件",
        "analyze_button": "分析文件",
        "file_not_found": "文件不存在，请重新选择。",
        "analysis_success": "分析完成，结果已保存到 {}\n",
        "no_save_path": "未选择保存路径，结果未保存。",
        "analysis_error": "分析文件时出错: {}",
        "switch_language": "切换语言",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "explanation": {
            "各轮评分均值": "每一轮专家评分的平均值",
            "各轮评分标准差": "每一轮专家评分的离散程度",
            "最终共识评分": "经过多轮反馈后，专家达成共识的评分",
            "评分收敛情况": "判断专家评分是否达到收敛标准"
        },
        "interpretation": {
            "各轮评分均值": "反映每一轮专家对问题的整体评价",
            "各轮评分标准差": "标准差越小，说明专家意见越集中",
            "最终共识评分": "作为最终的决策参考",
            "评分收敛情况": "若收敛，则表示专家意见达成一致；否则，需要进一步讨论"
        }
    },
    "en": {
        "title": "Delphi Method Analysis",
        "select_button": "Select File",
        "analyze_button": "Analyze File",
        "file_not_found": "The file does not exist. Please select again.",
        "analysis_success": "Analysis completed. The results have been saved to {}\n",
        "no_save_path": "No save path selected. The results were not saved.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "switch_language": "Switch Language",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "explanation": {
            "各轮评分均值": "The average score of experts in each round",
            "各轮评分标准差": "The dispersion degree of experts' scores in each round",
            "最终共识评分": "The consensus score reached by experts after multiple rounds of feedback",
            "评分收敛情况": "Determine whether the experts' scores have reached the convergence criterion"
        },
        "interpretation": {
            "各轮评分均值": "Reflects the overall evaluation of experts on the problem in each round",
            "各轮评分标准差": "The smaller the standard deviation, the more concentrated the experts' opinions",
            "最终共识评分": "As the final decision - making reference",
            "评分收敛情况": "If it converges, it means that the experts have reached an agreement; otherwise, further discussion is needed"
        }
    }
}


class DelphiMethodAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def delphi_analysis(self, data):
        """
        进行德尔菲专家法分析
        :param data: 专家多轮评分数据，每一行代表一轮评分
        :return: 各轮评分均值、各轮评分标准差、最终共识评分、评分收敛情况
        """
        means = np.mean(data, axis=1)
        stds = np.std(data, axis=1)
        # 简单假设最后一轮评分的均值为最终共识评分
        final_consensus_score = means[-1]
        # 简单收敛标准：最后一轮标准差小于某个阈值
        convergence_threshold = 1.0
        convergence = stds[-1] < convergence_threshold
        return means, stds, final_consensus_score, convergence

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 进行德尔菲分析
            means, stds, final_consensus_score, convergence = self.delphi_analysis(data)

            # 整理数据
            result_data = [
                ["各轮评分均值", means.tolist(), ""],
                ["各轮评分标准差", stds.tolist(), ""],
                ["最终共识评分", [final_consensus_score], ""],
                ["评分收敛情况", [convergence], ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(result_data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["各轮评分均值", "各轮评分标准差", "最终共识评分", "评分收敛情况"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["各轮评分均值", "各轮评分标准差", "最终共识评分", "评分收敛情况"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(languages[self.current_language]['title'], 0)

                # 添加分析结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for col, header in enumerate(headers):
                    hdr_cells[col].text = header
                for row in df.values:
                    row_cells = table.add_row().cells
                    for col, value in enumerate(row):
                        row_cells[col].text = str(value)

                # 添加解释说明表格
                doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", 1)
                table = doc.add_table(rows=1, cols=len(explanation_df.columns))
                hdr_cells = table.rows[0].cells
                for col, header in enumerate(explanation_df.columns):
                    hdr_cells[col].text = header
                for row in explanation_df.values:
                    row_cells = table.add_row().cells
                    for col, value in enumerate(row):
                        row_cells[col].text = str(value)

                # 添加结果解读表格
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Interpretation", 1)
                table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
                hdr_cells = table.rows[0].cells
                for col, header in enumerate(interpretation_df.columns):
                    hdr_cells[col].text = header
                for row in interpretation_df.values:
                    row_cells = table.add_row().cells
                    for col, value in enumerate(row):
                        row_cells[col].text = str(value)

                # 生成最后一轮评分分布柱状图
                fig, ax = plt.subplots()
                ax.hist(data[-1], bins=10)
                ax.set_title(
                    '最后一轮评分分布柱状图' if self.current_language == 'zh' else 'Histogram of Scores in the Last Round')
                ax.set_xlabel('评分' if self.current_language == 'zh' else 'Scores')
                ax.set_ylabel('频数' if self.current_language == 'zh' else 'Frequency')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_score_distribution.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading("最后一轮评分分布柱状图" if self.current_language == 'zh' else 'Histogram of Scores in the Last Round', 1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文件
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(languages[self.current_language]['title'])
        self.select_button.config(text=languages[self.current_language]['select_button'])
        self.analyze_button.config(text=languages[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=languages[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = DelphiMethodAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()