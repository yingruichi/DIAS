import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from sklearn.cluster import KMeans
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "聚类分析 K-Means",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'explanation': {
            "聚类标签": "每个样本所属的聚类类别",
            "聚类中心": "每个聚类的中心位置",
            "聚类结果可视化": "将聚类结果以散点图形式展示"
        },
        'interpretation': {
            "聚类标签": "可用于区分不同样本所属的类别",
            "聚类中心": "代表每个聚类的典型特征",
            "聚类结果可视化": "直观展示样本的聚类分布情况"
        }
    },
    'en': {
        'title': "Clustering Analysis K-Means",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'explanation': {
            "聚类标签": "The cluster label to which each sample belongs",
            "聚类中心": "The center position of each cluster",
            "聚类结果可视化": "Visualize the clustering results in a scatter plot"
        },
        'interpretation': {
            "聚类标签": "Can be used to distinguish the categories to which different samples belong",
            "聚类中心": "Represents the typical characteristics of each cluster",
            "聚类结果可视化": "Visually show the clustering distribution of samples"
        }
    }
}

# 当前语言
current_language = 'en'


def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        file_entry.configure(style="TEntry")


def clustering_analysis(data, n_clusters=3):
    """
    进行聚类分析
    :param data: 输入数据
    :param n_clusters: 聚类的数量
    :return: 聚类标签，聚类中心
    """
    kmeans = KMeans(n_clusters=n_clusters, random_state=42)
    kmeans.fit(data)
    labels = kmeans.labels_
    centers = kmeans.cluster_centers_
    return labels, centers


def analyze_file():
    global current_language
    file_path = file_entry.get()
    if file_path == "请输入待分析 Excel 文件的完整路径" or file_path == "Please enter the full path of the Excel file to be analyzed":
        file_path = ""
    if not os.path.exists(file_path):
        result_label.config(text=LANGUAGES[current_language]['file_not_found'])
        return
    try:
        # 打开 Excel 文件
        df = pd.read_excel(file_path, header=None)
        original_data = df.values

        # 进行聚类分析
        labels, centers = clustering_analysis(original_data)

        # 整理数据
        data = [
            ["聚类标签", labels.tolist(), ""],
            ["聚类中心", centers.tolist(), ""]
        ]
        headers = ["统计量", "统计量值", "p值"]
        df = pd.DataFrame(data, columns=headers)

        # 添加解释说明
        explanations = LANGUAGES[current_language]['explanation']
        interpretations = LANGUAGES[current_language]['interpretation']
        explanation_df = pd.DataFrame([explanations])
        explanation_df = explanation_df.reindex(columns=["聚类标签", "聚类中心", "聚类结果可视化"])
        explanation_df.insert(0, "统计量_解释说明", "解释说明" if current_language == 'zh' else "Explanation")

        # 添加分析结果解读
        interpretation_df = pd.DataFrame([interpretations])
        interpretation_df = interpretation_df.reindex(columns=["聚类标签", "聚类中心", "聚类结果可视化"])
        interpretation_df.insert(0, "统计量_结果解读", "结果解读" if current_language == 'zh' else "Interpretation")

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建 Word 文档
            doc = Document()

            # 添加表格数据
            doc.add_heading('分析结果', 1)
            table = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for col_idx, col_name in enumerate(df.columns):
                hdr_cells[col_idx].text = col_name
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 添加解释说明
            doc.add_heading('解释说明', 2)
            table = doc.add_table(rows=1, cols=len(explanation_df.columns))
            hdr_cells = table.rows[0].cells
            for col_idx, col_name in enumerate(explanation_df.columns):
                hdr_cells[col_idx].text = col_name
            for _, row in explanation_df.iterrows():
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 添加结果解读
            doc.add_heading('结果解读', 2)
            table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
            hdr_cells = table.rows[0].cells
            for col_idx, col_name in enumerate(interpretation_df.columns):
                hdr_cells[col_idx].text = col_name
            for _, row in interpretation_df.iterrows():
                row_cells = table.add_row().cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 生成聚类结果可视化图片
            if original_data.shape[1] >= 2:
                fig, ax = plt.subplots()
                # 确保 labels 是一维数组
                labels = np.asarray(labels).ravel()
                ax.scatter(original_data[:, 0], original_data[:, 1], c=labels, cmap='viridis')
                ax.scatter(centers[:, 0], centers[:, 1], marker='X', s=200, c='red')
                ax.set_title(
                    '聚类结果可视化' if current_language == 'zh' else 'Visualization of Clustering Results')
                ax.set_xlabel('特征 1' if current_language == 'zh' else 'Feature 1')
                ax.set_ylabel('特征 2' if current_language == 'zh' else 'Feature 2')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_clustering_visualization.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading('聚类结果可视化', 2)
                doc.add_picture(img_path, width=Inches(6))

            # 保存 Word 文档
            doc.save(save_path)

            result_msg = LANGUAGES[current_language]['analysis_success'].format(save_path)
            result_label.config(text=result_msg, wraplength=400)
        else:
            result_label.config(text=LANGUAGES[current_language]['no_save_path'])

    except Exception as e:
        result_label.config(text=LANGUAGES[current_language]['analysis_error'].format(str(e)))


def switch_language():
    global current_language
    current_language = 'en' if current_language == 'zh' else 'zh'
    root.title(LANGUAGES[current_language]['title'])
    select_button.config(text=LANGUAGES[current_language]['select_button'])
    analyze_button.config(text=LANGUAGES[current_language]['analyze_button'])
    language_label.config(text=LANGUAGES[current_language]['switch_language'])
    # 切换语言时更新提示信息
    file_entry.delete(0, tk.END)
    if current_language == 'zh':
        file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
        file_entry.configure(style="Gray.TEntry")
    else:
        file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
        file_entry.configure(style="Gray.TEntry")


def on_entry_click(event):
    """当用户点击输入框时，清除提示信息"""
    if file_entry.get() == "请输入待分析 Excel 文件的完整路径" or file_entry.get() == "Please enter the full path of the Excel file to be analyzed":
        file_entry.delete(0, tk.END)
        file_entry.configure(style="TEntry")


def on_focusout(event):
    """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
    if file_entry.get() == "":
        if current_language == 'zh':
            file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
            file_entry.configure(style="Gray.TEntry")
        else:
            file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
            file_entry.configure(style="Gray.TEntry")


# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(LANGUAGES[current_language]['title'])

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 设置窗口的宽度和高度
window_width = 500
window_height = 300

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建自定义样式
style = ttk.Style()
style.configure("Gray.TEntry", foreground="gray")

# 创建文件选择按钮
select_button = ttk.Button(root, text=LANGUAGES[current_language]['select_button'], command=select_file,
                           bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建文件路径输入框
file_entry = ttk.Entry(root, width=50, style="Gray.TEntry")
if current_language == 'zh':
    file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
else:
    file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
file_entry.pack(pady=5)
file_entry.bind("<FocusIn>", on_entry_click)
file_entry.bind("<FocusOut>", on_focusout)

# 创建分析按钮
analyze_button = ttk.Button(root, text=LANGUAGES[current_language]['analyze_button'], command=analyze_file,
                            bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建语言切换标签
language_label = ttk.Label(root, text=LANGUAGES[current_language]['switch_language'], cursor="hand2")
language_label.pack(pady=10)
language_label.bind("<Button-1>", lambda event: switch_language())

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
root.mainloop()