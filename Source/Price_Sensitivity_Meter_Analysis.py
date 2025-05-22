import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import matplotlib.pyplot as plt
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from docx import Document
from docx.shared import Inches

# 定义语言字典
languages = {
    "zh": {
        "title": "价格敏感度测试模型分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，PSM 图已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "columns_stats": ["价格点", "太便宜比例", "便宜比例", "贵比例", "太贵比例"],
        "switch_language_button_text": "切换语言",
        "column_name_hint": "列名应为 TooCheap, Cheap, Expensive, TooExpensive"
    },
    "en": {
        "title": "Price Sensitivity Meter (PSM) Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the PSM plot has been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "columns_stats": ["Price Point", "Too Cheap Ratio", "Cheap Ratio", "Expensive Ratio", "Too Expensive Ratio"],
        "switch_language_button_text": "Switch Language",
        "column_name_hint": "Column names should be TooCheap, Cheap, Expensive, TooExpensive"
    }
}


class PriceSensitivityMeterAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 读取 Excel 文件
            df = pd.read_excel(file_path)

            # 假设数据集中包含 'TooCheap', 'Cheap', 'Expensive', 'TooExpensive' 列
            price_points = df.index
            too_cheap_ratio = df['TooCheap'] / df['TooCheap'].sum()
            cheap_ratio = df['Cheap'] / df['Cheap'].sum()
            expensive_ratio = df['Expensive'] / df['Expensive'].sum()
            too_expensive_ratio = df['TooExpensive'] / df['TooExpensive'].sum()

            # 计算交叉点
            indifference_point = None
            optimal_price_point = None
            lower_bound = None
            upper_bound = None

            for i in range(len(price_points) - 1):
                if cheap_ratio[i] < expensive_ratio[i] and cheap_ratio[i + 1] > expensive_ratio[i + 1]:
                    indifference_point = price_points[i]
                if too_cheap_ratio[i] < too_expensive_ratio[i] and too_cheap_ratio[i + 1] > too_expensive_ratio[i + 1]:
                    optimal_price_point = price_points[i]
                if too_cheap_ratio[i] < cheap_ratio[i] and too_cheap_ratio[i + 1] > cheap_ratio[i + 1]:
                    lower_bound = price_points[i]
                if too_expensive_ratio[i] < expensive_ratio[i] and too_expensive_ratio[i + 1] > expensive_ratio[i + 1]:
                    upper_bound = price_points[i]

            # 绘制 PSM 图
            plt.figure(figsize=(10, 6))
            plt.plot(price_points, too_cheap_ratio, label='Too Cheap')
            plt.plot(price_points, cheap_ratio, label='Cheap')
            plt.plot(price_points, expensive_ratio, label='Expensive')
            plt.plot(price_points, too_expensive_ratio, label='Too Expensive')

            if indifference_point:
                plt.axvline(x=indifference_point, color='r', linestyle='--',
                            label=f'Indifference Point: {indifference_point}')
            if optimal_price_point:
                plt.axvline(x=optimal_price_point, color='g', linestyle='--',
                            label=f'Optimal Price Point: {optimal_price_point}')
            if lower_bound:
                plt.axvline(x=lower_bound, color='b', linestyle='--', label=f'Lower Bound: {lower_bound}')
            if upper_bound:
                plt.axvline(x=upper_bound, color='m', linestyle='--', label=f'Upper Bound: {upper_bound}')

            plt.title('Price Sensitivity Meter (PSM)')
            plt.xlabel('Price')
            plt.ylabel('Ratio')
            plt.legend()

            psm_plot_path = os.path.splitext(file_path)[0] + '_psm_plot.png'
            plt.savefig(psm_plot_path)
            plt.close()

            # 保存结果到 DataFrame
            columns_stats = languages[self.current_language]["columns_stats"]
            data = {
                columns_stats[0]: price_points,
                columns_stats[1]: too_cheap_ratio,
                columns_stats[2]: cheap_ratio,
                columns_stats[3]: expensive_ratio,
                columns_stats[4]: too_expensive_ratio
            }
            result_df = pd.DataFrame(data)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('Price Sensitivity Meter (PSM) Analysis Results', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(result_df.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(result_df.columns):
                    hdr_cells[i].text = col

                for index, row in result_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加 PSM 图
                doc.add_picture(psm_plot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 设置 wraplength 属性让文本自动换行
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        self.column_name_hint_label.config(text=languages[self.current_language]["column_name_hint"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建列名提示标签
        self.column_name_hint_label = ttk.Label(frame, text=languages[self.current_language]["column_name_hint"],
                                                foreground="gray")
        self.column_name_hint_label.pack(pady=5)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = PriceSensitivityMeterAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()