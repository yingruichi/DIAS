import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "功效系数分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "各指标实际值": "每个评价指标的实际测量值",
            "各指标不允许值": "每个评价指标的最低可接受值",
            "各指标满意值": "每个评价指标的理想值",
            "功效系数向量": "根据各指标实际值、不允许值和满意值计算得到的功效系数",
            "综合功效系数": "所有指标功效系数的加权平均值"
        },
        'interpretation': {
            "各指标实际值": "反映各指标的实际表现",
            "各指标不允许值": "作为指标表现的下限参考",
            "各指标满意值": "作为指标表现的上限参考",
            "功效系数向量": "值越高，说明该指标表现越好",
            "综合功效系数": "综合反映所有指标的整体表现，值越高越好"
        }
    },
    'en': {
        'title': "Efficacy Coefficient Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "各指标实际值": "The actual measured values of each evaluation indicator",
            "各指标不允许值": "The minimum acceptable values of each evaluation indicator",
            "各指标满意值": "The ideal values of each evaluation indicator",
            "功效系数向量": "The efficacy coefficients calculated based on the actual values, unacceptable values, and satisfactory values of each indicator",
            "综合功效系数": "The weighted average of the efficacy coefficients of all indicators"
        },
        'interpretation': {
            "各指标实际值": "Reflects the actual performance of each indicator",
            "各指标不允许值": "Serves as the lower limit reference for indicator performance",
            "各指标满意值": "Serves as the upper limit reference for indicator performance",
            "功效系数向量": "The higher the value, the better the performance of the indicator",
            "综合功效系数": "Comprehensively reflects the overall performance of all indicators. The higher the value, the better"
        }
    }
}

class EfficacyCoefficientAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def efficacy_coefficient_analysis(self, actual_values, unacceptable_values, satisfactory_values, weights):
        """
        进行功效系数分析
        :param actual_values: 各指标实际值
        :param unacceptable_values: 各指标不允许值
        :param satisfactory_values: 各指标满意值
        :param weights: 各指标权重
        :return: 功效系数向量和综合功效系数
        """
        efficacy_coefficients = (actual_values - unacceptable_values) / (satisfactory_values - unacceptable_values) * 40 + 60
        comprehensive_efficacy_coefficient = np.dot(efficacy_coefficients, weights)
        return efficacy_coefficients, comprehensive_efficacy_coefficient

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 假设第一行为各指标实际值，第二行为各指标不允许值，第三行为各指标满意值，第四行为各指标权重
            actual_values = data[0]
            unacceptable_values = data[1]
            satisfactory_values = data[2]
            weights = data[3]

            # 进行功效系数分析
            efficacy_coefficients, comprehensive_efficacy_coefficient = self.efficacy_coefficient_analysis(actual_values,
                                                                                                          unacceptable_values,
                                                                                                          satisfactory_values,
                                                                                                          weights)

            # 整理数据
            data = [
                ["各指标实际值", actual_values.tolist(), ""],
                ["各指标不允许值", unacceptable_values.tolist(), ""],
                ["各指标满意值", satisfactory_values.tolist(), ""],
                ["功效系数向量", efficacy_coefficients.tolist(), ""],
                ["综合功效系数", [comprehensive_efficacy_coefficient], ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["各指标实际值", "各指标不允许值", "各指标满意值", "功效系数向量", "综合功效系数"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["各指标实际值", "各指标不允许值", "各指标满意值", "功效系数向量", "综合功效系数"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('功效系数分析结果', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header

                # 添加数据行
                for index, row in combined_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 生成功效系数向量柱状图
                fig, ax = plt.subplots()
                ax.bar(range(len(efficacy_coefficients)), efficacy_coefficients)
                ax.set_title(
                    '功效系数向量柱状图' if self.current_language == 'zh' else 'Bar Chart of Efficacy Coefficient Vector')
                ax.set_xlabel('指标' if self.current_language == 'zh' else 'Indicators')
                ax.set_ylabel('功效系数' if self.current_language == 'zh' else 'Efficacy Coefficient')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_efficacy_coefficient.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = EfficacyCoefficientAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()