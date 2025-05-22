import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Pt

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "Friedman检验分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "Friedman检验": "用于比较三个或更多相关样本的分布是否有显著差异。",
            "样本量": "每个样本中的观测值数量。",
            "中位数": "样本数据的中间值，将数据分为上下两部分。"
        },
        'interpretation': {
            "统计量": "Friedman检验的统计量值，用于衡量样本之间的差异程度。",
            "p值": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为样本之间存在显著差异；否则，接受原假设，认为样本之间无显著差异。",
            "样本量": "样本量的大小会影响统计检验的功效，较大的样本量通常能提供更准确的结果。",
            "中位数": "中位数反映了数据的中心位置，可用于比较不同样本的集中趋势。"
        }
    },
    'en': {
        'title': "Friedman Test Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "Friedman检验": "Used to compare whether the distributions of three or more related samples are significantly different.",
            "样本量": "The number of observations in each sample.",
            "中位数": "The middle value of the sample data, dividing the data into two parts."
        },
        'interpretation': {
            "统计量": "The test statistic value of the Friedman test, used to measure the degree of difference between samples.",
            "p值": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between samples; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "样本量": "The sample size affects the power of the statistical test. A larger sample size usually provides more accurate results.",
            "中位数": "The median reflects the central position of the data and can be used to compare the central tendencies of different samples."
        }
    }
}


class FriedmanTestApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行Friedman检验。")

            # 进行Friedman检验
            stat, p_value = stats.friedmanchisquare(*numerical_df.T.values)

            # 计算样本量和中位数
            sample_sizes = numerical_df.count()
            medians = numerical_df.median()

            # 整理数据
            data = [
                ["Friedman检验", stat, p_value],
                ["样本量", sample_sizes.to_dict(), ""],
                ["中位数", medians.to_dict(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["Friedman检验", "样本量", "中位数"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["统计量", "p值", "样本量", "中位数"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('Friedman检验分析结果', 0)

                # 添加分析结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
                        row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

                # 添加解释说明表格
                doc.add_heading('统计量解释说明', level=1)
                exp_table = doc.add_table(rows=1, cols=len(explanation_df.columns))
                exp_hdr_cells = exp_table.rows[0].cells
                for i, col in enumerate(explanation_df.columns):
                    exp_hdr_cells[i].text = col
                    exp_hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

                for _, row in explanation_df.iterrows():
                    exp_row_cells = exp_table.add_row().cells
                    for i, value in enumerate(row):
                        exp_row_cells[i].text = str(value)
                        exp_row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

                # 添加结果解读表格
                doc.add_heading('统计量结果解读', level=1)
                int_table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
                int_hdr_cells = int_table.rows[0].cells
                for i, col in enumerate(interpretation_df.columns):
                    int_hdr_cells[i].text = col
                    int_hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

                for _, row in interpretation_df.iterrows():
                    int_row_cells = int_table.add_row().cells
                    for i, value in enumerate(row):
                        int_row_cells[i].text = str(value)
                        int_row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = FriedmanTestApp()
    app.run()


if __name__ == "__main__":
    run_app()