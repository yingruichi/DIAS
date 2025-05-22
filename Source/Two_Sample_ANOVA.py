import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
from scipy import stats
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文显示
plt.rcParams['font.family'] = 'SimHei'  # 指定支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    'zh': {
        'title': "双样本方差分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'explanation': {
            "双样本方差分析": "用于比较两个独立样本的均值是否有显著差异。",
            "样本量": "每个样本中的观测值数量。",
            "均值": "样本数据的平均值。",
            "t统计量": "衡量两个样本均值差异的统计量。",
            "自由度": "用于计算t分布的参数。",
            "p值": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为两个样本均值存在显著差异；否则，接受原假设，认为两个样本均值无显著差异。",
            "效应量": "反映两个样本均值差异的程度。"
        },
        'interpretation': {
            "t统计量": "t统计量的绝对值越大，说明两个样本均值之间的差异越显著。",
            "p值": "用于判断两个样本均值之间是否存在显著差异。",
            "自由度": "影响t分布的形状，进而影响p值的计算。",
            "样本量": "样本量的大小会影响统计检验的功效，较大的样本量通常能提供更准确的结果。",
            "均值": "反映样本数据的平均水平。",
            "效应量": "效应量越大，说明两个样本均值之间的差异越大。"
        }
    },
    'en': {
        'title': "Two-sample ANOVA",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'explanation': {
            "双样本方差分析": "Used to compare whether the means of two independent samples are significantly different.",
            "样本量": "The number of observations in each sample.",
            "均值": "The average value of the sample data.",
            "t统计量": "A statistic that measures the difference between the means of two samples.",
            "自由度": "Parameters used to calculate the t-distribution.",
            "p值": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between the means of the two samples; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "效应量": "Reflects the degree of difference between the means of two samples."
        },
        'interpretation': {
            "t统计量": "The larger the absolute value of the t-statistic, the more significant the difference between the means of the two samples.",
            "p值": "Used to determine whether there is a significant difference between the means of the two samples.",
            "自由度": "Affects the shape of the t-distribution, which in turn affects the calculation of the p-value.",
            "样本量": "The sample size affects the power of the statistical test. A larger sample size usually provides more accurate results.",
            "均值": "Reflects the average level of the sample data.",
            "效应量": "The larger the effect size, the greater the difference between the means of the two samples."
        }
    }
}

class TwoSampleANOVAApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行方差分析。")

            # 确保有两列数值数据
            if len(numerical_df.columns) < 2:
                raise ValueError("数据中至少需要两列数值数据才能进行双样本方差分析。")

            # 取前两列数据进行双样本 t 检验
            sample1 = numerical_df.iloc[:, 0]
            sample2 = numerical_df.iloc[:, 1]
            t_stat, p_value = stats.ttest_ind(sample1, sample2)

            # 计算自由度
            df_value = len(sample1) + len(sample2) - 2

            # 计算效应量（Cohen's d）
            pooled_std = np.sqrt(((len(sample1) - 1) * sample1.std()**2 + (len(sample2) - 1) * sample2.std()**2) / (len(sample1) + len(sample2) - 2))
            cohen_d = (sample1.mean() - sample2.mean()) / pooled_std

            # 计算样本量和均值
            sample_sizes = [len(sample1), len(sample2)]
            means = [sample1.mean(), sample2.mean()]

            # 整理数据
            data = [
                ["方差分析", t_stat, df_value, p_value, cohen_d],
                ["样本量", sample_sizes, "", "", ""],
                ["均值", means, "", "", ""]
            ]
            headers = ["统计量", "t统计量", "自由度", "p值", "效应量（Cohen's d）"]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["双样本方差分析", "样本量", "均值", "t统计量", "自由度", "p值", "效应量"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["t统计量", "p值", "自由度", "样本量", "均值", "效应量"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                doc.add_heading('分析结果', level=1)
                table = doc.add_table(rows=df_result.shape[0] + 1, cols=df_result.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header
                for row_idx, row in df_result.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明表格
                doc.add_heading('解释说明', level=1)
                table = doc.add_table(rows=explanation_df.shape[0] + 1, cols=explanation_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(explanation_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in explanation_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加结果解读表格
                doc.add_heading('结果解读', level=1)
                table = doc.add_table(rows=interpretation_df.shape[0] + 1, cols=interpretation_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(interpretation_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in interpretation_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 绘制箱线图
                plt.figure(figsize=(10, 6))
                numerical_df.iloc[:, :2].boxplot()
                plt.title('箱线图' if self.current_language == 'zh' else 'Box Plot')
                plt.xlabel('组' if self.current_language == 'zh' else 'Groups')
                plt.ylabel('数值' if self.current_language == 'zh' else 'Values')
                boxplot_path = save_path.replace('.docx', '_boxplot.png')
                plt.savefig(boxplot_path)
                plt.close()

                # 绘制柱状图
                plt.figure(figsize=(10, 6))
                bars = plt.bar(['样本1', '样本2'] if self.current_language == 'zh' else ['Sample 1', 'Sample 2'], means)
                for bar in bars:
                    height = bar.get_height()
                    plt.annotate(f'{height:.2f}',
                                 xy=(bar.get_x() + bar.get_width() / 2, height),
                                 xytext=(0, 3),  # 3 points vertical offset
                                 textcoords="offset points",
                                 ha='center', va='bottom')
                plt.title('柱状图' if self.current_language == 'zh' else 'Bar Chart')
                plt.xlabel('组' if self.current_language == 'zh' else 'Groups')
                plt.ylabel('均值' if self.current_language == 'zh' else 'Mean')
                barplot_path = save_path.replace('.docx', '_barplot.png')
                plt.savefig(barplot_path)
                plt.close()

                # 在 Word 文档中添加图片
                doc.add_heading('箱线图', level=1)
                doc.add_picture(boxplot_path, width=Inches(6))
                doc.add_heading('柱状图', level=1)
                doc.add_picture(barplot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = TwoSampleANOVAApp()
    app.run()

if __name__ == "__main__":
    run_app()