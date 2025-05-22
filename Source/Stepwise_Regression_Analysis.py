import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_squared_error, r2_score
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches

# 定义语言字典
languages = {
    "zh": {
        "title": "逐步回归分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}\n",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "images_saved": "图片已保存到 {}",
        "switch_language_button_text": "切换语言",
        "explanation": {
            "Coefficients": "回归系数，表示每个自变量对因变量的影响程度。",
            "Intercept": "截距，是当所有自变量为 0 时因变量的预测值。",
            "Mean Squared Error (MSE)": "均方误差，衡量预测值与真实值之间的平均误差。",
            "R-squared (R²)": "决定系数，取值范围在 0 到 1 之间，越接近 1 表示模型拟合效果越好。",
            "Adjusted R-squared": "调整决定系数，考虑了模型中自变量的数量，对模型的拟合优度进行了调整。",
            "F-value": "F 统计量，用于检验整个回归模型的显著性。",
            "t-value": "t 统计量，用于检验每个自变量的显著性。",
            "p-value": "p 值，用于判断自变量的显著性，p 值越小，自变量越显著。"
        }
    },
    "en": {
        "title": "Stepwise Regression Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}\n",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "images_saved": "Images have been saved to {}",
        "switch_language_button_text": "Switch Language",
        "explanation": {
            "Coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "Intercept": "Intercept, which is the predicted value of the dependent variable when all independent variables are 0.",
            "Mean Squared Error (MSE)": "Mean squared error, measuring the average error between the predicted and actual values.",
            "R-squared (R²)": "Coefficient of determination, ranging from 0 to 1. A value closer to 1 indicates a better fit of the model.",
            "Adjusted R-squared": "Adjusted coefficient of determination, which takes into account the number of independent variables in the model and adjusts the goodness of fit of the model.",
            "F-value": "F statistic, used to test the significance of the entire regression model.",
            "t-value": "t statistic, used to test the significance of each independent variable.",
            "p-value": "p value, used to determine the significance of the independent variable. The smaller the p value, the more significant the independent variable."
        }
    }
}


class StepwiseRegressionAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def stepwise_selection(self, X, y, initial_list=[], threshold_in=0.05, threshold_out=0.10, verbose=True):
        """
        逐步回归选择自变量的函数
        """
        included = list(initial_list)
        while True:
            changed = False
            # 尝试添加变量
            excluded = list(set(X.columns) - set(included))
            new_pval = pd.Series(index=excluded)
            for new_column in excluded:
                model = sm.OLS(y, sm.add_constant(pd.DataFrame(X[included + [new_column]]))).fit()
                new_pval[new_column] = model.pvalues[new_column]
            best_pval = new_pval.min()
            if best_pval < threshold_in:
                best_feature = new_pval.idxmin()
                included.append(best_feature)
                changed = True
                if verbose:
                    print('Add  {:30} with p-value {:.6}'.format(best_feature, best_pval))

            # 尝试移除变量
            model = sm.OLS(y, sm.add_constant(pd.DataFrame(X[included]))).fit()
            # 排除常数项
            pvalues = model.pvalues.iloc[1:]
            worst_pval = pvalues.max()
            if worst_pval > threshold_out:
                changed = True
                worst_feature = pvalues.idxmax()
                included.remove(worst_feature)
                if verbose:
                    print('Drop {:30} with p-value {:.6}'.format(worst_feature, worst_pval))
            if not changed:
                break
        return included

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设最后一列是因变量，其余列是自变量
            X = df.iloc[:, :-1]
            y = df.iloc[:, -1].values

            # 逐步回归分析
            selected_features = self.stepwise_selection(X, y)
            X_selected = X[selected_features]

            model = LinearRegression()
            model.fit(X_selected, y)
            y_pred = model.predict(X_selected)

            # 计算指标
            coefficients = model.coef_
            intercept = model.intercept_
            mse = mean_squared_error(y, y_pred)
            r2 = r2_score(y, y_pred)
            n = len(y)
            p = X_selected.shape[1]
            adjusted_r2 = 1 - (1 - r2) * (n - 1) / (n - p - 1)

            # 计算 t 值和 p 值
            X_with_const = sm.add_constant(X_selected)
            sm_model = sm.OLS(y, X_with_const).fit()
            t_values = sm_model.tvalues
            p_values = sm_model.pvalues

            # 计算 F 值
            f_value = sm_model.fvalue

            results = [["Stepwise Model"] + list(coefficients) + [intercept, mse, r2, adjusted_r2, f_value] + list(
                t_values) + list(p_values)]

            # 准备数据
            columns_stats = ["Coefficients", "t-value", "p-value", "R-squared (R²)", "Adjusted R-squared", "F-value"]
            explanations = languages[self.current_language]['explanation']
            headers = ["Model"] + [f"Coefficient_{i + 1}" for i in range(len(selected_features))] + ["Intercept", "MSE",
                                                                                                    "R-squared (R²)",
                                                                                                    "Adjusted R-squared",
                                                                                                    "F-value"] + [
                          f"t-value_{i + 1}" for i in range(len(selected_features))] + [
                          f"p-value_{i + 1}" for i in range(len(selected_features))]

            df_results = pd.DataFrame(results, columns=headers)

            # 添加解释说明
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=columns_stats)
            explanation_df.insert(0, "Model", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 合并数据和解释说明
            combined_df = pd.concat([df_results, explanation_df], ignore_index=True)

            # 转置数据框
            transposed_df = combined_df.set_index('Model').T.reset_index().rename(columns={'index': 'Model'})

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加表格
                table = doc.add_table(rows=transposed_df.shape[0] + 1, cols=transposed_df.shape[1])

                # 添加表头
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(transposed_df.columns):
                    hdr_cells[col_idx].text = header

                # 添加表格数据
                for row_idx, row in transposed_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 获取保存路径的目录
                save_dir = os.path.dirname(save_path)

                # 生成散点图
                plt.figure(figsize=(10, 6))
                plt.scatter(y, y_pred)
                plt.plot([y.min(), y.max()], [y.min(), y.max()], 'r--', lw=2)
                plt.xlabel('Actual Values')
                plt.ylabel('Predicted Values')
                plt.title('Actual vs Predicted Values')
                img_name = "stepwise_regression_scatter.png"
                img_path = os.path.join(save_dir, img_name)
                plt.savefig(img_path)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_complete'].format(save_path)
                result_msg += languages[self.current_language]['images_saved'].format(save_dir)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = StepwiseRegressionAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()