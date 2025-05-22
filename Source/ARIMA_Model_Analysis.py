import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
# import pmdarima as pm
from statsmodels.tsa.arima.model import ARIMA
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    "zh": {
        "title": "ARIMA模型分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "columns_stats": ["变量名", "ARIMA参数(p, d, q)", "AIC值", "BIC值"],
        "switch_language_button_text": "切换语言"
    },
    "en": {
        "title": "ARIMA Model Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "columns_stats": ["Variable Name", "ARIMA Parameters (p, d, q)", "AIC Value", "BIC Value"],
        "switch_language_button_text": "Switch Language"
    }
}

class ARIMAModelAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()
        
    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    # 计算 ARIMA 模型的函数
    def calculate_arima(self, X):
        arima_data = pd.DataFrame()
        arima_data["Variable Name"] = X.columns
        arima_params = []
        aic_values = []
        bic_values = []

        for col in X.columns:
            try:
                # 自动选择 ARIMA 参数
                auto_arima = pm.auto_arima(X[col], seasonal=False, trace=False)
                p, d, q = auto_arima.order
                arima_params.append(f"({p}, {d}, {q})")

                # 拟合 ARIMA 模型
                model = ARIMA(X[col], order=(p, d, q))
                model_fit = model.fit()

                aic_values.append(model_fit.aic)
                bic_values.append(model_fit.bic)

                # 绘制时间序列图和预测图
                plt.figure()
                plt.plot(X[col], label='Original Data')
                forecast = model_fit.get_forecast(steps=10)
                forecast_mean = forecast.predicted_mean
                plt.plot(pd.date_range(start=X.index[-1], periods=11, freq=X.index.freq)[1:], forecast_mean, label='Forecast')
                plt.xlabel('Date')
                plt.ylabel(col)
                plt.title(f'ARIMA Model for {col}')
                plt.legend()

                # 保存图片
                image_path = os.path.splitext(self.file_entry.get())[0] + f'_arima_{col}.png'
                plt.savefig(image_path)
                plt.close()

            except Exception as e:
                arima_params.append("Error")
                aic_values.append(np.nan)
                bic_values.append(np.nan)
                print(f"Error analyzing {col}: {e}")

        arima_data["ARIMA Parameters (p, d, q)"] = arima_params
        arima_data["AIC Value"] = aic_values
        arima_data["BIC Value"] = bic_values
        return arima_data

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, index_col=0, parse_dates=True)

            # 假设所有列都是时间序列数据
            X = df

            # 计算 ARIMA 模型
            arima_data = self.calculate_arima(X)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('ARIMA Model Results', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(arima_data.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(arima_data.columns):
                    hdr_cells[i].text = col

                # 添加数据到表格
                for index, row in arima_data.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加图片
                for col in X.columns:
                    image_path = os.path.splitext(file_path)[0] + f'_arima_{col}.png'
                    if os.path.exists(image_path):
                        doc.add_heading(f'ARIMA Model for {col}', level=1)
                        doc.add_picture(image_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                # 设置 wraplength 属性让文本自动换行
                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])
        
    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"], 
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ARIMAModelAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()