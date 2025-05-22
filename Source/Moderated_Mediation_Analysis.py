import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import tkinter.simpledialog
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "调节中介作用分析",
      'select_button_text': "选择文件",
      'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
      'analyze_button_text': "分析文件",
      'no_file_selected': "请选择有效的文件路径。",
      'file_not_exists': "文件不存在，请重新选择。",
      'analysis_error': "分析文件时出错: {}",
      'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
      'no_save_path_selected': "未选择保存路径，结果未保存。",
      'switch_language_button_text': "切换语言",
      'input_info': "输入信息",
      'input_ind_var': "请输入自变量的列名",
      'input_med_var': "请输入中介变量的列名",
      'input_dep_var': "请输入因变量的列名",
      'input_mod_var': "请输入调节变量的列名",
      'input_incomplete': "未输入完整的变量名，分析取消。",
        'explanation': {
            "自变量对因变量的总效应": "自变量直接对因变量产生的影响。",
            "自变量对中介变量的效应": "自变量对中介变量产生的影响。",
            "中介变量对因变量的效应（控制自变量）": "在控制自变量的情况下，中介变量对因变量产生的影响。",
            "调节变量对自变量 - 中介变量关系的调节效应": "调节变量对自变量与中介变量之间关系的影响。",
            "调节变量对中介变量 - 因变量关系的调节效应": "调节变量对中介变量与因变量之间关系的影响。",
            "中介效应": "自变量通过中介变量对因变量产生的间接影响。",
            "样本量": "参与分析的样本数量。"
        },
        'interpretation': {
            "自变量对因变量的总效应": "总效应显著表示自变量对因变量有直接影响。",
            "自变量对中介变量的效应": "该效应显著表示自变量能够影响中介变量。",
            "中介变量对因变量的效应（控制自变量）": "此效应显著表示中介变量在控制自变量后仍对因变量有影响。",
            "调节变量对自变量 - 中介变量关系的调节效应": "该调节效应显著表示调节变量会影响自变量与中介变量之间的关系。",
            "调节变量对中介变量 - 因变量关系的调节效应": "该调节效应显著表示调节变量会影响中介变量与因变量之间的关系。",
            "中介效应": "中介效应显著表示自变量通过中介变量对因变量产生了间接影响。",
            "样本量": "样本量的大小会影响统计结果的可靠性，较大的样本量通常能提供更可靠的结果。"
        }
    },
    'en': {
        'title': "Moderated Mediation Analysis",
      'select_button_text': "Select File",
      'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
      'analyze_button_text': "Analyze File",
      'no_file_selected': "Please select a valid file path.",
      'file_not_exists': "The file does not exist. Please select again.",
      'analysis_error': "An error occurred while analyzing the file: {}",
      'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
      'no_save_path_selected': "No save path selected. The results were not saved.",
      'switch_language_button_text': "Switch Language",
      'input_info': "Input Information",
      'input_ind_var': "Please enter the column name of the independent variable",
      'input_med_var': "Please enter the column name of the mediator variable",
      'input_dep_var': "Please enter the column name of the dependent variable",
      'input_mod_var': "Please enter the column name of the moderator variable",
      'input_incomplete': "Incomplete variable names entered, analysis canceled.",
        'explanation': {
            "自变量对因变量的总效应": "The total effect of the independent variable on the dependent variable.",
            "自变量对中介变量的效应": "The effect of the independent variable on the mediator variable.",
            "中介变量对因变量的效应（控制自变量）": "The effect of the mediator variable on the dependent variable while controlling for the independent variable.",
            "调节变量对自变量 - 中介变量关系的调节效应": "The moderating effect of the moderator variable on the relationship between the independent variable and the mediator variable.",
            "调节变量对中介变量 - 因变量关系的调节效应": "The moderating effect of the moderator variable on the relationship between the mediator variable and the dependent variable.",
            "中介效应": "The indirect effect of the independent variable on the dependent variable through the mediator variable.",
            "样本量": "The number of samples involved in the analysis."
        },
        'interpretation': {
            "自变量对因变量的总效应": "A significant total effect indicates that the independent variable has a direct impact on the dependent variable.",
            "自变量对中介变量的效应": "A significant effect indicates that the independent variable can influence the mediator variable.",
            "中介变量对因变量的效应（控制自变量）": "A significant effect indicates that the mediator variable still has an impact on the dependent variable after controlling for the independent variable.",
            "调节变量对自变量 - 中介变量关系的调节效应": "A significant moderating effect indicates that the moderator variable affects the relationship between the independent variable and the mediator variable.",
            "调节变量对中介变量 - 因变量关系的调节效应": "A significant moderating effect indicates that the moderator variable affects the relationship between the mediator variable and the dependent variable.",
            "中介效应": "A significant mediation effect indicates that the independent variable has an indirect impact on the dependent variable through the mediator variable.",
            "样本量": "The sample size affects the reliability of the statistical results. A larger sample size usually provides more reliable results."
        }
    }
}

class ModeratedMediationAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def moderated_mediation_analysis(self, data, ind_var, med_var, dep_var, mod_var):
        # 第一步：自变量对因变量的总效应
        X1 = data[ind_var]
        X1 = sm.add_constant(X1)
        model1 = sm.OLS(data[dep_var], X1).fit()
        total_effect = model1.params[ind_var]
        p_value_total = model1.pvalues[ind_var]

        # 第二步：自变量对中介变量的效应
        X2 = data[ind_var]
        X2 = sm.add_constant(X2)
        model2 = sm.OLS(data[med_var], X2).fit()
        effect_ind_med = model2.params[ind_var]
        p_value_ind_med = model2.pvalues[ind_var]

        # 第三步：调节变量对自变量 - 中介变量关系的调节效应
        data['ind_mod'] = data[ind_var] * data[mod_var]
        X3 = data[[ind_var, mod_var, 'ind_mod']]
        X3 = sm.add_constant(X3)
        model3 = sm.OLS(data[med_var], X3).fit()
        effect_mod_ind_med = model3.params['ind_mod']
        p_value_mod_ind_med = model3.pvalues['ind_mod']

        # 第四步：中介变量对因变量的效应（控制自变量）
        X4 = data[[ind_var, med_var]]
        X4 = sm.add_constant(X4)
        model4 = sm.OLS(data[dep_var], X4).fit()
        effect_med_dep = model4.params[med_var]
        p_value_med_dep = model4.pvalues[med_var]

        # 第五步：调节变量对中介变量 - 因变量关系的调节效应
        data['med_mod'] = data[med_var] * data[mod_var]
        X5 = data[[ind_var, med_var, mod_var, 'med_mod']]
        X5 = sm.add_constant(X5)
        model5 = sm.OLS(data[dep_var], X5).fit()
        effect_mod_med_dep = model5.params['med_mod']
        p_value_mod_med_dep = model5.pvalues['med_mod']

        # 第六步：中介效应
        mediation_effect = effect_ind_med * effect_med_dep

        sample_size = len(data)

        return total_effect, p_value_total, effect_ind_med, p_value_ind_med, effect_mod_ind_med, p_value_mod_ind_med, effect_med_dep, p_value_med_dep, effect_mod_med_dep, p_value_mod_med_dep, mediation_effect, sample_size

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 让用户输入自变量、中介变量、因变量和调节变量的列名
            ind_var = tkinter.simpledialog.askstring(languages[self.current_language]['input_info'],
                                                     languages[self.current_language]['input_ind_var'])
            med_var = tkinter.simpledialog.askstring(languages[self.current_language]['input_info'],
                                                     languages[self.current_language]['input_med_var'])
            dep_var = tkinter.simpledialog.askstring(languages[self.current_language]['input_info'],
                                                     languages[self.current_language]['input_dep_var'])
            mod_var = tkinter.simpledialog.askstring(languages[self.current_language]['input_info'],
                                                     languages[self.current_language]['input_mod_var'])

            if not ind_var or not med_var or not dep_var or not mod_var:
                self.result_label.config(text=languages[self.current_language]['input_incomplete'])
                return

            # 进行调节中介作用分析
            total_effect, p_value_total, effect_ind_med, p_value_ind_med, effect_mod_ind_med, p_value_mod_ind_med, effect_med_dep, p_value_med_dep, effect_mod_med_dep, p_value_mod_med_dep, mediation_effect, sample_size = self.moderated_mediation_analysis(
                df, ind_var, med_var, dep_var, mod_var)

            # 整理数据
            data = [
                ["自变量对因变量的总效应", total_effect, p_value_total],
                ["自变量对中介变量的效应", effect_ind_med, p_value_ind_med],
                ["调节变量对自变量 - 中介变量关系的调节效应", effect_mod_ind_med, p_value_mod_ind_med],
                ["中介变量对因变量的效应（控制自变量）", effect_med_dep, p_value_med_dep],
                ["调节变量对中介变量 - 因变量关系的调节效应", effect_mod_med_dep, p_value_mod_med_dep],
                ["中介效应", mediation_effect, ""],
                ["样本量", sample_size, ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('调节中介作用分析结果', 0)

                # 添加分析结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for index, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明
                doc.add_heading('统计量解释说明', level=1)
                for key, value in explanations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 添加分析结果解读
                doc.add_heading('统计量结果解读', level=1)
                for key, value in interpretations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 生成图片（中介效应柱状图）
                fig, ax = plt.subplots()
                effects = [total_effect, effect_ind_med, effect_mod_ind_med, effect_med_dep, effect_mod_med_dep,
                           mediation_effect]
                labels = ["自变量对因变量总效应", "自变量对中介变量效应", "调节变量对自变量 - 中介变量关系调节效应", "中介变量对因变量效应",
                          "调节变量对中介变量 - 因变量关系调节效应", "中介效应"] if self.current_language == 'zh' else [
                    "Total Effect of Independent on Dependent", "Effect of Independent on Mediator",
                    "Moderating Effect of Moderator on Independent - Mediator",
                    "Effect of Mediator on Dependent (Controlling Independent)",
                    "Moderating Effect of Moderator on Mediator - Dependent", "Mediation Effect"]
                ax.bar(labels, effects)
                ax.set_title('调节中介作用分析结果' if self.current_language == 'zh' else 'Moderated Mediation Analysis Results')
                ax.set_ylabel('效应值' if self.current_language == 'zh' else 'Effect Value')
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入 Word 文档
                doc.add_heading('中介效应柱状图', level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]['analysis_complete'].format(save_path), wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", wraplength=400)
        self.result_label.pack(pady=10)

        # 创建语言切换标签
        self.switch_language_label = ttk.Label(self.root, text=languages[self.current_language]["switch_language_button_text"], cursor="hand2")
        self.switch_language_label.pack(pady=10)
        self.switch_language_label.bind("<Button-1>", self.switch_language)


if __name__ == "__main__":
    app = ModeratedMediationAnalysisApp()
    app.root.mainloop()
