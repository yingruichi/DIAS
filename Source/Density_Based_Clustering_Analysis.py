import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from sklearn.cluster import DBSCAN
from sklearn.metrics import silhouette_score

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "密度聚类分析",
        "select_button": "选择文件",
        "analyze_button": "分析文件",
        "file_not_found": "文件不存在，请重新选择。",
        "analysis_success": "分析完成，结果已保存到 {}\n",
        "no_save_path": "未选择保存路径，结果未保存。",
        "analysis_error": "分析文件时出错: {}",
        "switch_language": "切换语言",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "explanation": {
            "聚类结果": "每个样本所属的聚类类别",
            "聚类散点图": "展示样本在二维空间中的分布及聚类结果",
            "轮廓系数": "衡量聚类效果的指标，值越接近1表示聚类效果越好"
        },
        "interpretation": {
            "聚类结果": "可用于区分不同样本所属的类别，-1 表示噪声点",
            "聚类散点图": "直观展示样本之间的聚类关系和分布情况",
            "轮廓系数": "若值接近1，说明聚类紧凑且分离度高；若值接近 -1，说明聚类效果差"
        }
    },
    "en": {
        "title": "Density-Based Clustering Analysis",
        "select_button": "Select File",
        "analyze_button": "Analyze File",
        "file_not_found": "The file does not exist. Please select again.",
        "analysis_success": "Analysis completed. The results have been saved to {}\n",
        "no_save_path": "No save path selected. The results were not saved.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "switch_language": "Switch Language",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "explanation": {
            "聚类结果": "The cluster label to which each sample belongs",
            "聚类散点图": "Show the distribution of samples in a two-dimensional space and the clustering results",
            "轮廓系数": "An index to measure the clustering effect. A value closer to 1 indicates better clustering."
        },
        "interpretation": {
            "聚类结果": "Can be used to distinguish the categories to which different samples belong. -1 represents noise points.",
            "聚类散点图": "Visually show the clustering relationship and distribution of samples.",
            "轮廓系数": "If the value is close to 1, the clusters are compact and well-separated. If close to -1, the clustering effect is poor."
        }
    }
}


class DensityBasedClusteringAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def density_clustering(self, data, eps=0.5, min_samples=5):
        """
        进行密度聚类分析
        :param data: 输入数据
        :param eps: 邻域的最大距离
        :param min_samples: 形成核心点所需的最小样本数
        :return: 聚类标签
        """
        model = DBSCAN(eps=eps, min_samples=min_samples)
        labels = model.fit_predict(data)
        return labels, model

    def plot_clustering(self, data, labels, save_path):
        """
        绘制聚类散点图
        :param data: 输入数据
        :param labels: 聚类标签
        :param save_path: 图片保存路径
        """
        plt.figure(figsize=(10, 5))
        unique_labels = set(labels)
        colors = [plt.cm.Spectral(each) for each in np.linspace(0, 1, len(unique_labels))]
        for k, col in zip(unique_labels, colors):
            if k == -1:
                # 噪声点用黑色表示
                col = [0, 0, 0, 1]
            class_member_mask = (labels == k)
            xy = data[class_member_mask]
            plt.plot(xy[:, 0], xy[:, 1], 'o', markerfacecolor=tuple(col),
                     markeredgecolor='k', markersize=6)
        plt.title('密度聚类散点图' if self.current_language == 'zh' else 'Density-Based Clustering Scatter Plot')
        plt.xlabel('特征1' if self.current_language == 'zh' else 'Feature 1')
        plt.ylabel('特征2' if self.current_language == 'zh' else 'Feature 2')
        img_path = os.path.splitext(save_path)[0] + '_density_clustering_scatter.png'
        plt.savefig(img_path)
        plt.close()
        return img_path

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            original_data = df.values  # 使用新变量存储原始数据

            # 进行密度聚类分析
            labels, model = self.density_clustering(original_data)

            # 计算轮廓系数
            try:
                silhouette_avg = silhouette_score(original_data, labels)
            except ValueError:
                silhouette_avg = "无法计算（可能只有一个聚类或所有样本为噪声点）"

            # 整理数据
            result_data = [
                ["聚类结果", labels.tolist(), ""],
                ["轮廓系数", silhouette_avg, ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(result_data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["聚类结果", "聚类散点图", "轮廓系数"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["聚类结果", "聚类散点图", "轮廓系数"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                document = Document()

                # 添加标题
                document.add_heading('密度聚类分析结果', 0)

                # 添加分析结果表格
                table = document.add_table(rows=1, cols=len(df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(df.columns):
                    hdr_cells[col_idx].text = header
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明表格
                document.add_heading('解释说明', level=1)
                table = document.add_table(rows=1, cols=len(explanation_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(explanation_df.columns):
                    hdr_cells[col_idx].text = header
                for _, row in explanation_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加结果解读表格
                document.add_heading('结果解读', level=1)
                table = document.add_table(rows=1, cols=len(interpretation_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(interpretation_df.columns):
                    hdr_cells[col_idx].text = header
                for _, row in interpretation_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 生成聚类散点图
                img_path = None
                if original_data.shape[1] >= 2:
                    img_path = self.plot_clustering(original_data, labels, save_path)
                    if img_path:
                        document.add_heading('聚类散点图', level=1)
                        document.add_picture(img_path)

                # 保存 Word 文档
                document.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                if img_path is None:
                    result_msg += "数据维度小于2，无法绘制散点图。"
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(languages[self.current_language]['title'])
        self.select_button.config(text=languages[self.current_language]['select_button'])
        self.analyze_button.config(text=languages[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=languages[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = DensityBasedClusteringAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()