import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
import statsmodels.api as sm
from statsmodels.genmod.families import Poisson
from statsmodels.genmod.cov_struct import Exchangeable
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "广义估计方程分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "广义估计方程": "用于处理具有相关性的纵向数据或聚类数据，能在考虑数据相关性的情况下估计回归系数。",
        },
        'interpretation': {
            "回归系数": "表示自变量对因变量的影响程度，系数的正负表示影响方向，绝对值大小表示影响强度。",
            "标准误": "衡量回归系数估计值的抽样误差大小，标准误越小，估计越精确。",
            "p值": "若 p 值小于显著性水平（通常为 0.05），则认为该自变量对因变量有显著影响。"
        }
    },
    'en': {
        'title': "Generalized Estimating Equations Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "广义估计方程": "Used to handle correlated longitudinal or clustered data, and can estimate regression coefficients while considering data correlation.",
        },
        'interpretation': {
            "回归系数": "Indicates the degree of influence of the independent variable on the dependent variable. The sign of the coefficient represents the direction of the influence, and the absolute value represents the strength of the influence.",
            "标准误": "Measures the sampling error of the regression coefficient estimate. A smaller standard error indicates a more precise estimate.",
            "p值": "If the p-value is less than the significance level (usually 0.05), it is considered that the independent variable has a significant influence on the dependent variable."
        }
    }
}


class GeneralizedEstimatingEquationsAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.configure(style="TEntry")

    def on_entry_click(self, event):
        """当用户点击输入框时，清除提示信息"""
        if self.file_entry.get() == LANGUAGES[self.current_language]['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.configure(style="TEntry")

    def on_focusout(self, event):
        """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]['file_entry_placeholder']:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设第一列是聚类标识，最后一列是因变量，其余列是自变量
            cluster_id = df.iloc[:, 0]
            y = df.iloc[:, -1]
            X = df.iloc[:, 1:-1]
            X = sm.add_constant(X)

            # 进行广义估计方程分析
            fam = Poisson()
            ind = Exchangeable()
            model = sm.GEE(y, X, groups=cluster_id, cov_struct=ind, family=fam)
            result = model.fit()

            # 提取结果
            summary = result.summary()
            summary_df = pd.DataFrame(summary.tables[1].data[1:], columns=summary.tables[1].data[0])

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["广义估计方程" if self.current_language == 'zh' else "Generalized Estimating Equations"])
            explanation_df.insert(0, "统计量", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=[
                "回归系数" if self.current_language == 'zh' else "Regression Coefficient",
                "标准误" if self.current_language == 'zh' else "Standard Error",
                "p值" if self.current_language == 'zh' else "p-value"
            ])
            interpretation_df.insert(0, "统计量", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                doc.add_heading('分析结果', level=1)
                table = doc.add_table(rows=1, cols=len(summary_df.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(summary_df.columns):
                    hdr_cells[i].text = col
                for _, row in summary_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明表格
                doc.add_heading('解释说明', level=1)
                table = doc.add_table(rows=1, cols=len(explanation_df.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(explanation_df.columns):
                    hdr_cells[i].text = col
                for _, row in explanation_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加结果解读表格
                doc.add_heading('结果解读', level=1)
                table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(interpretation_df.columns):
                    hdr_cells[i].text = col
                for _, row in interpretation_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 生成结果图片（回归系数可视化）
                plot_path = os.path.splitext(save_path)[0] + '_gee_regression_coefficients.png'
                plt.figure()
                coefs = result.params[1:]
                variables = X.columns[1:]
                plt.bar(variables, coefs)
                plt.xlabel('自变量' if self.current_language == 'zh' else 'Independent Variables')
                plt.ylabel('回归系数' if self.current_language == 'zh' else 'Regression Coefficients')
                plt.title(
                    '广义估计方程回归系数' if self.current_language == 'zh' else 'Generalized Estimating Equations Regression Coefficients')
                plt.xticks(rotation=45)
                plt.savefig(plot_path)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_heading('回归系数可视化', level=1)
                doc.add_picture(plot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.configure(style="Gray.TEntry")

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建文件选择按钮
        self.select_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file,
                                        bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(self.root, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack(pady=5)
        self.file_entry.bind("<FocusIn>", self.on_entry_click)
        self.file_entry.bind("<FocusOut>", self.on_focusout)

        # 创建分析按钮
        self.analyze_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file,
                                         bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建语言切换标签
        self.switch_language_label = ttk.Label(self.root, text=LANGUAGES[self.current_language]['switch_language'],
                                               cursor="hand2")
        self.switch_language_label.pack(pady=10)
        self.switch_language_label.bind("<Button-1>", self.switch_language)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = GeneralizedEstimatingEquationsAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()