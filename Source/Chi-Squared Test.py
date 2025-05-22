import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "卡方检验",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'explanation': {
            "Pearson卡方": "当2*2列联表中n>=40且所有期望频数E>=5时使用，衡量实际频数与理论频数的差异。",
            "Yates校正卡方": "当2*2列联表中n>=40但有一个格子的期望频数满足1<=E<5时使用，对Pearson卡方的校正。",
            "似然比卡方": "当R*C列联表中期望频数不满足使用Pearson卡方的条件时使用。",
            "Fisher卡方": "当2*2列联表中任何一格子出现E<1或n<40时使用。",
            "Phi系数": "用于衡量2*2列联表的效应量。",
            "Cramer's V": "用于衡量R*C列联表的效应量。",
            "趋势卡方": "用于检验变量之间是否存在线性趋势。"
        },
        'interpretation': {
            "卡方值": "卡方值越大，说明实际频数与理论频数之间的差异越大。",
            "p值": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为变量之间存在显著关联；否则，接受原假设，认为变量之间无显著关联。",
            "自由度": "自由度反映了数据的独立变化程度，用于计算卡方分布的临界值。",
            "显著性（α=0.05）": "表示在0.05的显著性水平下，变量之间是否存在显著关联。",
            "校正p值（Bonferroni）": "经过Bonferroni校正后的p值，用于多重比较，校正后的p值小于显著性水平时，拒绝原假设。",
            "Phi系数": "Phi系数的绝对值越接近1，说明2*2列联表中两个变量之间的关联越强。",
            "Cramer's V": "Cramer's V的值越接近1，说明R*C列联表中两个变量之间的关联越强。",
            "趋势卡方值": "趋势卡方值越大，说明变量之间存在线性趋势的可能性越大。",
            "趋势卡方p值": "趋势卡方p值小于显著性水平时，说明变量之间存在线性趋势；否则，不存在线性趋势。"
        }
    },
    'en': {
        'title': "Chi-square test",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'explanation': {
            "Pearson卡方": "Used when n>=40 and all expected frequencies E>=5 in a 2*2 contingency table, measuring the difference between observed and expected frequencies.",
            "Yates校正卡方": "Used when n>=40 but there is one cell with 1<=E<5 in a 2*2 contingency table, a correction to the Pearson chi-square.",
            "似然比卡方": "Used when the expected frequencies in an R*C contingency table do not meet the conditions for using the Pearson chi-square.",
            "Fisher卡方": "Used when any cell has E<1 or n<40 in a 2*2 contingency table.",
            "Phi系数": "Used to measure the effect size of a 2*2 contingency table.",
            "Cramer's V": "Used to measure the effect size of an R*C contingency table.",
            "趋势卡方": "Used to test if there is a linear trend between variables."
        },
        'interpretation': {
            "卡方值": "A larger chi-square value indicates a greater difference between the observed and expected frequencies.",
            "p值": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant association between variables; otherwise, the null hypothesis is accepted, indicating no significant association.",
            "自由度": "The degrees of freedom reflect the independent variation of the data and are used to calculate the critical value of the chi-square distribution.",
            "显著性（α=0.05）": "Indicates whether there is a significant association between variables at the 0.05 significance level.",
            "校正p值（Bonferroni）": "The p-value after Bonferroni correction, used for multiple comparisons. When the corrected p-value is less than the significance level, the null hypothesis is rejected.",
            "Phi系数": "The closer the absolute value of the Phi coefficient is to 1, the stronger the association between the two variables in the 2*2 contingency table.",
            "Cramer's V": "The closer the value of Cramer's V is to 1, the stronger the association between the two variables in the R*C contingency table.",
            "趋势卡方值": "A larger trend chi-square value indicates a greater possibility of a linear trend between variables.",
            "趋势卡方p值": "When the trend chi-square p-value is less than the significance level, there is a linear trend between variables; otherwise, there is no linear trend."
        }
    }
}

# 当前语言
current_language = 'en'


def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        file_entry.configure(style="TEntry")  # 恢复默认样式


def analyze_file():
    global current_language
    file_path = file_entry.get()
    if file_path == "请输入待分析 Excel 文件的完整路径" or file_path == "Please enter the full path of the Excel file to be analyzed":
        file_path = ""
    if not os.path.exists(file_path):
        result_label.config(text=LANGUAGES[current_language]['file_not_found'])
        return
    try:
        # 打开 Excel 文件
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active

        rows = sheet.max_row
        columns = sheet.max_column

        data = []
        chi_types = ["Pearson卡方", "Yates校正卡方", "似然比卡方", "Fisher卡方",
                     "趋势卡方"]
        explanations = LANGUAGES[current_language]['explanation']
        interpretations = LANGUAGES[current_language]['interpretation']

        # 读取数据
        table = []
        for row_idx in range(1, rows + 1):
            row = []
            for col_idx in range(1, columns + 1):
                cell_value = sheet.cell(row=row_idx, column=col_idx).value
                if isinstance(cell_value, (int, float)):
                    row.append(cell_value)
            if row:
                table.append(row)
        table = np.array(table)

        n = table.sum()
        R, C = table.shape

        # 计算百分比
        total_percentage = table / n * 100
        percentage_df = pd.DataFrame(total_percentage, columns=[f'列 {i + 1} 百分比' for i in range(C)],
                                     index=[f'行 {i + 1} 百分比' for i in range(R)])

        chi_type = None
        chi2 = None
        p = None
        dof = None
        expected = None
        chi_choice_reason = ""

        if R == 2 and C == 2:
            # 2*2列联表
            chi2, p, dof, expected = stats.chi2_contingency(table)
            if n >= 40 and np.all(expected >= 5):
                chi_type = chi_types[0]
                chi2, p = stats.chi2_contingency(table)[:2]
                chi_choice_reason = "由于是2*2列联表，且样本量n>=40，所有期望频数E>=5，所以选择Pearson卡方检验。" if current_language == 'zh' else "Since it is a 2*2 contingency table, the sample size n>=40, and all expected frequencies E>=5, the Pearson chi-square test is selected."
            elif n >= 40 and np.any((expected >= 1) & (expected < 5)):
                chi_type = chi_types[1]
                chi2, p = stats.chi2_contingency(table, correction=True)[:2]
                chi_choice_reason = "由于是2*2列联表，样本量n>=40，但有一个格子的期望频数满足1<=E<5，所以选择Yates校正卡方检验。" if current_language == 'zh' else "Since it is a 2*2 contingency table, the sample size n>=40, but there is one cell with an expected frequency satisfying 1<=E<5, the Yates corrected chi-square test is selected."
            else:
                chi_type = chi_types[3]
                odds_ratio, p = stats.fisher_exact(table)
                chi2 = None
                chi_choice_reason = "由于是2*2列联表，任何一格子出现E<1或n<40，所以选择Fisher卡方检验。" if current_language == 'zh' else "Since it is a 2*2 contingency table and any cell has E<1 or n<40, the Fisher chi-square test is selected."

            # 计算 Phi 系数
            if chi2 is not None:
                phi = np.sqrt(chi2 / n)
                data.append(["Phi系数（phi/varphi）", phi])

            # 计算卡方检验统计量过程值
            if chi2 is not None:
                chi2_process = ((table - expected) ** 2 / expected).sum()
                data.append(["Pearson卡方统计量过程值", chi2_process])

        else:
            # R*C列联表
            chi2, p, dof, expected = stats.chi2_contingency(table)
            if np.all(expected > 1) and np.sum((expected >= 1) & (expected < 5)) / (R * C) < 0.2:
                chi_type = chi_types[0]
                chi_choice_reason = "由于是R*C列联表，所有期望频数E>1，且期望频数满足1<=E<5的格子占比小于20%，所以选择Pearson卡方检验。" if current_language == 'zh' else "Since it is an R*C contingency table, all expected frequencies E>1, and the proportion of cells with expected frequencies satisfying 1<=E<5 is less than 20%, the Pearson chi-square test is selected."
            else:
                chi_type = chi_types[2]
                g, p = stats.chi2_contingency(table, lambda_="log-likelihood")[:2]
                chi2 = g
                chi_choice_reason = "由于是R*C列联表，期望频数不满足使用Pearson卡方的条件，所以选择似然比卡方检验。" if current_language == 'zh' else "Since it is an R*C contingency table and the expected frequencies do not meet the conditions for using the Pearson chi-square, the likelihood ratio chi-square test is selected."

            # 计算 Cramer's V
            if min(R, C) > 1:  # 确保分母不为零
                v = np.sqrt(chi2 / (n * (min(R, C) - 1)))
                data.append(["Cramer's V", v])
            else:
                data.append(["Cramer's V", "无法计算（R或C等于1）"])

            # 计算卡方检验统计量过程值
            chi2_process = ((table - expected) ** 2 / expected).sum()
            data.append(["Pearson卡方统计量过程值", chi2_process])

        # 趋势卡方检验
        if R > 1 and C > 1:
            scores = np.arange(1, C + 1)
            row_sums = table.sum(axis=1)
            col_sums = table.sum(axis=0)
            total_sum = table.sum()
            expected_values = np.outer(row_sums, col_sums) / total_sum
            chi2_trend = ((table - expected_values) ** 2 / expected_values).sum()
            p_trend = 1 - stats.chi2.cdf(chi2_trend, (R - 1) * (C - 1))
            data.append(["趋势卡方值（Trend χ²）", chi2_trend])
            data.append(["趋势卡方p值", p_trend])

        data.append(["卡方检验类型", chi_type])
        if chi2 is not None:
            data.append(["卡方值（χ²）", chi2])
        data.append(["自由度", dof])  # 添加自由度信息
        data.append(["p值", p])

        # 显著性判断
        alpha = 0.05
        significant = p < alpha
        data.append(["显著性（α=0.05）", "显著" if significant else "不显著"])

        # 多重比较（Bonferroni 校正）
        corrected_p = p * (R * C - 1)  # 假设进行了 R*C - 1 次比较
        data.append(["校正p值（Bonferroni）", corrected_p])

        headers = ["统计量", "值"]
        df = pd.DataFrame(data, columns=headers)

        # 添加解释说明
        explanation_df = pd.DataFrame([explanations])
        explanation_df = explanation_df.reindex(columns=chi_types + ["Phi系数（phi/varphi）", "Cramer's V"])
        explanation_df.insert(0, "统计量", "解释说明" if current_language == 'zh' else "Explanation")

        # 添加分析结果解读
        interpretation_df = pd.DataFrame([interpretations])
        interpretation_df = interpretation_df.reindex(columns=[
            "卡方值（χ²）", "p值", "自由度", "显著性（α=0.05）", "校正p值（Bonferroni）",
            "Phi系数（phi/varphi）", "Cramer's V", "趋势卡方值（Trend χ²）", "趋势卡方p值"
        ])
        interpretation_df.insert(0, "统计量", "结果解读" if current_language == 'zh' else "Interpretation")

        # 添加选择卡方检验方法的原因
        reason_df = pd.DataFrame([{"统计量": "选择卡方检验方法的原因", "值": chi_choice_reason}])

        # 合并数据、解释说明和结果解读
        combined_df = pd.concat([df, explanation_df, interpretation_df, reason_df], ignore_index=True)

        # 转置数据框
        transposed_df = combined_df.set_index('统计量').T.reset_index().rename(columns={'index': '统计量'})

        # 合并百分比数据
        percentage_df = percentage_df.reset_index().rename(columns={'index': '统计量'})
        final_df = pd.concat([transposed_df, percentage_df], ignore_index=True)

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建 Word 文档
            doc = Document()

            # 添加标题
            doc.add_heading('卡方检验分析结果', 0)

            # 添加表格
            table = doc.add_table(rows=final_df.shape[0]+1, cols=final_df.shape[1])
            hdr_cells = table.rows[0].cells
            for col_idx, col_name in enumerate(final_df.columns):
                hdr_cells[col_idx].text = col_name

            for row_idx, row in final_df.iterrows():
                row_cells = table.rows[row_idx+1].cells
                for col_idx, value in enumerate(row):
                    row_cells[col_idx].text = str(value)

            # 生成卡方交叉图
            desktop_path = pathlib.Path.home() / 'Desktop'
            plot_path = desktop_path / 'chi_square_plot.png'
            plt.figure(figsize=(8, 6))
            plt.imshow(table, cmap='YlGnBu', interpolation='nearest')
            plt.colorbar()
            for i in range(table.shape[0]):
                for j in range(table.shape[1]):
                    plt.text(j, i, str(table[i, j]), ha='center', va='center', color='black')
            plt.title('Chi-Square Contingency Table')
            plt.xlabel('Columns')
            plt.ylabel('Rows')
            plt.savefig(plot_path)
            plt.close()

            # 将图片插入到 Word 文档中
            doc.add_picture(str(plot_path), width=Inches(6))

            # 保存 Word 文档
            doc.save(save_path)

            result_msg = LANGUAGES[current_language]['analysis_success'].format(
                save_path) + f"卡方交叉图已保存到 {plot_path}"
            result_label.config(text=result_msg, wraplength=400)
        else:
            result_label.config(text=LANGUAGES[current_language]['no_save_path'])

    except Exception as e:
        result_label.config(text=LANGUAGES[current_language]['analysis_error'].format(str(e)))


def switch_language():
    global current_language
    current_language = 'en' if current_language == 'zh' else 'zh'
    root.title(LANGUAGES[current_language]['title'])
    select_button.config(text=LANGUAGES[current_language]['select_button'])
    analyze_button.config(text=LANGUAGES[current_language]['analyze_button'])
    language_label.config(text=LANGUAGES[current_language]['switch_language'])
    # 切换语言时更新提示信息
    file_entry.delete(0, tk.END)
    if current_language == 'zh':
        file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
        file_entry.configure(style="Gray.TEntry")
    else:
        file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
        file_entry.configure(style="Gray.TEntry")


def on_entry_click(event):
    """当用户点击输入框时，清除提示信息"""
    if file_entry.get() == "请输入待分析 Excel 文件的完整路径" or file_entry.get() == "Please enter the full path of the Excel file to be analyzed":
        file_entry.delete(0, tk.END)
        file_entry.configure(style="TEntry")  # 恢复默认样式


def on_focusout(event):
    """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
    if file_entry.get() == "":
        if current_language == 'zh':
            file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
            file_entry.configure(style="Gray.TEntry")
        else:
            file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
            file_entry.configure(style="Gray.TEntry")


# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(LANGUAGES[current_language]['title'])

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 设置窗口的宽度和高度
window_width = 500
window_height = 300

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建自定义样式
style = ttk.Style()
style.configure("Gray.TEntry", foreground="gray")

# 创建文件选择按钮
select_button = ttk.Button(root, text=LANGUAGES[current_language]['select_button'], command=select_file,
                           bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建文件路径输入框
file_entry = ttk.Entry(root, width=50, style="Gray.TEntry")
if current_language == 'zh':
    file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
else:
    file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
file_entry.pack(pady=5)
file_entry.bind("<FocusIn>", on_entry_click)
file_entry.bind("<FocusOut>", on_focusout)

# 创建分析按钮
analyze_button = ttk.Button(root, text=LANGUAGES[current_language]['analyze_button'], command=analyze_file,
                            bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建语言切换标签
language_label = ttk.Label(root, text=LANGUAGES[current_language]['switch_language'], cursor="hand2")
language_label.pack(pady=10)
language_label.bind("<Button-1>", lambda event: switch_language())

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
root.mainloop()