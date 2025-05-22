import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import openpyxl
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "极差分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "switch_language_button_text": "切换语言",
        "explanation": {
            "极差": "极差反映了数据的离散程度，在极差分析中，极差越大说明该因素对试验结果的影响越大。"
        },
        "interpretation": {
            "极差": "极差越大，表明该因素对试验结果的影响越显著。",
            "均值": "各水平下试验结果的平均值，用于比较不同水平对试验结果的影响。"
        }
    },
    "en": {
        "title": "Range Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "switch_language_button_text": "Switch Language",
        "explanation": {
            "极差": "The range reflects the dispersion of the data. In range analysis, a larger range indicates that the factor has a greater influence on the test results."
        },
        "interpretation": {
            "极差": "A larger range indicates that the factor has a more significant influence on the test results.",
            "均值": "The average value of the test results at each level, used to compare the influence of different levels on the test results."
        }
    }
}

class RangeAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            workbook = openpyxl.load_workbook(file_path)
            sheet = workbook.active

            rows = sheet.max_row
            columns = sheet.max_column

            data = []
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 读取数据
            table = []
            for row_idx in range(1, rows + 1):
                row = []
                for col_idx in range(1, columns + 1):
                    cell_value = sheet.cell(row=row_idx, column=col_idx).value
                    if isinstance(cell_value, (int, float)):
                        row.append(cell_value)
                if row:
                    table.append(row)
            table = np.array(table)

            num_factors = columns - 1  # 减去结果列
            levels = np.unique(table[:, :-1], axis=0).shape[0]
            results = table[:, -1]

            # 计算各因素各水平下的均值和极差
            factor_means = []
            factor_ranges = []
            for factor in range(num_factors):
                level_means = []
                for level in np.unique(table[:, factor]):
                    level_results = results[table[:, factor] == level]
                    level_mean = np.mean(level_results)
                    level_means.append(level_mean)
                factor_means.append(level_means)
                factor_range = np.max(level_means) - np.min(level_means)
                factor_ranges.append(factor_range)

            # 整理数据
            for i in range(num_factors):
                for j in range(levels):
                    data.append([f"因素{i + 1} 水平{j + 1} 均值", factor_means[i][j]])
                data.append([f"因素{i + 1} 极差", factor_ranges[i]])

            headers = ["统计量", "值"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["极差"])
            explanation_df.insert(0, "统计量", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["极差", "均值"])
            interpretation_df.insert(0, "统计量", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df, explanation_df, interpretation_df], ignore_index=True)

            # 转置数据框
            transposed_df = combined_df.set_index('统计量').T.reset_index().rename(columns={'index': '统计量'})

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表格
                table = doc.add_table(rows=transposed_df.shape[0], cols=transposed_df.shape[1])
                for i, row in enumerate(transposed_df.values):
                    for j, value in enumerate(row):
                        table.cell(i, j).text = str(value)

                # 生成极差分析图
                desktop_path = pathlib.Path.home() / 'Desktop'
                plot_path = os.path.splitext(save_path)[0] + '_range_analysis_plot.png'
                plt.figure(figsize=(10, 6))
                for i in range(num_factors):
                    plt.plot(np.arange(1, levels + 1), factor_means[i], marker='o', label=f'Factor {i + 1}')
                plt.title('Range Analysis - Mean Values by Factor and Level')
                plt.xlabel('Level')
                plt.ylabel('Mean Value')
                plt.legend()
                plt.grid(True)
                plt.savefig(plot_path)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_picture(plot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_complete'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = RangeAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()