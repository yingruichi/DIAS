import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import tkinter.simpledialog
import pingouin as pg
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    "zh": {
        "title": "组内评分者信度rwg分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "switch_language_button_text": "切换语言",
        "group_column_prompt": "请输入分组列的列名",
        "rating_column_prompt": "请输入评分列的列名（点击取消结束输入）",
        "no_group_column": "未输入有效的分组列名，分析取消。",
        "no_rating_columns": "未输入有效的评分列名，分析取消。",
        "empty_rating_column": "输入的列名不能为空，请重新输入。",
        "group_less_than_2": "组 {} 的样本数少于2，跳过...",
        "explanation": {
            "rwg值": "组内评分者信度rwg用于评估组内成员评分的一致性，值越接近1表示一致性越高。",
            "Rwg值标准差SD": "Rwg值的标准差，反映了Rwg值的离散程度。",
            "P25": "Rwg值的第25百分位数。",
            "中位数": "Rwg值的中位数。",
            "P75": "Rwg值的第75百分位数。",
            "ICC1": "组内相关系数1，用于衡量组内评分者之间的一致性。",
            "ICC2": "组内相关系数2，考虑了评分者和项目的交互作用。",
            "MSB": "组间均方，反映了组间差异。",
            "MSW": "组内均方，反映了组内差异。",
            "F值": "F检验统计量，用于检验组间差异是否显著。",
            "p值": "F检验的p值，用于判断组间差异是否显著。"
        },
        "interpretation": {
            "rwg值": "rwg值越接近1，说明组内成员的评分越一致；值越低，说明组内成员的评分差异越大。",
            "Rwg值标准差SD": "标准差越大，说明Rwg值的离散程度越大。",
            "P25": "第25百分位数较低表示有25%的Rwg值低于该值。",
            "中位数": "中位数反映了Rwg值的中间水平。",
            "P75": "第75百分位数较高表示有75%的Rwg值低于该值。",
            "ICC1": "ICC1值越接近1，组内评分者之间的一致性越高。",
            "ICC2": "ICC2值越接近1，考虑交互作用后组内评分者之间的一致性越高。",
            "MSB": "MSB值越大，组间差异越明显。",
            "MSW": "MSW值越大，组内差异越明显。",
            "F值": "F值越大，说明组间差异越可能显著。",
            "p值": "p值小于0.05时，说明组间差异显著。"
        }
    },
    "en": {
        "title": "Within-Group Inter-Rater Reliability rwg Analysis",
        "select_button_text": "Select Files",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze Files",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "switch_language_button_text": "Switch Language",
        "group_column_prompt": "Please enter the name of the grouping column",
        "rating_column_prompt": "Please enter the name of the rating column (click Cancel to finish input)",
        "no_group_column": "No valid grouping column name entered. Analysis canceled.",
        "no_rating_columns": "No valid rating column names entered. Analysis canceled.",
        "empty_rating_column": "The column name entered cannot be empty. Please re-enter.",
        "group_less_than_2": "Group {} has less than 2 samples. Skipping...",
        "explanation": {
            "rwg值": "The within-group inter-rater reliability rwg is used to evaluate the consistency of ratings within a group. A value closer to 1 indicates higher consistency.",
            "Rwg值标准差SD": "The standard deviation of the rwg values, reflecting the dispersion of the rwg values.",
            "P25": "The 25th percentile of the rwg values.",
            "中位数": "The median of the rwg values.",
            "P75": "The 75th percentile of the rwg values.",
            "ICC1": "Intraclass correlation coefficient 1, used to measure the consistency between raters within a group.",
            "ICC2": "Intraclass correlation coefficient 2, considering the interaction between raters and items.",
            "MSB": "Mean square between groups, reflecting the differences between groups.",
            "MSW": "Mean square within groups, reflecting the differences within groups.",
            "F值": "F-test statistic, used to test whether the differences between groups are significant.",
            "p值": "The p-value of the F-test, used to determine whether the differences between groups are significant."
        },
        "interpretation": {
            "rwg值": "The closer the rwg value is to 1, the more consistent the ratings within the group; the lower the value, the greater the difference in ratings within the group.",
            "Rwg值标准差SD": "A larger standard deviation indicates a greater dispersion of the rwg values.",
            "P25": "A lower 25th percentile means that 25% of the rwg values are below this value.",
            "中位数": "The median reflects the middle level of the rwg values.",
            "P75": "A higher 75th percentile means that 75% of the rwg values are below this value.",
            "ICC1": "The closer the ICC1 value is to 1, the higher the consistency between raters within the group.",
            "ICC2": "The closer the ICC2 value is to 1, the higher the consistency between raters within the group considering the interaction.",
            "MSB": "A larger MSB value indicates more obvious differences between groups.",
            "MSW": "A larger MSW value indicates more obvious differences within groups.",
            "F值": "A larger F value indicates that the differences between groups are more likely to be significant.",
            "p值": "When the p-value is less than 0.05, the differences between groups are significant."
        }
    }
}


class WithinGroupInterRaterReliabilityRwgAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_paths = filedialog.askopenfilenames(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_paths:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, ", ".join(file_paths))
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def rwg_analysis(self, data, group_column, rating_columns):
        all_results = []
        rwg_values = []
        groups = data[group_column].unique()

        for group in groups:
            group_data = data[data[group_column] == group][rating_columns]
            # 检查数据是否为空或者只有一个样本
            if group_data.shape[0] < 2:
                print(languages[self.current_language]["group_less_than_2"].format(group))
                continue
            k = group_data.shape[1]  # 评分者数量
            n = group_data.shape[0]  # 项目数量
            var_within = group_data.var(axis=1).mean()
            expected_var = (k ** 2 - 1) / 12
            rwg = 1 - (var_within / expected_var)
            rwg_values.append(rwg)
            result = {
                f"{group}_rwg值": rwg
            }
            all_results.append(result)

        # 计算 Rwg 值的统计量
        if rwg_values:
            rwg_sd = np.std(rwg_values)
            rwg_p25 = np.percentile(rwg_values, 25)
            rwg_median = np.median(rwg_values)
            rwg_p75 = np.percentile(rwg_values, 75)

            # 计算 ICC1 和 ICC2
            icc_data = pd.melt(data, id_vars=[group_column], value_vars=rating_columns)
            icc_data.columns = ['Group', 'Rater', 'Score']
            icc = pg.intraclass_corr(data=icc_data, targets='Group', raters='Rater', ratings='Score')
            icc1 = icc[icc['Type'] == 'ICC1']['ICC'].values[0]
            icc2 = icc[icc['Type'] == 'ICC2']['ICC'].values[0]

            # 计算 MSB, MSW, F 值, p 值
            anova = pg.anova(data=icc_data, dv='Score', between='Group')
            msb = anova['MS'][0]
            msw = anova['MS'][1]
            f_value = anova['F'][0]
            p_value = anova['p-unc'][0]

            additional_stats = {
                "Rwg值标准差SD": rwg_sd,
                "P25": rwg_p25,
                "中位数": rwg_median,
                "P75": rwg_p75,
                "ICC1": icc1,
                "ICC2": icc2,
                "MSB": msb,
                "MSW": msw,
                "F值": f_value,
                "p值": p_value
            }
            all_results.append(additional_stats)

        return all_results, rwg_values

    def analyze_file(self):
        file_paths = self.file_entry.get().split(", ")
        if file_paths[0] == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return

        for file_path in file_paths:
            if not os.path.exists(file_path):
                self.result_label.config(text=languages[self.current_language]["file_not_exists"])
                return

        try:
            # 获取分组列名
            group_column = tkinter.simpledialog.askstring("输入信息",
                                                          languages[self.current_language]["group_column_prompt"])
            if not group_column:
                self.result_label.config(text=languages[self.current_language]["no_group_column"])
                return

            # 获取评分列名
            rating_columns = []
            while True:
                rating_column = tkinter.simpledialog.askstring("输入信息",
                                                               languages[self.current_language]["rating_column_prompt"])
                if rating_column is None:
                    break
                if rating_column.strip():
                    rating_columns.append(rating_column.strip())
                else:
                    print(languages[self.current_language]["empty_rating_column"])

            if not rating_columns:
                self.result_label.config(text=languages[self.current_language]["no_rating_columns"])
                return

            all_results = []
            file_names = []
            rwg_values_all = []
            group_names_all = []

            for file_path in file_paths:
                # 打开 Excel 文件
                df = pd.read_excel(file_path)

                # 进行rwg分析
                rwg_results, rwg_values = self.rwg_analysis(df, group_column, rating_columns)
                all_results.extend(rwg_results)
                file_names.extend([os.path.basename(file_path)] * len(rwg_results))

                # 收集rwg值和组名用于绘图
                if rwg_values:
                    for i, result in enumerate(rwg_results):
                        if i < len(rwg_results) - 1:  # 最后一个是统计结果
                            key = list(result.keys())[0]
                            group = key.split('_')[0]
                            rwg_values_all.append(result[key])
                            group_names_all.append(f"{os.path.basename(file_path)}_{group}")

            # 整理数据
            all_data = []
            for i, results in enumerate(all_results):
                if results:
                    data = []
                    for key, value in results.items():
                        data.append([f"{file_names[i]}_{key}", value])
                    all_data.extend(data)

            headers = ["指标", "数值"]
            df_result = pd.DataFrame(all_data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["rwg值", "Rwg值标准差SD", "P25", "中位数", "P75", "ICC1", "ICC2", "MSB", "MSW", "F值", "p值"])
            explanation_df.insert(0, "指标_解释说明",
                                  "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["rwg值", "Rwg值标准差SD", "P25", "中位数", "P75", "ICC1", "ICC2", "MSB", "MSW", "F值", "p值"])
            interpretation_df.insert(0, "指标_结果解读",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word files", "*.docx")])

            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加分析结果表格
                table = doc.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = headers[0]
                hdr_cells[1].text = headers[1]
                for index, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row[headers[0]])
                    row_cells[1].text = str(row[headers[1]])

                # 添加解释说明表格
                doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", level=2)
                table = doc.add_table(rows=1, cols=len(explanation_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(explanation_df.columns):
                    hdr_cells[col_idx].text = col_name
                for index, row in explanation_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, col_name in enumerate(explanation_df.columns):
                        row_cells[col_idx].text = str(row[col_name])

                # 添加分析结果解读表格
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Interpretation", level=2)
                table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(interpretation_df.columns):
                    hdr_cells[col_idx].text = col_name
                for index, row in interpretation_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, col_name in enumerate(interpretation_df.columns):
                        row_cells[col_idx].text = str(row[col_name])

                # 生成rwg值柱状图
                if rwg_values_all:
                    fig, ax = plt.subplots(figsize=(10, 6))
                    ax.bar(group_names_all, rwg_values_all)
                    ax.set_title(
                        '组内评分者信度rwg值' if self.current_language == 'zh' else 'Within-Group Inter-Rater Reliability rwg Values')
                    ax.set_ylabel('rwg值' if self.current_language == 'zh' else 'rwg Value')
                    ax.set_xlabel('分组' if self.current_language == 'zh' else 'Group')
                    plt.xticks(rotation=45, ha='right')
                    plt.tight_layout()

                    # 保存图片
                    img_path = os.path.splitext(save_path)[0] + '_rwg.png'
                    plt.savefig(img_path)
                    plt.close()

                    # 将图片插入 Word 文档
                    doc.add_heading(
                        "组内评分者信度rwg值柱状图" if self.current_language == 'zh' else "Within-Group Inter-Rater Reliability rwg Values Bar Chart",
                        level=2)
                    doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = WithinGroupInterRaterReliabilityRwgAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()