import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "秩和比(RSR)分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "switch_language_button_text": "切换语言",
        "explanation": {
            "秩矩阵": "将原始数据转换为秩次后得到的矩阵",
            "秩和比(RSR)": "反映各评价对象综合水平的统计量",
            "RSR 分布直方图": "展示 RSR 值分布情况的直方图",
            "回归方程": "用于拟合 RSR 值与概率单位之间关系的方程",
            "RSR 排序结果": "根据 RSR 值对各评价对象进行排序的结果"
        },
        "interpretation": {
            "秩矩阵": "便于后续计算秩和比",
            "秩和比(RSR)": "值越大，综合水平越高",
            "RSR 分布直方图": "直观观察 RSR 值的分布特征",
            "回归方程": "用于进一步分析 RSR 值的变化规律",
            "RSR 排序结果": "排名越靠前，综合水平越高"
        }
    },
    "en": {
        "title": "Rank - Sum Ratio (RSR) Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "switch_language_button_text": "Switch Language",
        "explanation": {
            "秩矩阵": "The matrix obtained after converting the original data into ranks",
            "秩和比(RSR)": "A statistic reflecting the comprehensive level of each evaluation object",
            "RSR 分布直方图": "A histogram showing the distribution of RSR values",
            "回归方程": "An equation used to fit the relationship between RSR values and probability units",
            "RSR 排序结果": "The result of ranking each evaluation object according to RSR values"
        },
        "interpretation": {
            "秩矩阵": "Facilitate the subsequent calculation of the rank - sum ratio",
            "秩和比(RSR)": "The larger the value, the higher the comprehensive level",
            "RSR 分布直方图": "Visually observe the distribution characteristics of RSR values",
            "回归方程": "Used to further analyze the change rule of RSR values",
            "RSR 排序结果": "The higher the ranking, the higher the comprehensive level"
        }
    }
}

class RankSumRatioRSRAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def rsr_method(self, data):
        """
        实现秩和比(RSR)法
        :param data: 原始数据矩阵
        :return: 秩矩阵, 秩和比(RSR), RSR 排序结果
        """
        # 计算秩矩阵
        rank_matrix = np.apply_along_axis(lambda x: pd.Series(x).rank().values, 0, data)

        # 计算秩和比(RSR)
        RSR = rank_matrix.sum(axis=1) / (rank_matrix.shape[0] * rank_matrix.shape[1])

        # 对 RSR 进行排序
        ranking = np.argsort(-RSR) + 1

        return rank_matrix, RSR, ranking

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 将数据转换为浮点类型
            data = data.astype(float)

            # 进行 RSR 分析
            rank_matrix, RSR, ranking = self.rsr_method(data)

            # 整理数据
            result_data = [
                ["秩矩阵", rank_matrix.tolist(), ""],
                ["秩和比(RSR)", RSR.tolist(), ""],
                ["RSR 排序结果", ranking.tolist(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            result_df = pd.DataFrame(result_data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["秩矩阵", "秩和比(RSR)", "RSR 分布直方图", "回归方程", "RSR 排序结果"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["秩矩阵", "秩和比(RSR)", "RSR 分布直方图", "回归方程", "RSR 排序结果"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 生成 RSR 分布直方图
            fig, ax = plt.subplots()
            ax.hist(RSR, bins=10)
            ax.set_title(
                'RSR 分布直方图' if self.current_language == 'zh' else 'Histogram of RSR Distribution')
            ax.set_xlabel('秩和比(RSR)' if self.current_language == 'zh' else 'Rank - Sum Ratio (RSR)')
            ax.set_ylabel('频数' if self.current_language == 'zh' else 'Frequency')

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header
                for _, row in result_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明表格
                doc.add_heading('解释说明', level=2)
                exp_table = doc.add_table(rows=1, cols=len(explanation_df.columns))
                exp_hdr_cells = exp_table.rows[0].cells
                for col_idx, header in enumerate(explanation_df.columns):
                    exp_hdr_cells[col_idx].text = header
                for _, row in explanation_df.iterrows():
                    exp_row_cells = exp_table.add_row().cells
                    for col_idx, value in enumerate(row):
                        exp_row_cells[col_idx].text = str(value)

                # 添加结果解读表格
                doc.add_heading('结果解读', level=2)
                int_table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
                int_hdr_cells = int_table.rows[0].cells
                for col_idx, header in enumerate(interpretation_df.columns):
                    int_hdr_cells[col_idx].text = header
                for _, row in interpretation_df.iterrows():
                    int_row_cells = int_table.add_row().cells
                    for col_idx, value in enumerate(row):
                        int_row_cells[col_idx].text = str(value)

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_rsr_histogram.png'
                plt.savefig(img_path)
                plt.close()

                # 添加图片到 Word 文档
                doc.add_heading('RSR 分布直方图', level=2)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_complete'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = RankSumRatioRSRAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()