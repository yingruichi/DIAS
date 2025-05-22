import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import pathlib
import pingouin as pg
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'
# 用于解决负号显示问题
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "重复测量方差分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "重复测量方差分析": "用于分析同一组对象在不同时间点或条件下的测量数据，判断不同处理水平之间是否存在显著差异。",
        },
        'interpretation': {
            "F值": "用于检验组间差异的显著性，F值越大，组间差异越显著。",
            "p值": "若 p 值小于显著性水平（通常为 0.05），则拒绝原假设，认为组间存在显著差异。",
            "偏 eta 平方": "反映了组间差异在总变异中所占的比例。"
        }
    },
    'en': {
        'title': "Repeated Measures ANOVA",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "重复测量方差分析": "Used to analyze the measurement data of the same group of objects at different time points or conditions to determine whether there are significant differences between different treatment levels.",
        },
        'interpretation': {
            "F值": "Used to test the significance of the differences between groups. A larger F value indicates more significant differences between groups.",
            "p值": "If the p value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating that there are significant differences between groups.",
            "偏 eta 平方": "Reflects the proportion of the variance explained by the group differences in the total variance."
        }
    }
}


class RepeatedMeasuresANOVAApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设第一列是被试编号，其余列是不同处理水平下的测量值
            subject = df.iloc[:, 0]
            measures = df.iloc[:, 1:]

            # 将数据转换为长格式
            long_df = pd.melt(df, id_vars=df.columns[0], var_name='Treatment', value_name='Value')

            # 进行重复测量方差分析
            anova = pg.rm_anova(data=long_df, dv='Value', within='Treatment', subject=df.columns[0])

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["重复测量方差分析" if self.current_language == 'zh' else "Repeated Measures ANOVA"])
            explanation_df.insert(0, "统计量", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=[
                "F值" if self.current_language == 'zh' else "F value",
                "p值" if self.current_language == 'zh' else "p value",
                "偏 eta 平方" if self.current_language == 'zh' else "Partial eta squared"
            ])
            interpretation_df.insert(0, "统计量", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([anova, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的Word文档
                doc = Document()

                # 添加表格到文档
                table = doc.add_table(rows=combined_df.shape[0] + 1, cols=combined_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(combined_df.columns):
                    hdr_cells[col_idx].text = col_name
                for row_idx in range(combined_df.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(combined_df.iloc[row_idx]):
                        row_cells[col_idx].text = str(value)

                # 生成结果图片
                desktop_path = pathlib.Path.home() / 'Desktop'
                plot_path = desktop_path / 'repeated_measures_anova_plot.png'
                plt.figure()
                for subj in subject.unique():
                    subj_data = long_df[long_df[df.columns[0]] == subj]
                    plt.plot(subj_data['Treatment'], subj_data['Value'], marker='o', label=f'Subject {subj}')
                plt.xlabel('Treatment')
                plt.ylabel('Value')
                plt.title(
                    '重复测量方差分析结果' if self.current_language == 'zh' else 'Repeated Measures ANOVA Results')
                plt.legend()
                plt.savefig(plot_path)
                plt.close()

                # 将图片插入到文档中
                doc.add_picture(str(plot_path), width=Inches(6))

                # 保存Word文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path) + f"结果图片已保存到 {plot_path}" if self.current_language == 'zh' else f"The result image has been saved to {plot_path}"
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        try:
            self.root.mainloop()
        except KeyboardInterrupt:
            print("程序被手动中断。")
            self.root.destroy()  # 销毁主窗口


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = RepeatedMeasuresANOVAApp()
    app.run()


if __name__ == "__main__":
    run_app()