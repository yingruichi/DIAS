import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "模糊综合评价分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "因素权重向量": "各评价因素的权重分配",
            "模糊评价矩阵": "对各因素在不同评价等级下的隶属度矩阵",
            "综合评价结果向量": "综合考虑各因素权重和模糊评价矩阵得到的最终评价结果向量"
        },
        'interpretation': {
            "因素权重向量": "权重越大，该因素在综合评价中越重要",
            "模糊评价矩阵": "反映各因素在不同评价等级下的隶属程度",
            "综合评价结果向量": "向量中值最大的对应评价等级为最终评价结果"
        }
    },
    'en': {
        'title': "Fuzzy Comprehensive Evaluation Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "因素权重向量": "The weight distribution of each evaluation factor",
            "模糊评价矩阵": "The membership matrix of each factor under different evaluation levels",
            "综合评价结果向量": "The final evaluation result vector obtained by comprehensively considering the weights of each factor and the fuzzy evaluation matrix"
        },
        'interpretation': {
            "因素权重向量": "The larger the weight, the more important the factor is in the comprehensive evaluation",
            "模糊评价矩阵": "Reflects the membership degree of each factor under different evaluation levels",
            "综合评价结果向量": "The evaluation level corresponding to the largest value in the vector is the final evaluation result"
        }
    }
}

class FuzzyComprehensiveEvaluationAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.configure(style="TEntry")

    def on_entry_click(self, event):
        """当用户点击输入框时，清除提示信息"""
        if self.file_entry.get() == LANGUAGES[self.current_language]['file_entry_placeholder']:
            self.file_entry.delete(0, tk.END)
            self.file_entry.configure(style="TEntry")

    def on_focusout(self, event):
        """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
            self.file_entry.configure(style="Gray.TEntry")

    def fuzzy_comprehensive_evaluation(self, weight_vector, evaluation_matrix):
        """
        进行模糊综合评价
        :param weight_vector: 因素权重向量
        :param evaluation_matrix: 模糊评价矩阵
        :return: 综合评价结果向量
        """
        result_vector = np.dot(weight_vector, evaluation_matrix)
        result_vector = result_vector / np.sum(result_vector)
        return result_vector

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]['file_entry_placeholder']:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 假设第一行为因素权重向量，其余行为模糊评价矩阵
            weight_vector = data[0]
            evaluation_matrix = data[1:]

            # 进行模糊综合评价
            result_vector = self.fuzzy_comprehensive_evaluation(weight_vector, evaluation_matrix)

            # 整理数据
            data = [
                ["因素权重向量", weight_vector.tolist(), ""],
                ["模糊评价矩阵", evaluation_matrix.tolist(), ""],
                ["综合评价结果向量", result_vector.tolist(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["因素权重向量", "模糊评价矩阵", "综合评价结果向量"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["因素权重向量", "模糊评价矩阵", "综合评价结果向量"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('模糊综合评价分析结果', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for col_index, header in enumerate(headers):
                    hdr_cells[col_index].text = header

                # 添加数据行
                for row_data in combined_df.values.tolist():
                    row_cells = table.add_row().cells
                    for col_index, value in enumerate(row_data):
                        row_cells[col_index].text = str(value)

                # 生成综合评价结果向量柱状图
                fig, ax = plt.subplots()
                ax.bar(range(len(result_vector)), result_vector)
                ax.set_title(
                    '综合评价结果向量柱状图' if self.current_language == 'zh' else 'Bar Chart of Comprehensive Evaluation Result Vector')
                ax.set_xlabel('评价等级' if self.current_language == 'zh' else 'Evaluation Levels')
                ax.set_ylabel('隶属度' if self.current_language == 'zh' else 'Membership Degree')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_result_vector.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_picture(img_path, width=Inches(6))

                # 保存文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.configure(style="Gray.TEntry")

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建文件选择按钮
        self.select_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file,
                                        bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(self.root, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.pack(pady=5)
        self.file_entry.bind("<FocusIn>", self.on_entry_click)
        self.file_entry.bind("<FocusOut>", self.on_focusout)

        # 创建分析按钮
        self.analyze_button = ttk.Button(self.root, text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file,
                                         bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建语言切换标签
        self.switch_language_label = ttk.Label(self.root, text=LANGUAGES[self.current_language]['switch_language'],
                                               cursor="hand2")
        self.switch_language_label.pack(pady=10)
        self.switch_language_label.bind("<Button-1>", self.switch_language)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = FuzzyComprehensiveEvaluationAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()