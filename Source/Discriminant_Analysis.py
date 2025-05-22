import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import pathlib
from sklearn.discriminant_analysis import LinearDiscriminantAnalysis
from sklearn.metrics import classification_report, confusion_matrix
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "判别分析",
        "select_button": "选择文件",
        "analyze_button": "分析文件",
        "file_not_found": "文件不存在，请重新选择。",
        "analysis_success": "分析完成，结果已保存到 {}\n",
        "no_save_path": "未选择保存路径，结果未保存。",
        "analysis_error": "分析文件时出错: {}",
        "switch_language": "切换语言",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "explanation": {
            "判别分析": "根据已知类别的样本数据建立判别函数，对新样本进行分类。",
        },
        "interpretation": {
            "准确率": "分类正确的样本数占总样本数的比例。",
            "精确率": "预测为某一类别的样本中，实际为该类别的样本比例。",
            "召回率": "实际为某一类别的样本中，被正确预测为该类别的样本比例。",
            "F1值": "精确率和召回率的调和平均数，综合衡量分类性能。"
        }
    },
    "en": {
        "title": "Discriminant Analysis",
        "select_button": "Select File",
        "analyze_button": "Analyze File",
        "file_not_found": "The file does not exist. Please select again.",
        "analysis_success": "Analysis completed. The results have been saved to {}\n",
        "no_save_path": "No save path selected. The results were not saved.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "switch_language": "Switch Language",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "explanation": {
            "Discriminant Analysis": "Establish a discriminant function based on sample data of known categories to classify new samples.",
        },
        "interpretation": {
            "Accuracy": "The proportion of correctly classified samples to the total number of samples.",
            "Precision": "The proportion of samples actually belonging to a certain category among the samples predicted to belong to that category.",
            "Recall": "The proportion of samples correctly predicted to belong to a certain category among the samples actually belonging to that category.",
            "F1-score": "The harmonic mean of precision and recall, comprehensively measuring the classification performance."
        }
    }
}


class DiscriminantAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设最后一列是类别变量，其余列是特征变量
            X = df.iloc[:, :-1]
            y = df.iloc[:, -1]

            # 进行判别分析
            lda = LinearDiscriminantAnalysis()
            lda.fit(X, y)
            y_pred = lda.predict(X)

            # 生成分类报告
            report = classification_report(y, y_pred, output_dict=True)
            report_df = pd.DataFrame(report).transpose()

            # 生成混淆矩阵
            cm = confusion_matrix(y, y_pred)
            cm_df = pd.DataFrame(cm, index=lda.classes_, columns=lda.classes_)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["判别分析" if self.current_language == 'zh' else "Discriminant Analysis"])
            explanation_df.insert(0, "统计量", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=[
                "准确率" if self.current_language == 'zh' else "Accuracy",
                "精确率" if self.current_language == 'zh' else "Precision",
                "召回率" if self.current_language == 'zh' else "Recall",
                "F1值" if self.current_language == 'zh' else "F1-score"
            ])
            interpretation_df.insert(0, "统计量", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                document = Document()

                # 添加标题
                document.add_heading('判别分析结果报告', 0)

                # 添加分类报告表格
                document.add_heading('分类报告', level=1)
                table = document.add_table(rows=1, cols=len(report_df.columns))
                hdr_cells = table.rows[0].cells
                for col, key in enumerate(report_df.columns):
                    hdr_cells[col].text = key
                for index, row in report_df.iterrows():
                    row_cells = table.add_row().cells
                    for col, value in enumerate(row):
                        row_cells[col].text = str(value)

                # 添加混淆矩阵表格
                document.add_heading('混淆矩阵', level=1)
                table = document.add_table(rows=1, cols=len(cm_df.columns))
                hdr_cells = table.rows[0].cells
                for col, key in enumerate(cm_df.columns):
                    hdr_cells[col].text = key
                for index, row in cm_df.iterrows():
                    row_cells = table.add_row().cells
                    for col, value in enumerate(row):
                        row_cells[col].text = str(value)

                # 添加解释说明表格
                document.add_heading('解释说明', level=1)
                table = document.add_table(rows=1, cols=len(explanation_df.columns))
                hdr_cells = table.rows[0].cells
                for col, key in enumerate(explanation_df.columns):
                    hdr_cells[col].text = key
                for index, row in explanation_df.iterrows():
                    row_cells = table.add_row().cells
                    for col, value in enumerate(row):
                        row_cells[col].text = str(value)

                # 添加结果解读表格
                document.add_heading('结果解读', level=1)
                table = document.add_table(rows=1, cols=len(interpretation_df.columns))
                hdr_cells = table.rows[0].cells
                for col, key in enumerate(interpretation_df.columns):
                    hdr_cells[col].text = key
                for index, row in interpretation_df.iterrows():
                    row_cells = table.add_row().cells
                    for col, value in enumerate(row):
                        row_cells[col].text = str(value)

                # 生成结果图片（混淆矩阵可视化）
                desktop_path = pathlib.Path.home() / 'Desktop'
                plot_path = os.path.splitext(save_path)[0] + '_confusion_matrix.png'
                plt.figure()
                plt.imshow(cm, interpolation='nearest', cmap=plt.cm.Blues)
                plt.title('判别分析混淆矩阵' if self.current_language == 'zh' else 'Confusion Matrix of Discriminant Analysis')
                plt.colorbar()
                tick_marks = np.arange(len(lda.classes_))
                plt.xticks(tick_marks, lda.classes_, rotation=45)
                plt.yticks(tick_marks, lda.classes_)

                thresh = cm.max() / 2.
                for i in range(cm.shape[0]):
                    for j in range(cm.shape[1]):
                        plt.text(j, i, format(cm[i, j], 'd'),
                                 horizontalalignment="center",
                                 color="white" if cm[i, j] > thresh else "black")

                plt.tight_layout()
                plt.ylabel('真实类别' if self.current_language == 'zh' else 'True label')
                plt.xlabel('预测类别' if self.current_language == 'zh' else 'Predicted label')
                plt.savefig(plot_path)
                plt.close()

                # 将图片插入到 Word 文档中
                document.add_heading('混淆矩阵可视化', level=1)
                document.add_picture(plot_path, width=Inches(6))

                # 保存 Word 文档
                document.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                result_msg += f"\n结果图片已保存到 {plot_path}" if self.current_language == 'zh' else f"\nThe result image has been saved to {plot_path}"
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(languages[self.current_language]['title'])
        self.select_button.config(text=languages[self.current_language]['select_button'])
        self.analyze_button.config(text=languages[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=languages[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = DiscriminantAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()