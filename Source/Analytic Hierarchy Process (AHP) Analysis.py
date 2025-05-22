import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "层次分析法 AHP 分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'explanation': {
            "特征向量": "反映各因素相对重要性的向量",
            "一致性指标 CI": "衡量判断矩阵一致性的指标",
            "随机一致性指标 RI": "根据矩阵阶数确定的随机一致性指标",
            "一致性比率 CR": "CI 与 RI 的比值，判断矩阵是否具有满意一致性"
        },
        'interpretation': {
            "特征向量": "特征向量值越大，对应因素越重要",
            "一致性指标 CI": "CI 值越小，矩阵一致性越好",
            "随机一致性指标 RI": "不同阶数矩阵有对应标准值",
            "一致性比率 CR": "CR < 0.1 时，矩阵具有满意一致性，结果可信"
        }
    },
    'en': {
        'title': "Analytic Hierarchy Process (AHP) Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'explanation': {
            "特征向量": "A vector reflecting the relative importance of each factor",
            "一致性指标 CI": "An indicator to measure the consistency of the judgment matrix",
            "随机一致性指标 RI": "A random consistency indicator determined by the order of the matrix",
            "一致性比率 CR": "The ratio of CI to RI to determine if the matrix has satisfactory consistency"
        },
        'interpretation': {
            "特征向量": "The larger the value in the eigenvector, the more important the corresponding factor",
            "一致性指标 CI": "The smaller the CI value, the better the consistency of the matrix",
            "随机一致性指标 RI": "There are corresponding standard values for matrices of different orders",
            "一致性比率 CR": "When CR < 0.1, the matrix has satisfactory consistency and the results are reliable"
        }
    }
}

# 当前语言
current_language = 'en'

# 随机一致性指标 RI 表
RI_TABLE = {
    1: 0, 2: 0, 3: 0.58, 4: 0.90, 5: 1.12, 6: 1.24, 7: 1.32, 8: 1.41, 9: 1.45
}


def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        file_entry.configure(style="TEntry")


def ahp_analysis(data):
    """
    进行层次分析法 AHP 分析
    :param data: 判断矩阵数据
    :return: 特征向量、一致性指标 CI、一致性比率 CR
    """
    # 计算特征值和特征向量
    eigenvalues, eigenvectors = np.linalg.eig(data)
    max_eigenvalue = np.max(eigenvalues).real
    index = np.argmax(eigenvalues)
    eigenvector = eigenvectors[:, index].real
    eigenvector = eigenvector / np.sum(eigenvector)

    # 计算一致性指标 CI
    n = data.shape[0]
    CI = (max_eigenvalue - n) / (n - 1)

    # 计算随机一致性指标 RI
    RI = RI_TABLE.get(n, None)
    if RI is None:
        raise ValueError("判断矩阵阶数超出支持范围")

    # 计算一致性比率 CR
    CR = CI / RI

    return eigenvector, CI, RI, CR


def analyze_file():
    global current_language
    file_path = file_entry.get()
    if file_path == "请输入待分析 Excel 文件的完整路径" or file_path == "Please enter the full path of the Excel file to be analyzed":
        file_path = ""
    if not os.path.exists(file_path):
        result_label.config(text=LANGUAGES[current_language]['file_not_found'])
        return
    try:
        # 打开 Excel 文件
        df = pd.read_excel(file_path, header=None)
        data = df.values

        # 进行 AHP 分析
        eigenvector, CI, RI, CR = ahp_analysis(data)

        # 整理数据
        data = [
            ["特征向量", eigenvector.tolist(), ""],
            ["一致性指标 CI", CI, ""],
            ["随机一致性指标 RI", RI, ""],
            ["一致性比率 CR", CR, ""]
        ]
        headers = ["统计量", "统计量值", "p值"]
        df = pd.DataFrame(data, columns=headers)

        # 添加解释说明
        explanations = LANGUAGES[current_language]['explanation']
        interpretations = LANGUAGES[current_language]['interpretation']
        explanation_df = pd.DataFrame([explanations])
        explanation_df = explanation_df.reindex(
            columns=["特征向量", "一致性指标 CI", "随机一致性指标 RI", "一致性比率 CR"])
        explanation_df.insert(0, "统计量_解释说明", "解释说明" if current_language == 'zh' else "Explanation")

        # 添加分析结果解读
        interpretation_df = pd.DataFrame([interpretations])
        interpretation_df = interpretation_df.reindex(
            columns=["特征向量", "一致性指标 CI", "随机一致性指标 RI", "一致性比率 CR"])
        interpretation_df.insert(0, "统计量_结果解读", "结果解读" if current_language == 'zh' else "Interpretation")

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建Word文档
            doc = Document()

            # 添加标题
            doc.add_heading(
                '层次分析法 AHP 分析结果' if current_language == 'zh' else 'Analytic Hierarchy Process (AHP) Analysis Results',
                0)

            # 添加分析说明
            doc.add_paragraph('本报告展示了层次分析法(AHP)的分析结果，包括特征向量、一致性指标和一致性比率。')

            # 添加统计量表格
            doc.add_heading('统计量结果', 1)
            table = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for col_index, col_name in enumerate(df.columns):
                hdr_cells[col_index].text = col_name

            for index, row in df.iterrows():
                row_cells = table.add_row().cells
                for col_index, value in enumerate(row):
                    row_cells[col_index].text = str(value)

            # 添加解释说明表格
            doc.add_heading('统计量解释说明', 1)
            table = doc.add_table(rows=1, cols=len(explanation_df.columns))
            hdr_cells = table.rows[0].cells
            for col_index, col_name in enumerate(explanation_df.columns):
                hdr_cells[col_index].text = col_name

            for index, row in explanation_df.iterrows():
                row_cells = table.add_row().cells
                for col_index, value in enumerate(row):
                    row_cells[col_index].text = str(value)

            # 添加结果解读表格
            doc.add_heading('结果解读', 1)
            table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
            hdr_cells = table.rows[0].cells
            for col_index, col_name in enumerate(interpretation_df.columns):
                hdr_cells[col_index].text = col_name

            for index, row in interpretation_df.iterrows():
                row_cells = table.add_row().cells
                for col_index, value in enumerate(row):
                    row_cells[col_index].text = str(value)

            # 生成特征向量柱状图
            fig, ax = plt.subplots()
            ax.bar(range(len(eigenvector)), eigenvector)
            ax.set_title('特征向量柱状图' if current_language == 'zh' else 'Bar Chart of Eigenvector')
            ax.set_xlabel('因素' if current_language == 'zh' else 'Factors')
            ax.set_ylabel('权重' if current_language == 'zh' else 'Weights')

            # 保存图片
            img_path = os.path.splitext(save_path)[0] + '_eigenvector.png'
            plt.savefig(img_path)
            plt.close()

            # 在Word文档中添加图片
            doc.add_heading('特征向量可视化', 1)
            doc.add_picture(img_path, width=Inches(6))

            # 保存Word文档
            doc.save(save_path)

            result_msg = LANGUAGES[current_language]['analysis_success'].format(save_path)
            result_label.config(text=result_msg, wraplength=400)

        else:
            result_label.config(text=LANGUAGES[current_language]['no_save_path'])

    except Exception as e:
        result_label.config(text=LANGUAGES[current_language]['analysis_error'].format(str(e)))


def switch_language():
    global current_language
    current_language = 'en' if current_language == 'zh' else 'zh'
    root.title(LANGUAGES[current_language]['title'])
    select_button.config(text=LANGUAGES[current_language]['select_button'])
    analyze_button.config(text=LANGUAGES[current_language]['analyze_button'])
    language_label.config(text=LANGUAGES[current_language]['switch_language'])
    # 切换语言时更新提示信息
    file_entry.delete(0, tk.END)
    if current_language == 'zh':
        file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
        file_entry.configure(style="Gray.TEntry")
    else:
        file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
        file_entry.configure(style="Gray.TEntry")


def on_entry_click(event):
    """当用户点击输入框时，清除提示信息"""
    if file_entry.get() == "请输入待分析 Excel 文件的完整路径" or file_entry.get() == "Please enter the full path of the Excel file to be analyzed":
        file_entry.delete(0, tk.END)
        file_entry.configure(style="TEntry")


def on_focusout(event):
    """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
    if file_entry.get() == "":
        if current_language == 'zh':
            file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
            file_entry.configure(style="Gray.TEntry")
        else:
            file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
            file_entry.configure(style="Gray.TEntry")


# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(LANGUAGES[current_language]['title'])

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 设置窗口的宽度和高度
window_width = 500
window_height = 300

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建自定义样式
style = ttk.Style()
style.configure("Gray.TEntry", foreground="gray")

# 创建文件选择按钮
select_button = ttk.Button(root, text=LANGUAGES[current_language]['select_button'], command=select_file,
                           bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建文件路径输入框
file_entry = ttk.Entry(root, width=50, style="Gray.TEntry")
if current_language == 'zh':
    file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
else:
    file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
file_entry.pack(pady=5)
file_entry.bind("<FocusIn>", on_entry_click)
file_entry.bind("<FocusOut>", on_focusout)

# 创建分析按钮
analyze_button = ttk.Button(root, text=LANGUAGES[current_language]['analyze_button'], command=analyze_file,
                            bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建语言切换标签
language_label = ttk.Label(root, text=LANGUAGES[current_language]['switch_language'], cursor="hand2")
language_label.pack(pady=10)
language_label.bind("<Button-1>", lambda event: switch_language())

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
root.mainloop()