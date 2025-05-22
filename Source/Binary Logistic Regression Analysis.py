import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from sklearn.metrics import accuracy_score, roc_auc_score
import matplotlib.pyplot as plt
import statsmodels.api as sm
from docx import Document
from docx.shared import Pt

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "二元逻辑回归分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'images_saved': "图片已保存到 {}",
        'switch_language': "切换语言",
        'explanation': {
            "Coefficients": "回归系数，表示每个自变量对因变量的影响程度。",
            "Intercept": "截距，是当所有自变量为 0 时因变量的预测概率。",
            "Accuracy": "准确率，衡量模型预测正确的比例。",
            "AUC": "ROC 曲线下面积，衡量模型区分正例和反例的能力。",
            "z-value": "z 统计量，用于检验每个自变量的显著性。",
            "p-value": "p 值，用于判断自变量的显著性，p 值越小，自变量越显著。"
        }
    },
    'en': {
        'title': "Binary Logistic Regression Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'images_saved': "Images have been saved to {}",
        'switch_language': "Switch Language",
        'explanation': {
            "Coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "Intercept": "Intercept, which is the predicted probability of the dependent variable when all independent variables are 0.",
            "Accuracy": "Accuracy, measuring the proportion of correct predictions of the model.",
            "AUC": "Area Under the ROC Curve, measuring the ability of the model to distinguish between positive and negative examples.",
            "z-value": "z statistic, used to test the significance of each independent variable.",
            "p-value": "p value, used to determine the significance of the independent variable. The smaller the p value, the more significant the independent variable."
        }
    }
}

# 当前语言
current_language = 'en'

def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        file_entry.configure(style="TEntry")  # 恢复默认样式

def analyze_file():
    global current_language
    file_path = file_entry.get()
    if file_path == "请输入待分析 Excel 文件的完整路径" or file_path == "Please enter the full path of the Excel file to be analyzed":
        file_path = ""
    if not os.path.exists(file_path):
        result_label.config(text=LANGUAGES[current_language]['file_not_found'])
        return
    try:
        # 打开 Excel 文件
        df = pd.read_excel(file_path)

        # 假设最后一列是因变量，其余列是自变量
        X = df.iloc[:, :-1].values
        y = df.iloc[:, -1].values

        # 检查因变量是否在 [0, 1] 区间
        if not all(0 <= val <= 1 for val in y):
            raise ValueError("因变量的值必须在 [0, 1] 区间内。")

        # 添加常数项
        X_with_const = sm.add_constant(X)

        # 进行二元逻辑回归分析
        logit_model = sm.Logit(y, X_with_const).fit()
        y_pred_proba = logit_model.predict(X_with_const)
        y_pred = (y_pred_proba > 0.5).astype(int)

        # 计算指标
        coefficients = logit_model.params[1:]
        intercept = logit_model.params[0]
        accuracy = accuracy_score(y, y_pred)
        auc = roc_auc_score(y, y_pred_proba)

        # 计算 z 值和 p 值
        z_values = logit_model.tvalues[1:]
        p_values = logit_model.pvalues[1:]

        # 准备数据
        columns_stats = ["Coefficients", "z-value", "p-value", "Accuracy", "AUC"]
        explanations = LANGUAGES[current_language]['explanation']
        data = [["Binary Logistic Regression"] + list(coefficients) + [intercept, accuracy, auc] + list(z_values) + list(p_values)]
        headers = ["Model"] + [f"Coefficient_{i+1}" for i in range(len(coefficients))] + ["Intercept"] + columns_stats[3:] + [f"z-value_{i+1}" for i in range(len(z_values))] + [f"p-value_{i+1}" for i in range(len(p_values))]

        df = pd.DataFrame(data, columns=headers)

        # 添加解释说明
        explanation_df = pd.DataFrame([explanations])
        explanation_df = explanation_df.reindex(columns=columns_stats)
        explanation_df.insert(0, "Model", "解释说明" if current_language == 'zh' else "Explanation")

        # 合并数据和解释说明
        combined_df = pd.concat([df, explanation_df], ignore_index=True)

        # 转置数据框
        transposed_df = combined_df.set_index('Model').T.reset_index().rename(columns={'index': 'Model'})

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建一个新的 Word 文档
            doc = Document()

            # 添加标题
            doc.add_heading('二元逻辑回归分析结果' if current_language == 'zh' else 'Binary Logistic Regression Analysis Results', 0)

            # 创建一个表格，行数为 DataFrame 的行数，列数为 DataFrame 的列数
            table = doc.add_table(rows=transposed_df.shape[0], cols=transposed_df.shape[1])

            # 填充表格数据
            for row_idx in range(transposed_df.shape[0]):
                row_cells = table.rows[row_idx].cells
                for col_idx in range(transposed_df.shape[1]):
                    cell = row_cells[col_idx]
                    p = cell.paragraphs[0]
                    run = p.add_run(str(transposed_df.iloc[row_idx, col_idx]))
                    run.font.size = Pt(10)

            # 调整表格列宽
            for column in table.columns:
                for cell in column.cells:
                    cell.width = Pt(80)

            # 获取保存路径的目录
            save_dir = os.path.dirname(save_path)

            # 生成 ROC 曲线
            from sklearn.metrics import roc_curve
            fpr, tpr, thresholds = roc_curve(y, y_pred_proba)
            plt.figure(figsize=(10, 6))
            plt.plot(fpr, tpr, label='ROC curve (area = %0.2f)' % auc)
            plt.plot([0, 1], [0, 1], 'k--')
            plt.xlim([0.0, 1.0])
            plt.ylim([0.0, 1.05])
            plt.xlabel('False Positive Rate')
            plt.ylabel('True Positive Rate')
            plt.title('Receiver Operating Characteristic')
            plt.legend(loc="lower right")
            img_name = "logistic_regression_roc.png"
            img_path = os.path.join(save_dir, img_name)
            plt.savefig(img_path)
            plt.close()

            # 在 Word 文档中插入图片
            doc.add_picture(img_path, width=Pt(400))

            # 保存 Word 文档
            doc.save(save_path)

            result_msg = LANGUAGES[current_language]['analysis_success'].format(save_path)
            result_msg += LANGUAGES[current_language]['images_saved'].format(save_dir)
            result_label.config(text=result_msg, wraplength=400)
        else:
            result_label.config(text=LANGUAGES[current_language]['no_save_path'])

    except Exception as e:
        result_label.config(text=LANGUAGES[current_language]['analysis_error'].format(str(e)))

def switch_language():
    global current_language
    current_language = 'en' if current_language == 'zh' else 'zh'
    root.title(LANGUAGES[current_language]['title'])
    select_button.config(text=LANGUAGES[current_language]['select_button'])
    analyze_button.config(text=LANGUAGES[current_language]['analyze_button'])
    language_label.config(text=LANGUAGES[current_language]['switch_language'])
    # 切换语言时更新提示信息
    file_entry.delete(0, tk.END)
    if current_language == 'zh':
        file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
        file_entry.configure(style="Gray.TEntry")
    else:
        file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
        file_entry.configure(style="Gray.TEntry")

def on_entry_click(event):
    """当用户点击输入框时，清除提示信息"""
    if file_entry.get() == "请输入待分析 Excel 文件的完整路径" or file_entry.get() == "Please enter the full path of the Excel file to be analyzed":
        file_entry.delete(0, tk.END)
        file_entry.configure(style="TEntry")  # 恢复默认样式

def on_focusout(event):
    """当用户离开输入框时，如果没有输入内容，显示提示信息"""
    if file_entry.get() == "":
        if current_language == 'zh':
            file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
        else:
            file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
        file_entry.configure(style="Gray.TEntry")

# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(LANGUAGES[current_language]['title'])

# 获取屏幕宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 计算窗口的宽度和高度
window_width = 500
window_height = 250

# 计算窗口的 x 和 y 坐标，使其居中
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建一个框架来包含按钮和输入框
frame = ttk.Frame(root)
frame.pack(expand=True)  # 使用 expand 选项使框架在上下方向上居中

# 创建文件选择按钮
select_button = ttk.Button(frame, text=LANGUAGES[current_language]['select_button'], command=select_file, bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建自定义样式
style = ttk.Style()
style.configure("Gray.TEntry", foreground="gray")
style.configure("Gray.TLabel", foreground="gray")

# 创建文件路径输入框
file_entry = ttk.Entry(frame, width=50, style="Gray.TEntry")
if current_language == 'zh':
    file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
else:
    file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
file_entry.pack(pady=5)
file_entry.bind("<FocusIn>", on_entry_click)
file_entry.bind("<FocusOut>", on_focusout)

# 创建分析按钮
analyze_button = ttk.Button(frame, text=LANGUAGES[current_language]['analyze_button'], command=analyze_file, bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建语言切换标签
language_label = ttk.Label(frame, text=LANGUAGES[current_language]['switch_language'], style="Gray.TLabel", cursor="hand2")
language_label.pack(pady=10)
language_label.bind("<Button-1>", lambda event: switch_language())

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
root.mainloop()