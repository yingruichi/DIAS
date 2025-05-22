import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "层次分析法 AHP 分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}\n",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'explanation': {
            "特征向量": "反映各因素相对重要性的向量",
            "一致性指标 CI": "衡量判断矩阵一致性的指标",
            "随机一致性指标 RI": "根据矩阵阶数确定的随机一致性指标",
            "一致性比率 CR": "CI 与 RI 的比值，判断矩阵是否具有满意一致性"
        },
        'interpretation': {
            "特征向量": "特征向量值越大，对应因素越重要",
            "一致性指标 CI": "CI 值越小，矩阵一致性越好",
            "随机一致性指标 RI": "不同阶数矩阵有对应标准值",
            "一致性比率 CR": "CR < 0.1 时，矩阵具有满意一致性，结果可信"
        }
    },
    'en': {
        'title': "Analytic Hierarchy Process (AHP) Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}\n",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'explanation': {
            "特征向量": "A vector reflecting the relative importance of each factor",
            "一致性指标 CI": "An indicator to measure the consistency of the judgment matrix",
            "随机一致性指标 RI": "A random consistency indicator determined by the order of the matrix",
            "一致性比率 CR": "The ratio of CI to RI to determine if the matrix has satisfactory consistency"
        },
        'interpretation': {
            "特征向量": "The larger the value in the eigenvector, the more important the corresponding factor",
            "一致性指标 CI": "The smaller the CI value, the better the consistency of the matrix",
            "随机一致性指标 RI": "There are corresponding standard values for matrices of different orders",
            "一致性比率 CR": "When CR < 0.1, the matrix has satisfactory consistency and the results are reliable"
        }
    }
}

# 随机一致性指标 RI 表
RI_TABLE = {
    1: 0, 2: 0, 3: 0.58, 4: 0.90, 5: 1.12, 6: 1.24, 7: 1.32, 8: 1.41, 9: 1.45
}

class AnalyticHierarchyProcessAHPApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"
        
        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])
            
        self.create_ui()
    
    def ahp_analysis(self, data):
        """
        进行层次分析法 AHP 分析
        :param data: 判断矩阵数据
        :return: 特征向量、一致性指标 CI、一致性比率 CR
        """
        # 计算特征值和特征向量
        eigenvalues, eigenvectors = np.linalg.eig(data)
        max_eigenvalue = np.max(eigenvalues).real
        index = np.argmax(eigenvalues)
        eigenvector = eigenvectors[:, index].real
        eigenvector = eigenvector / np.sum(eigenvector)

        # 计算一致性指标 CI
        n = data.shape[0]
        CI = (max_eigenvalue - n) / (n - 1)

        # 计算随机一致性指标 RI
        RI = RI_TABLE.get(n, None)
        if RI is None:
            raise ValueError("判断矩阵阶数超出支持范围")

        # 计算一致性比率 CR
        CR = CI / RI

        return eigenvector, CI, RI, CR
        
    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 进行 AHP 分析
            eigenvector, CI, RI, CR = self.ahp_analysis(data)

            # 整理数据
            data = [
                ["特征向量", eigenvector.tolist(), ""],
                ["一致性指标 CI", CI, ""],
                ["随机一致性指标 RI", RI, ""],
                ["一致性比率 CR", CR, ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["特征向量", "一致性指标 CI", "随机一致性指标 RI", "一致性比率 CR"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["特征向量", "一致性指标 CI", "随机一致性指标 RI", "一致性比率 CR"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建Word文档
                doc = Document()

                # 添加标题
                doc.add_heading(
                    '层次分析法 AHP 分析结果' if self.current_language == 'zh' else 'Analytic Hierarchy Process (AHP) Analysis Results',
                    0)

                # 添加分析说明
                doc.add_paragraph('本报告展示了层次分析法(AHP)的分析结果，包括特征向量、一致性指标和一致性比率。')

                # 添加统计量表格
                doc.add_heading('统计量结果', 1)
                table = doc.add_table(rows=1, cols=len(df.columns))
                hdr_cells = table.rows[0].cells
                for col_index, col_name in enumerate(df.columns):
                    hdr_cells[col_index].text = col_name

                for index, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for col_index, value in enumerate(row):
                        row_cells[col_index].text = str(value)

                # 添加解释说明表格
                doc.add_heading('统计量解释说明', 1)
                table = doc.add_table(rows=1, cols=len(explanation_df.columns))
                hdr_cells = table.rows[0].cells
                for col_index, col_name in enumerate(explanation_df.columns):
                    hdr_cells[col_index].text = col_name

                for index, row in explanation_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_index, value in enumerate(row):
                        row_cells[col_index].text = str(value)

                # 添加结果解读表格
                doc.add_heading('结果解读', 1)
                table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
                hdr_cells = table.rows[0].cells
                for col_index, col_name in enumerate(interpretation_df.columns):
                    hdr_cells[col_index].text = col_name

                for index, row in interpretation_df.iterrows():
                    row_cells = table.add_row().cells
                    for col_index, value in enumerate(row):
                        row_cells[col_index].text = str(value)

                # 生成特征向量柱状图
                fig, ax = plt.subplots()
                ax.bar(range(len(eigenvector)), eigenvector)
                ax.set_title('特征向量柱状图' if self.current_language == 'zh' else 'Bar Chart of Eigenvector')
                ax.set_xlabel('因素' if self.current_language == 'zh' else 'Factors')
                ax.set_ylabel('权重' if self.current_language == 'zh' else 'Weights')

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_eigenvector.png'
                plt.savefig(img_path)
                plt.close()

                # 在Word文档中添加图片
                doc.add_heading('特征向量可视化', 1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存Word文档
                doc.save(save_path)

                result_msg = languages[self.current_language]['analysis_complete'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path_selected'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(languages[self.current_language]['title'])
        self.select_button.config(text=languages[self.current_language]['select_button_text'])
        self.analyze_button.config(text=languages[self.current_language]['analyze_button_text'])
        self.switch_language_label.config(text=languages[self.current_language]['switch_language_button_text'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')
        
    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"], 
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50, style="Gray.TEntry")
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"], 
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language_button_text"],
                                              foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)
        
    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = AnalyticHierarchyProcessAHPApp()
    app.run()

if __name__ == "__main__":
    run_app()