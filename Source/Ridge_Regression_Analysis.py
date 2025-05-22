import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from sklearn.linear_model import Ridge
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import StandardScaler
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "岭回归分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'images_saved': "图片已保存到 {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "Coefficients": "回归系数，表示每个自变量对因变量的影响程度。",
            "Intercept": "截距，是当所有自变量为 0 时因变量的预测值。",
            "Mean Squared Error (MSE)": "均方误差，衡量预测值与真实值之间的平均误差。",
            "R-squared (R²)": "决定系数，取值范围在 0 到 1 之间，越接近 1 表示模型拟合效果越好。",
            "Adjusted R-squared": "调整决定系数，考虑了模型中自变量的数量，对模型的拟合优度进行了调整。",
            "F-value": "F 统计量，用于检验整个回归模型的显著性。",
            "t-value": "t 统计量，用于检验每个自变量的显著性。",
            "p-value": "p 值，用于判断自变量的显著性，p 值越小，自变量越显著。"
        }
    },
    'en': {
        'title': "Ridge Regression Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'images_saved': "Images have been saved to {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "Coefficients": "Regression coefficients, indicating the influence of each independent variable on the dependent variable.",
            "Intercept": "Intercept, which is the predicted value of the dependent variable when all independent variables are 0.",
            "Mean Squared Error (MSE)": "Mean squared error, measuring the average error between the predicted and actual values.",
            "R-squared (R²)": "Coefficient of determination, ranging from 0 to 1. A value closer to 1 indicates a better fit of the model.",
            "Adjusted R-squared": "Adjusted coefficient of determination, which takes into account the number of independent variables in the model and adjusts the goodness of fit of the model.",
            "F-value": "F statistic, used to test the significance of the entire regression model.",
            "t-value": "t statistic, used to test the significance of each independent variable.",
            "p-value": "p value, used to determine the significance of the independent variable. The smaller the p value, the more significant the independent variable."
        }
    }
}


class RidgeRegressionAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设最后一列是因变量，其余列是自变量
            X = df.iloc[:, :-1].values
            y = df.iloc[:, -1].values

            # 进行岭回归分析
            ridge = Ridge(alpha=1.0)
            ridge.fit(X, y)
            y_pred = ridge.predict(X)

            # 计算指标
            coefficients = ridge.coef_
            intercept = ridge.intercept_
            mse = mean_squared_error(y, y_pred)
            r2 = r2_score(y, y_pred)
            n = len(y)
            p = X.shape[1]
            adjusted_r2 = 1 - (1 - r2) * (n - 1) / (n - p - 1)

            # 计算 t 值和 p 值
            X_with_const = sm.add_constant(X)
            model = sm.OLS(y, X_with_const).fit()
            t_values = model.tvalues
            p_values = model.pvalues

            # 计算标准化系数
            scaler = StandardScaler()
            X_scaled = scaler.fit_transform(X)
            ridge_scaled = Ridge(alpha=1.0)
            ridge_scaled.fit(X_scaled, y)
            standardized_coefficients = ridge_scaled.coef_

            # 计算 F 值
            f_value = model.fvalue

            # 准备数据
            columns_stats = ["Coefficients", "Standardized Coefficients", "t-value", "p-value", "R-squared (R²)",
                             "Adjusted R-squared", "F-value"]
            explanations = LANGUAGES[self.current_language]['explanation']
            data = [["Ridge Regression"] + list(coefficients) + [intercept, mse, r2, adjusted_r2, f_value] + list(
                standardized_coefficients) + list(t_values) + list(p_values)]
            headers = ["Model"] + [f"Coefficient_{i + 1}" for i in range(len(coefficients))] + ["Intercept",
                                                                                                "MSE"] + columns_stats[
                                                                                                         4:] + [
                          f"Standardized Coefficient_{i + 1}" for i in range(len(standardized_coefficients))] + [
                          f"t-value_{i + 1}" for i in range(len(t_values))] + [f"p-value_{i + 1}" for i in
                                                                               range(len(p_values))]

            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=columns_stats)
            explanation_df.insert(0, "Model", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 合并数据和解释说明
            combined_df = pd.concat([df_result, explanation_df], ignore_index=True)

            # 转置数据框
            transposed_df = combined_df.set_index('Model').T.reset_index().rename(columns={'index': 'Model'})

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加表格
                table = doc.add_table(rows=transposed_df.shape[0] + 1, cols=transposed_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(transposed_df.columns):
                    hdr_cells[col_idx].text = col_name

                for row_idx, row in transposed_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 获取保存路径的目录
                save_dir = os.path.dirname(save_path)

                # 生成散点图
                plt.figure(figsize=(10, 6))
                plt.scatter(y, y_pred)
                plt.plot([y.min(), y.max()], [y.min(), y.max()], 'r--', lw=2)
                plt.xlabel('Actual Values')
                plt.ylabel('Predicted Values')
                plt.title('Actual vs Predicted Values')
                img_name = "ridge_regression_scatter.png"
                img_path = os.path.join(save_dir, img_name)
                plt.savefig(img_path)
                plt.close()

                # 生成岭迹图
                alphas = np.logspace(-2, 2, 50)
                coefs = []
                for a in alphas:
                    ridge = Ridge(alpha=a)
                    ridge.fit(X, y)
                    coefs.append(ridge.coef_)

                plt.figure(figsize=(10, 6))
                ax = plt.gca()
                ax.plot(alphas, coefs)
                ax.set_xscale('log')
                plt.xlabel('Alpha')
                plt.ylabel('Coefficients')
                plt.title('Ridge coefficients as a function of the regularization')
                plt.axis('tight')
                img_name_ridge_trace = "ridge_regression_trace.png"
                img_path_ridge_trace = os.path.join(save_dir, img_name_ridge_trace)
                plt.savefig(img_path_ridge_trace)
                plt.close()

                # 在 Word 文档中插入图片
                doc.add_heading('Actual vs Predicted Values', level=2)
                doc.add_picture(img_path, width=Inches(6))
                doc.add_heading('Ridge coefficients as a function of the regularization', level=2)
                doc.add_picture(img_path_ridge_trace, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                result_msg += LANGUAGES[self.current_language]['images_saved'].format(save_dir)
                self.result_label.config(text=result_msg, wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 250

        # 计算窗口的 x 和 y 坐标，使其居中
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)  # 使用 expand 选项使框架在上下方向上居中

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = RidgeRegressionAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()