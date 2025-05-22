import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
from statsmodels.sandbox.stats.runs import runstest_1samp
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'
# 用于解决负号显示问题
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "游程检验",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "游程检验": "用于检验数据的随机性。",
            "游程数": "数据中连续相同符号的段数。",
            "Z统计量": "用于衡量游程数与期望游程数之间的差异。",
            "p值": "p值小于显著性水平（通常为0.05）时，拒绝原假设，认为数据不具有随机性；否则，接受原假设，认为数据具有随机性。"
        },
        'interpretation': {
            "游程数": "游程数过多或过少可能表示数据不具有随机性。",
            "Z统计量": "Z统计量的绝对值越大，说明游程数与期望游程数之间的差异越显著。",
            "p值": "用于判断数据是否具有随机性。"
        }
    },
    'en': {
        'title': "Runs Test",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "游程检验": "Used to test the randomness of data.",
            "游程数": "The number of consecutive segments of the same symbol in the data.",
            "Z统计量": "Used to measure the difference between the number of runs and the expected number of runs.",
            "p值": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating that the data is not random; otherwise, the null hypothesis is accepted, indicating that the data is random."
        },
        'interpretation': {
            "游程数": "Too many or too few runs may indicate that the data is not random.",
            "Z统计量": "The larger the absolute value of the Z-statistic, the more significant the difference between the number of runs and the expected number of runs.",
            "p值": "Used to determine whether the data is random."
        }
    }
}


class RunsTestApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行游程检验。")

            # 对每列数据进行游程检验
            results = []
            for column in numerical_df.columns:
                data = numerical_df[column].dropna()
                z_stat, p_value = runstest_1samp(data)
                runs = runstest_1samp(data, cutoff='median')[1]
                results.append([column, runs, z_stat, p_value])

            # 整理数据
            headers = ["列名", "游程数", "Z统计量", "p值"]
            result_df = pd.DataFrame(results, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                for _, row in result_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)

                # 添加解释说明
                doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", level=2)
                for key, value in explanations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 添加分析结果解读
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Interpretation", level=2)
                for key, value in interpretations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 绘制折线图并插入到 Word 文档中
                for column in numerical_df.columns:
                    data = numerical_df[column].dropna()
                    plt.figure(figsize=(10, 6))
                    plt.plot(data)
                    plt.title(f'{column} 数据折线图' if self.current_language == 'zh' else f'{column} Data Line Plot')
                    plt.xlabel('序号' if self.current_language == 'zh' else 'Index')
                    plt.ylabel('数值' if self.current_language == 'zh' else 'Value')
                    plot_path = save_path.replace('.docx', f'_{column}_lineplot.png')
                    plt.savefig(plot_path)
                    plt.close()
                    doc.add_heading(
                        f'{column} 数据折线图' if self.current_language == 'zh' else f'{column} Data Line Plot',
                        level=2)
                    doc.add_picture(plot_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = RunsTestApp()
    app.run()


if __name__ == "__main__":
    run_app()