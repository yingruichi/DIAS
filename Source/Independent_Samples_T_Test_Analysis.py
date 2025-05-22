import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "独立样本 t 检验分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "独立样本 t 检验": "用于比较两个独立样本的均值是否有显著差异。",
            "样本量": "每个样本中的观测值数量。",
            "均值": "样本数据的平均值。",
            "标准差": "样本数据的离散程度。",
            "t 统计量": "用于衡量两个样本均值差异的程度。",
            "自由度": "在统计分析中能够自由取值的变量个数。",
            "p 值": "用于判断两个样本均值是否有显著差异的指标。",
            "置信区间": "均值差异的可能范围。"
        },
        'interpretation': {
            "t 统计量": "t 统计量的绝对值越大，说明两个样本均值的差异越显著。",
            "自由度": "自由度越大，t 分布越接近正态分布。",
            "p 值": "p 值小于显著性水平（通常为 0.05）时，拒绝原假设，认为两个样本均值存在显著差异；否则，接受原假设，认为两个样本均值无显著差异。",
            "置信区间": "如果置信区间不包含 0，说明两个样本均值存在显著差异。"
        }
    },
    'en': {
        'title': "Independent Samples T-Test Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "独立样本 t 检验": "Used to compare whether the means of two independent samples are significantly different.",
            "样本量": "The number of observations in each sample.",
            "均值": "The average value of the sample data.",
            "标准差": "The degree of dispersion of the sample data.",
            "t 统计量": "Used to measure the degree of difference between the means of two samples.",
            "自由度": "The number of variables that can take on independent values in a statistical analysis.",
            "p 值": "An indicator used to determine whether the means of two samples are significantly different.",
            "置信区间": "The possible range of the difference in means."
        },
        'interpretation': {
            "t 统计量": "The larger the absolute value of the t statistic, the more significant the difference between the means of the two samples.",
            "自由度": "The larger the degrees of freedom, the closer the t distribution is to the normal distribution.",
            "p 值": "When the p-value is less than the significance level (usually 0.05), the null hypothesis is rejected, indicating a significant difference between the means of the two samples; otherwise, the null hypothesis is accepted, indicating no significant difference.",
            "置信区间": "If the confidence interval does not contain 0, it indicates a significant difference between the means of the two samples."
        }
    }
}


class IndependentSamplesTTestAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def plot_results(self, sample1, sample2):
        # 柱状图
        plt.figure(figsize=(12, 8))
        plt.subplot(2, 2, 1)
        means = [sample1.mean(), sample2.mean()]
        labels = ['Sample 1', 'Sample 2']
        plt.bar(labels, means)
        plt.title('Bar Chart')
        plt.xlabel('Samples')
        plt.ylabel('Mean')

        # 误差线图
        plt.subplot(2, 2, 2)
        stds = [sample1.std(), sample2.std()]
        plt.errorbar(labels, means, yerr=stds, fmt='o')
        plt.title('Error Bar Chart')
        plt.xlabel('Samples')
        plt.ylabel('Mean')

        # 箱线图
        plt.subplot(2, 2, 3)
        plt.boxplot([sample1, sample2])
        plt.title('Box Plot')
        plt.xlabel('Samples')
        plt.ylabel('Value')

        # 折线图
        plt.subplot(2, 2, 4)
        plt.plot(sample1, label='Sample 1')
        plt.plot(sample2, label='Sample 2')
        plt.title('Line Chart')
        plt.xlabel('Index')
        plt.ylabel('Value')
        plt.legend()

        plt.tight_layout()
        plt.show()

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行独立样本 t 检验。")
            if len(numerical_df.columns) != 2:
                raise ValueError("数据必须包含两列数值数据，用于独立样本 t 检验。")

            # 进行独立样本 t 检验
            sample1 = numerical_df.iloc[:, 0]
            sample2 = numerical_df.iloc[:, 1]
            t_stat, p_value = stats.ttest_ind(sample1, sample2)
            df_value = len(sample1) + len(sample2) - 2
            mean_diff = sample1.mean() - sample2.mean()
            std_err = stats.sem(sample1 - sample2)
            conf_int = stats.t.interval(0.95, df_value, loc=mean_diff, scale=std_err)

            # 计算样本量、均值和标准差
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()
            stds = numerical_df.std()

            # 整理数据
            data = [
                ["独立样本 t 检验", t_stat, df_value, p_value, conf_int],
                ["样本量", sample_sizes.to_dict(), "", "", ""],
                ["均值", means.to_dict(), "", "", ""],
                ["标准差", stds.to_dict(), "", "", ""]
            ]
            headers = ["统计量", "t 统计量", "自由度", "p 值", "置信区间"]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["独立样本 t 检验", "样本量", "均值", "标准差", "t 统计量", "自由度", "p 值", "置信区间"])
            explanation_df.insert(0, "统计量_解释说明",
                                  "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["t 统计量", "自由度", "p 值", "置信区间"])
            interpretation_df.insert(0, "统计量_结果解读",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df_result, explanation_df, interpretation_df],
                                    ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                     filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的Word文档
                doc = Document()
                doc.add_heading('独立样本 t 检验分析结果', 0)

                # 添加表格
                table = doc.add_table(rows=combined_df.shape[0] + 1,
                                      cols=combined_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(combined_df.columns):
                    hdr_cells[col_idx].text = header

                for row_idx in range(combined_df.shape[0]):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(combined_df.iloc[row_idx]):
                        row_cells[col_idx].text = str(value)

                # 保存Word文档
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(
                    save_path)
                self.result_label.config(text=result_msg, wraplength=400)

                # 绘制图表
                self.plot_results(sample1, sample2)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建自定义样式
        style = ttk.Style()
        style.configure("Gray.TEntry", foreground="gray")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame,
                                        text=LANGUAGES[self.current_language]['select_button'],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame,
                                         text=LANGUAGES[self.current_language]['analyze_button'],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=LANGUAGES[self.current_language]['switch_language'],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = IndependentSamplesTTestAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()