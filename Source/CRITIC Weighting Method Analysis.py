import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "CRITIC 权重法分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'explanation': {
            "原始数据矩阵": "从 Excel 文件中读取的原始数据矩阵",
            "标准差矩阵": "各指标的标准差矩阵，反映指标的对比强度",
            "相关系数矩阵": "各指标之间的相关系数矩阵，反映指标之间的冲突性",
            "信息量矩阵": "结合标准差和相关系数计算得到的各指标信息量矩阵",
            "指标权重": "根据信息量矩阵计算得到的各指标权重",
        },
        'interpretation': {
            "原始数据矩阵": "用于后续分析的基础数据",
            "标准差矩阵": "标准差越大，该指标的对比强度越大，在综合评价中越重要",
            "相关系数矩阵": "相关系数越小，指标之间的冲突性越大，该指标在综合评价中越重要",
            "信息量矩阵": "反映各指标包含的信息量，信息量越大，该指标越重要",
            "指标权重": "各指标在综合评价中的相对重要程度，权重越大越重要"
        }
    },
    'en': {
        'title': "CRITIC Weighting Method Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'explanation': {
            "原始数据矩阵": "The original data matrix read from the Excel file",
            "标准差矩阵": "The standard deviation matrix of each indicator, reflecting the contrast intensity of the indicators",
            "相关系数矩阵": "The correlation coefficient matrix between each indicator, reflecting the conflict between the indicators",
            "信息量矩阵": "The information matrix of each indicator calculated by combining the standard deviation and correlation coefficient",
            "指标权重": "The weight of each indicator calculated based on the information matrix"
        },
        'interpretation': {
            "原始数据矩阵": "The basic data for subsequent analysis",
            "标准差矩阵": "The larger the standard deviation, the greater the contrast intensity of the indicator, and the more important it is in the comprehensive evaluation",
            "相关系数矩阵": "The smaller the correlation coefficient, the greater the conflict between the indicators, and the more important the indicator is in the comprehensive evaluation",
            "信息量矩阵": "Reflects the information contained in each indicator. The greater the information, the more important the indicator",
            "指标权重": "The relative importance of each indicator in the comprehensive evaluation. The larger the weight, the more important it is"
        }
    }
}

# 当前语言
current_language = 'en'


def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        file_entry.configure(style="TEntry")


def critic_weight_method(data):
    """
    实现 CRITIC 权重法
    :param data: 原始数据矩阵
    :return: 标准差矩阵, 相关系数矩阵, 信息量矩阵, 指标权重
    """
    # 计算标准差矩阵
    std_matrix = np.std(data, axis=0)

    # 计算相关系数矩阵
    corr_matrix = np.corrcoef(data, rowvar=False)

    # 计算冲突性
    conflict = 1 - corr_matrix

    # 计算信息量矩阵
    info_matrix = std_matrix * np.sum(conflict, axis=0)

    # 计算指标权重
    weights = info_matrix / np.sum(info_matrix)

    return std_matrix, corr_matrix, info_matrix, weights


def analyze_file():
    global current_language
    file_path = file_entry.get()
    if file_path == "请输入待分析 Excel 文件的完整路径" or file_path == "Please enter the full path of the Excel file to be analyzed":
        file_path = ""
    if not os.path.exists(file_path):
        result_label.config(text=LANGUAGES[current_language]['file_not_found'])
        return
    try:
        # 打开 Excel 文件
        df = pd.read_excel(file_path, header=None)
        data = df.values

        # 将数据转换为浮点类型
        data = data.astype(float)

        # 进行 CRITIC 权重法分析
        std_matrix, corr_matrix, info_matrix, weights = critic_weight_method(data)

        # 整理数据
        data = [
            ["原始数据矩阵", data.tolist(), ""],
            ["标准差矩阵", std_matrix.tolist(), ""],
            ["相关系数矩阵", corr_matrix.tolist(), ""],
            ["信息量矩阵", info_matrix.tolist(), ""],
            ["指标权重", weights.tolist(), ""]
        ]
        headers = ["统计量", "统计量值", "p值"]
        df = pd.DataFrame(data, columns=headers)

        # 添加解释说明
        explanations = LANGUAGES[current_language]['explanation']
        interpretations = LANGUAGES[current_language]['interpretation']
        explanation_df = pd.DataFrame([explanations])
        explanation_df = explanation_df.reindex(
            columns=["原始数据矩阵", "标准差矩阵", "相关系数矩阵", "信息量矩阵", "指标权重"])
        explanation_df.insert(0, "统计量_解释说明", "解释说明" if current_language == 'zh' else "Explanation")

        # 添加分析结果解读
        interpretation_df = pd.DataFrame([interpretations])
        interpretation_df = interpretation_df.reindex(
            columns=["原始数据矩阵", "标准差矩阵", "相关系数矩阵", "信息量矩阵", "指标权重"])
        interpretation_df.insert(0, "统计量_结果解读", "结果解读" if current_language == 'zh' else "Interpretation")

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建 Word 文档
            doc = Document()

            # 添加分析结果表格
            doc.add_heading('分析结果', level=1)
            table = doc.add_table(rows=1, cols=len(headers))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # 添加解释说明表格
            doc.add_heading('解释说明', level=1)
            table = doc.add_table(rows=1, cols=len(explanation_df.columns))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(explanation_df.columns):
                hdr_cells[i].text = header
            for _, row in explanation_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # 添加结果解读表格
            doc.add_heading('结果解读', level=1)
            table = doc.add_table(rows=1, cols=len(interpretation_df.columns))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(interpretation_df.columns):
                hdr_cells[i].text = header
            for _, row in interpretation_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # 生成指标权重柱状图
            fig, ax = plt.subplots()
            ax.bar(range(len(weights)), weights)
            ax.set_title(
                '指标权重柱状图' if current_language == 'zh' else 'Bar Chart of Indicator Weights')
            ax.set_xlabel('指标编号' if current_language == 'zh' else 'Indicator Number')
            ax.set_ylabel('指标权重' if current_language == 'zh' else 'Indicator Weight')
            # 保存图片
            img_path = os.path.splitext(save_path)[0] + '_indicator_weights.png'
            plt.savefig(img_path)
            plt.close()

            # 在 Word 文档中插入图片
            doc.add_heading('指标权重柱状图', level=1)
            doc.add_picture(img_path, width=Inches(6))

            # 保存 Word 文档
            doc.save(save_path)

            result_msg = LANGUAGES[current_language]['analysis_success'].format(save_path)
            result_label.config(text=result_msg, wraplength=400)

        else:
            result_label.config(text=LANGUAGES[current_language]['no_save_path'])

    except Exception as e:
        result_label.config(text=LANGUAGES[current_language]['analysis_error'].format(str(e)))


def switch_language():
    global current_language
    current_language = 'en' if current_language == 'zh' else 'zh'
    root.title(LANGUAGES[current_language]['title'])
    select_button.config(text=LANGUAGES[current_language]['select_button'])
    analyze_button.config(text=LANGUAGES[current_language]['analyze_button'])
    language_label.config(text=LANGUAGES[current_language]['switch_language'])
    # 切换语言时更新提示信息
    file_entry.delete(0, tk.END)
    if current_language == 'zh':
        file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
        file_entry.configure(style="Gray.TEntry")
    else:
        file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
        file_entry.configure(style="Gray.TEntry")


def on_entry_click(event):
    """当用户点击输入框时，清除提示信息"""
    if file_entry.get() == "请输入待分析 Excel 文件的完整路径" or file_entry.get() == "Please enter the full path of the Excel file to be analyzed":
        file_entry.delete(0, tk.END)
        file_entry.configure(style="TEntry")


def on_focusout(event):
    """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
    if file_entry.get() == "":
        if current_language == 'zh':
            file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
            file_entry.configure(style="Gray.TEntry")
        else:
            file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
            file_entry.configure(style="Gray.TEntry")


# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(LANGUAGES[current_language]['title'])

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 设置窗口的宽度和高度
window_width = 500
window_height = 300

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建自定义样式
style = ttk.Style()
style.configure("Gray.TEntry", foreground="gray")

# 创建文件选择按钮
select_button = ttk.Button(root, text=LANGUAGES[current_language]['select_button'], command=select_file,
                           bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建文件路径输入框
file_entry = ttk.Entry(root, width=50, style="Gray.TEntry")
if current_language == 'zh':
    file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
else:
    file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
file_entry.pack(pady=5)
file_entry.bind("<FocusIn>", on_entry_click)
file_entry.bind("<FocusOut>", on_focusout)

# 创建分析按钮
analyze_button = ttk.Button(root, text=LANGUAGES[current_language]['analyze_button'], command=analyze_file,
                            bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建语言切换标签
language_label = ttk.Label(root, text=LANGUAGES[current_language]['switch_language'], cursor="hand2")
language_label.pack(pady=10)
language_label.bind("<Button-1>", lambda event: switch_language())

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
root.mainloop()