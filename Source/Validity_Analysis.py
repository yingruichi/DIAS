import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
from factor_analyzer import FactorAnalyzer
from factor_analyzer.factor_analyzer import calculate_kmo, calculate_bartlett_sphericity
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'  # 设置字体为黑体，可根据系统情况修改为其他支持中文的字体
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
languages = {
    "zh": {
        "title": "效度分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "switch_language_button_text": "切换语言",
        "explanation": {
            "KMO检验值": "Kaiser-Meyer-Olkin检验用于衡量数据是否适合进行因子分析，取值范围在0 - 1之间，越接近1越适合。",
            "Bartlett球形检验p值": "用于检验变量之间是否存在相关性，p值小于0.05表示变量之间存在相关性，适合进行因子分析。",
            "因子载荷矩阵": "反映了每个变量与每个因子之间的相关性。",
            "样本量": "每个样本中的观测值数量。",
            "均值": "样本数据的平均值。"
        },
        "interpretation": {
            "KMO检验值": "KMO检验值越接近1，说明变量之间的相关性越强，越适合进行因子分析。",
            "Bartlett球形检验p值": "若Bartlett球形检验p值小于0.05，则拒绝原假设，表明变量之间存在相关性，适合进行因子分析。",
            "因子载荷矩阵": "因子载荷的绝对值越大，说明该变量与对应因子的相关性越强。",
            "样本量": "样本量的大小会影响统计检验的稳定性，较大的样本量通常能提供更可靠的结果。",
            "均值": "均值反映了数据的平均水平，可用于比较不同变量的集中趋势。"
        }
    },
    "en": {
        "title": "Validity Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "switch_language_button_text": "Switch Language",
        "explanation": {
            "KMO检验值": "The Kaiser-Meyer-Olkin (KMO) test measures whether the data is suitable for factor analysis. The value ranges from 0 to 1, and the closer it is to 1, the more suitable it is.",
            "Bartlett球形检验p值": "Used to test whether there is a correlation between variables. A p-value less than 0.05 indicates that there is a correlation between variables, which is suitable for factor analysis.",
            "因子载荷矩阵": "Reflects the correlation between each variable and each factor.",
            "样本量": "The number of observations in each sample.",
            "均值": "The average value of the sample data."
        },
        "interpretation": {
            "KMO检验值": "The closer the KMO test value is to 1, the stronger the correlation between variables, and the more suitable it is for factor analysis.",
            "Bartlett球形检验p值": "If the p-value of the Bartlett's test of sphericity is less than 0.05, the null hypothesis is rejected, indicating that there is a correlation between variables, which is suitable for factor analysis.",
            "因子载荷矩阵": "The larger the absolute value of the factor loading, the stronger the correlation between the variable and the corresponding factor.",
            "样本量": "The sample size affects the stability of the statistical test. A larger sample size usually provides more reliable results.",
            "均值": "The mean reflects the average level of the data and can be used to compare the central tendencies of different variables."
        }
    }
}


class ValidityAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def validity_analysis(self, data):
        # 进行KMO检验和Bartlett球形检验
        kmo_all, kmo_model = calculate_kmo(data)
        chi_square_value, p_value = calculate_bartlett_sphericity(data)

        # 进行因子分析
        fa = FactorAnalyzer(n_factors=len(data.columns), rotation=None)
        fa.fit(data)
        loadings = fa.loadings_

        return kmo_model, p_value, loadings

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 检查数据是否为数值类型
            numerical_df = df.select_dtypes(include=[np.number])
            if numerical_df.empty:
                raise ValueError("数据中没有数值列，无法进行效度分析。")

            # 进行效度分析
            kmo, bartlett_p, loadings = self.validity_analysis(numerical_df)

            # 计算样本量和均值
            sample_sizes = numerical_df.count()
            means = numerical_df.mean()

            # 整理数据
            data = [
                ["KMO检验值", kmo, ""],
                ["Bartlett球形检验p值", bartlett_p, ""],
                ["因子载荷矩阵", pd.DataFrame(loadings, index=numerical_df.columns).to_csv(sep='\t'), ""],
                ["样本量", sample_sizes.to_dict(), ""],
                ["均值", means.to_dict(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading(languages[self.current_language]["title"], 0)

                # 添加统计结果表格
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header

                for _, row in df_result.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        if isinstance(value, dict):
                            value_str = '\n'.join([f"{k}: {v}" for k, v in value.items()])
                            row_cells[i].text = value_str
                        else:
                            row_cells[i].text = str(value)

                # 添加解释说明
                doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", 1)
                for key, value in explanations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 添加分析结果解读
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Interpretation", 1)
                for key, value in interpretations.items():
                    doc.add_paragraph(f"{key}: {value}")

                # 生成图片（均值柱状图）
                fig, ax = plt.subplots()
                means.plot(kind='bar', ax=ax)
                ax.set_title('变量均值柱状图' if self.current_language == 'zh' else 'Bar Chart of Variable Means')
                ax.set_xlabel('变量' if self.current_language == 'zh' else 'Variables')
                ax.set_ylabel('均值' if self.current_language == 'zh' else 'Mean')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading("变量均值柱状图" if self.current_language == 'zh' else 'Bar Chart of Variable Means', 1)
                doc.add_picture(img_path)

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = ValidityAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()