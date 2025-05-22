import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from statsmodels.sandbox.regression.gmm import GMM
from docx import Document

# 设置支持中文的字体
plt.rcParams['font.family'] = 'SimHei'  # 使用黑体字体，可根据系统情况修改
plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

# 定义语言字典
LANGUAGES = {
    "zh": {
        "title": "GMM估计分析",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，相关图片已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "columns_stats": ["参数", "估计值", "标准误差"],
        "switch_language_button_text": "切换语言"
    },
    "en": {
        "title": "GMM Estimation Analysis",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "columns_stats": ["Parameter", "Estimated Value", "Standard Error"],
        "switch_language_button_text": "Switch Language"
    }
}


class GMMEstimationAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    # 定义 GMM 模型的矩条件函数
    def moment_condition(self, params, exog, endog):
        beta = params
        error = endog - np.dot(exog, beta)
        moments = exog * error[:, np.newaxis]
        return moments

    # 计算 GMM 估计的函数
    def calculate_gmm(self, X, y):
        nobs = X.shape[0]
        nvar = X.shape[1]

        # 创建 GMM 模型实例
        model = GMM(y, X, None, self.moment_condition)

        # 初始参数猜测
        beta0 = np.zeros(nvar)

        # 进行 GMM 估计
        result = model.fit(beta0, maxiter=100, optim_method='nm', wargs=dict(centered=False))

        # 提取估计结果
        params = result.params
        std_errors = result.bse

        # 创建结果 DataFrame
        gmm_data = pd.DataFrame({
            LANGUAGES[self.current_language]["columns_stats"][0]: [f"beta_{i}" for i in range(nvar)],
            LANGUAGES[self.current_language]["columns_stats"][1]: params,
            LANGUAGES[self.current_language]["columns_stats"][2]: std_errors
        })

        # 绘制参数估计值的柱状图
        plt.figure()
        plt.bar(gmm_data[LANGUAGES[self.current_language]["columns_stats"][0]],
                gmm_data[LANGUAGES[self.current_language]["columns_stats"][1]])
        plt.title("GMM 参数估计值" if self.current_language == "zh" else "GMM Parameter Estimates")
        plt.xlabel("参数" if self.current_language == "zh" else "Parameters")
        plt.ylabel("估计值" if self.current_language == "zh" else "Estimated Values")

        # 保存图片
        image_path = os.path.splitext(self.file_entry.get())[0] + '_gmm_plot.png'
        plt.savefig(image_path)
        plt.close()

        return gmm_data

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=LANGUAGES[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)

            # 假设最后一列是因变量，其余列是自变量
            y = df.iloc[:, -1].values
            X = df.iloc[:, :-1].values

            # 计算 GMM 估计
            gmm_data = self.calculate_gmm(X, y)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('GMM 估计分析结果', 0)

                # 添加表格标题
                doc.add_heading('GMM 估计结果', level=1)

                # 创建表格
                table = doc.add_table(rows=1, cols=len(gmm_data.columns))
                hdr_cells = table.rows[0].cells
                for col_idx, col_name in enumerate(gmm_data.columns):
                    hdr_cells[col_idx].text = col_name

                # 填充表格数据
                for _, row in gmm_data.iterrows():
                    row_cells = table.add_row().cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加图片
                image_path = os.path.splitext(self.file_entry.get())[0] + '_gmm_plot.png'
                if os.path.exists(image_path):
                    doc.add_heading('参数估计值柱状图', level=1)
                    doc.add_picture(image_path)

                # 保存 Word 文档
                doc.save(save_path)

                # 设置 wraplength 属性让文本自动换行
                self.result_label.config(text=LANGUAGES[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=LANGUAGES[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = "en" if self.current_language == "zh" else "zh"

        # 更新界面文字
        self.root.title(LANGUAGES[self.current_language]["title"])
        self.select_button.config(text=LANGUAGES[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=LANGUAGES[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=LANGUAGES[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = GMMEstimationAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()