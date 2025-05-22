import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
from scipy import stats
import matplotlib.pyplot as plt
import pathlib
import pingouin as pg
from docx import Document
from docx.shared import Inches

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "协方差分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'explanation': {
            "协方差分析": "在控制一个或多个协变量的影响下，分析不同组之间因变量的均值是否存在显著差异。",
        },
        'interpretation': {
            "F": "F 统计量，用于检验组间差异的显著性。",
            "p-unc": "未经校正的 p 值，小于显著性水平（通常为 0.05）时，认为组间差异显著。",
            "np2": "偏 eta 平方，反映了组间差异在总变异中所占的比例。"
        }
    },
    'en': {
        'title': "Analysis of Covariance (ANCOVA)",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'explanation': {
            "Analysis of Covariance (ANCOVA)": "Analyze whether there are significant differences in the means of the dependent variable between different groups while controlling for the effects of one or more covariates.",
        },
        'interpretation': {
            "F": "F statistic, used to test the significance of the differences between groups.",
            "p-unc": "Uncorrected p-value. When it is less than the significance level (usually 0.05), the differences between groups are considered significant.",
            "np2": "Partial eta squared, reflecting the proportion of the variance explained by the group differences in the total variance."
        }
    }
}

# 当前语言
current_language = 'en'


def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        file_entry.configure(style="TEntry")


def analyze_file():
    global current_language
    file_path = file_entry.get()
    if file_path == "请输入待分析 Excel 文件的完整路径" or file_path == "Please enter the full path of the Excel file to be analyzed":
        file_path = ""
    if not os.path.exists(file_path):
        result_label.config(text=LANGUAGES[current_language]['file_not_found'])
        return
    try:
        # 打开 Excel 文件
        df = pd.read_excel(file_path)

        # 假设第一列是分组变量，最后一列是因变量，其余列是协变量
        group_var = df.columns[0]
        dep_var = df.columns[-1]
        covar_vars = df.columns[1:-1]

        # 进行协方差分析
        ancova = pg.ancova(data=df, dv=dep_var, between=group_var, covar=covar_vars.tolist())

        # 添加解释说明
        explanations = LANGUAGES[current_language]['explanation']
        interpretations = LANGUAGES[current_language]['interpretation']
        explanation_df = pd.DataFrame([explanations])
        explanation_df = explanation_df.reindex(
            columns=["协方差分析" if current_language == 'zh' else "Analysis of Covariance (ANCOVA)"])
        explanation_df.insert(0, "统计量", "解释说明" if current_language == 'zh' else "Explanation")

        # 添加分析结果解读
        interpretation_df = pd.DataFrame([interpretations])
        interpretation_df = interpretation_df.reindex(columns=[
            "F", "p-unc", "np2"
        ])
        interpretation_df.insert(0, "统计量", "结果解读" if current_language == 'zh' else "Interpretation")

        # 合并数据、解释说明和结果解读
        combined_df = pd.concat([ancova, explanation_df, interpretation_df], ignore_index=True)

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建一个新的Word文档
            doc = Document()

            # 添加标题
            doc.add_heading('协方差分析结果' if current_language == 'zh' else 'Analysis of Covariance (ANCOVA) Results', 0)

            # 添加表格
            table = doc.add_table(rows=1, cols=len(combined_df.columns))
            hdr_cells = table.rows[0].cells
            for col_index, col_name in enumerate(combined_df.columns):
                hdr_cells[col_index].text = col_name

            # 添加数据行
            for index, row in combined_df.iterrows():
                row_cells = table.add_row().cells
                for col_index, value in enumerate(row):
                    row_cells[col_index].text = str(value)

            # 生成结果图片
            desktop_path = pathlib.Path.home() / 'Desktop'
            plot_path = desktop_path / 'ancova_plot.png'
            plt.figure()
            for group in df[group_var].unique():
                group_data = df[df[group_var] == group]
                plt.scatter(group_data[covar_vars[0]], group_data[dep_var], label=group)
            plt.xlabel(covar_vars[0])
            plt.ylabel(dep_var)
            plt.title('协方差分析结果' if current_language == 'zh' else 'Analysis of Covariance (ANCOVA) Results')
            plt.legend()
            plt.savefig(plot_path)
            plt.close()

            # 将图片插入到Word文档中
            doc.add_picture(str(plot_path), width=Inches(6))

            # 保存Word文档
            doc.save(save_path)

            result_msg = LANGUAGES[current_language]['analysis_success'].format(
                save_path) + f"结果图片已保存到 {plot_path}" if current_language == 'zh' else f"The result image has been saved to {plot_path}"
            result_label.config(text=result_msg, wraplength=400)
        else:
            result_label.config(text=LANGUAGES[current_language]['no_save_path'])

    except Exception as e:
        result_label.config(text=LANGUAGES[current_language]['analysis_error'].format(str(e)))


def switch_language():
    global current_language
    current_language = 'en' if current_language == 'zh' else 'zh'
    root.title(LANGUAGES[current_language]['title'])
    select_button.config(text=LANGUAGES[current_language]['select_button'])
    analyze_button.config(text=LANGUAGES[current_language]['analyze_button'])
    language_label.config(text=LANGUAGES[current_language]['switch_language'])
    # 切换语言时更新提示信息
    file_entry.delete(0, tk.END)
    if current_language == 'zh':
        file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
        file_entry.configure(style="Gray.TEntry")
    else:
        file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
        file_entry.configure(style="Gray.TEntry")


def on_entry_click(event):
    """当用户点击输入框时，清除提示信息"""
    if file_entry.get() == "请输入待分析 Excel 文件的完整路径" or file_entry.get() == "Please enter the full path of the Excel file to be analyzed":
        file_entry.delete(0, tk.END)
        file_entry.configure(style="TEntry")


def on_focusout(event):
    """当用户离开输入框时，如果没有输入内容，恢复提示信息"""
    if file_entry.get() == "":
        if current_language == 'zh':
            file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
            file_entry.configure(style="Gray.TEntry")
        else:
            file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
            file_entry.configure(style="Gray.TEntry")


# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(LANGUAGES[current_language]['title'])

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 设置窗口的宽度和高度
window_width = 500
window_height = 300

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建自定义样式
style = ttk.Style()
style.configure("Gray.TEntry", foreground="gray")

# 创建文件选择按钮
select_button = ttk.Button(root, text=LANGUAGES[current_language]['select_button'], command=select_file,
                           bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建文件路径输入框
file_entry = ttk.Entry(root, width=50, style="Gray.TEntry")
if current_language == 'zh':
    file_entry.insert(0, "请输入待分析 Excel 文件的完整路径")
else:
    file_entry.insert(0, "Please enter the full path of the Excel file to be analyzed")
file_entry.pack(pady=5)
file_entry.bind("<FocusIn>", on_entry_click)
file_entry.bind("<FocusOut>", on_focusout)

# 创建分析按钮
analyze_button = ttk.Button(root, text=LANGUAGES[current_language]['analyze_button'], command=analyze_file,
                            bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建语言切换标签
language_label = ttk.Label(root, text=LANGUAGES[current_language]['switch_language'], cursor="hand2")
language_label.pack(pady=10)
language_label.bind("<Button-1>", lambda event: switch_language())

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
try:
    root.mainloop()
except KeyboardInterrupt:
    print("程序被手动中断。")
    root.destroy()  # 销毁主窗口