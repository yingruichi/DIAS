import tkinter as tk
from tkinter import filedialog
import os
import numpy as np
import pandas as pd
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    "zh": {
        "title": "DEMATEL 分析",
        "select_button": "选择文件",
        "analyze_button": "分析文件",
        "file_not_found": "文件不存在，请重新选择。",
        "analysis_success": "分析完成，结果已保存到 {}\n",
        "no_save_path": "未选择保存路径，结果未保存。",
        "analysis_error": "分析文件时出错: {}",
        "switch_language": "切换语言",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "explanation": {
            "综合影响矩阵": "反映因素之间综合影响关系的矩阵",
            "原因度": "衡量因素对其他因素影响程度的指标",
            "中心度": "衡量因素在系统中重要程度的指标"
        },
        "interpretation": {
            "综合影响矩阵": "矩阵元素值越大，对应因素之间的影响越强",
            "原因度": "原因度为正，该因素为原因因素；原因度为负，该因素为结果因素",
            "中心度": "中心度越大，该因素在系统中越重要"
        }
    },
    "en": {
        "title": "DEMATEL Analysis",
        "select_button": "Select File",
        "analyze_button": "Analyze File",
        "file_not_found": "The file does not exist. Please select again.",
        "analysis_success": "Analysis completed. The results have been saved to {}\n",
        "no_save_path": "No save path selected. The results were not saved.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "switch_language": "Switch Language",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "explanation": {
            "综合影响矩阵": "A matrix reflecting the comprehensive influence relationship between factors",
            "原因度": "An indicator to measure the influence degree of a factor on other factors",
            "中心度": "An indicator to measure the importance of a factor in the system"
        },
        "interpretation": {
            "综合影响矩阵": "The larger the matrix element value, the stronger the influence between corresponding factors",
            "原因度": "If the causal degree is positive, the factor is a causal factor; if negative, it is a result factor",
            "中心度": "The larger the centrality, the more important the factor in the system"
        }
    }
}

class DEMATELAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def dematel_analysis(self, data):
        """
        进行 DEMATEL 分析
        :param data: 直接影响矩阵数据
        :return: 综合影响矩阵、原因度、中心度
        """
        # 归一化直接影响矩阵
        n = data.shape[0]
        max_sum_row = np.max(np.sum(data, axis=1))
        max_sum_col = np.max(np.sum(data, axis=0))
        max_value = max(max_sum_row, max_sum_col)
        D = data / max_value

        # 计算综合影响矩阵
        I = np.eye(n)
        T = np.dot(D, np.linalg.inv(I - D))

        # 计算原因度和中心度
        sum_row = np.sum(T, axis=1)
        sum_col = np.sum(T, axis=0)
        causal_degree = sum_row - sum_col
        centrality = sum_row + sum_col

        return T, causal_degree, centrality

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 进行 DEMATEL 分析
            T, causal_degree, centrality = self.dematel_analysis(data)

            # 整理数据
            factors = [f"因素{i + 1}" for i in range(data.shape[0])]
            T_df = pd.DataFrame(T, index=factors, columns=factors)
            causal_degree_df = pd.DataFrame(causal_degree, index=factors, columns=["原因度"])
            centrality_df = pd.DataFrame(centrality, index=factors, columns=["中心度"])

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建一个新的 Word 文档
                document = Document()

                # 添加标题
                title = document.add_heading('DEMATEL 分析结果', level=1)
                title.alignment = 1  # 居中对齐

                # 添加综合影响矩阵
                document.add_heading('综合影响矩阵', level=2)
                document.add_paragraph(explanations["综合影响矩阵"])
                document.add_paragraph(interpretations["综合影响矩阵"])
                table = document.add_table(rows=len(T_df) + 1, cols=len(T_df.columns) + 1)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = ''
                for col_idx, col_name in enumerate(T_df.columns):
                    hdr_cells[col_idx + 1].text = col_name
                for row_idx, row in enumerate(T_df.values):
                    row_cells = table.rows[row_idx + 1].cells
                    row_cells[0].text = T_df.index[row_idx]
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx + 1].text = str(value)

                # 添加原因度
                document.add_heading('原因度', level=2)
                document.add_paragraph(explanations["原因度"])
                document.add_paragraph(interpretations["原因度"])
                table = document.add_table(rows=len(causal_degree_df) + 1, cols=len(causal_degree_df.columns) + 1)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '因素'
                hdr_cells[1].text = '原因度'
                for row_idx, row in enumerate(causal_degree_df.values):
                    row_cells = table.rows[row_idx + 1].cells
                    row_cells[0].text = causal_degree_df.index[row_idx]
                    row_cells[1].text = str(row[0])

                # 添加中心度
                document.add_heading('中心度', level=2)
                document.add_paragraph(explanations["中心度"])
                document.add_paragraph(interpretations["中心度"])
                table = document.add_table(rows=len(centrality_df) + 1, cols=len(centrality_df.columns) + 1)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '因素'
                hdr_cells[1].text = '中心度'
                for row_idx, row in enumerate(centrality_df.values):
                    row_cells = table.rows[row_idx + 1].cells
                    row_cells[0].text = centrality_df.index[row_idx]
                    row_cells[1].text = str(row[0])

                # 生成原因度和中心度柱状图
                fig, axes = plt.subplots(2, 1, figsize=(8, 10))
                axes[0].bar(factors, causal_degree)
                axes[0].set_title('原因度柱状图' if self.current_language == 'zh' else 'Bar Chart of Causal Degree')
                axes[0].set_xlabel('因素' if self.current_language == 'zh' else 'Factors')
                axes[0].set_ylabel('原因度' if self.current_language == 'zh' else 'Causal Degree')

                axes[1].bar(factors, centrality)
                axes[1].set_title('中心度柱状图' if self.current_language == 'zh' else 'Bar Chart of Centrality')
                axes[1].set_xlabel('因素' if self.current_language == 'zh' else 'Factors')
                axes[1].set_ylabel('中心度' if self.current_language == 'zh' else 'Centrality')

                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_charts.png'
                plt.tight_layout()
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                document.add_heading('原因度和中心度柱状图', level=2)
                document.add_picture(img_path)

                # 保存 Word 文档
                document.save(save_path)

                result_msg = languages[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=languages[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(languages[self.current_language]['title'])
        self.select_button.config(text=languages[self.current_language]['select_button'])
        self.analyze_button.config(text=languages[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=languages[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=languages[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = DEMATELAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()