import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
import openpyxl
import os
import numpy as np
import pandas as pd
from tkinter import filedialog
import tkinter as tk
import matplotlib.pyplot as plt
import pathlib
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
LANGUAGES = {
    'zh': {
        'title': "熵值法分析",
        'select_button': "选择文件",
        'analyze_button': "分析文件",
        'file_not_found': "文件不存在，请重新选择。",
        'analysis_success': "分析完成，结果已保存到 {}\n",
        'no_save_path': "未选择保存路径，结果未保存。",
        'analysis_error': "分析文件时出错: {}",
        'switch_language': "切换语言",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'explanation': {
            "指标权重": "通过熵值法计算得到的各指标的权重",
            "指标熵值": "各指标的熵值，反映指标的信息无序程度"
        },
        'interpretation': {
            "指标权重": "权重越大，说明该指标在综合评价中越重要",
            "指标熵值": "熵值越大，说明该指标的信息无序程度越高，提供的有效信息越少"
        }
    },
    'en': {
        'title': "Entropy Method Analysis",
        'select_button': "Select File",
        'analyze_button': "Analyze File",
        'file_not_found': "The file does not exist. Please select again.",
        'analysis_success': "Analysis completed. The results have been saved to {}\n",
        'no_save_path': "No save path selected. The results were not saved.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'switch_language': "Switch Language",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'explanation': {
            "指标权重": "The weights of each indicator calculated by the entropy method",
            "指标熵值": "The entropy value of each indicator, reflecting the degree of information disorder of the indicator"
        },
        'interpretation': {
            "指标权重": "The larger the weight, the more important the indicator is in the comprehensive evaluation",
            "指标熵值": "The larger the entropy value, the higher the degree of information disorder of the indicator and the less effective information it provides"
        }
    }
}

class EntropyMethodAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(LANGUAGES[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(LANGUAGES[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def entropy_method(self, data):
        """
        进行熵值法分析
        :param data: 输入数据，每行代表一个样本，每列代表一个指标
        :return: 指标权重，指标熵值
        """
        # 数据标准化
        standardized_data = data / data.sum(axis=0)

        # 计算每个指标的熵值
        entropy = -np.sum(standardized_data * np.log(standardized_data + 1e-8), axis=0) / np.log(data.shape[0])

        # 计算指标的权重
        weight = (1 - entropy) / np.sum(1 - entropy)

        return weight, entropy

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == LANGUAGES[self.current_language]["file_entry_placeholder"]:
            file_path = ""
        if not os.path.exists(file_path):
            self.result_label.config(text=LANGUAGES[self.current_language]['file_not_found'])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path)
            data = df.values

            # 进行熵值法分析
            weight, entropy = self.entropy_method(data)

            # 整理数据
            data = [
                ["指标权重", weight.tolist(), ""],
                ["指标熵值", entropy.tolist(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = LANGUAGES[self.current_language]['explanation']
            interpretations = LANGUAGES[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(columns=["指标权重", "指标熵值"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(columns=["指标权重", "指标熵值"])
            interpretation_df.insert(0, "统计量_结果解读", "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 合并数据、解释说明和结果解读
            combined_df = pd.concat([df, explanation_df, interpretation_df], ignore_index=True)

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加标题
                doc.add_heading('熵值法分析结果', 0)

                # 添加表格
                table = doc.add_table(rows=1, cols=len(combined_df.columns))
                hdr_cells = table.rows[0].cells
                for col, header in enumerate(combined_df.columns):
                    hdr_cells[col].text = header

                for index, row in combined_df.iterrows():
                    row_cells = table.add_row().cells
                    for col, value in enumerate(row):
                        row_cells[col].text = str(value)

                # 生成指标权重柱状图
                fig, ax = plt.subplots()
                ax.bar(range(len(weight)), weight)
                ax.set_title('指标权重柱状图' if self.current_language == 'zh' else 'Bar Chart of Indicator Weights')
                ax.set_xlabel('指标' if self.current_language == 'zh' else 'Indicators')
                ax.set_ylabel('权重' if self.current_language == 'zh' else 'Weights')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_weights.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_heading('指标权重柱状图', level=1)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文件
                doc.save(save_path)

                result_msg = LANGUAGES[self.current_language]['analysis_success'].format(save_path)
                self.result_label.config(text=result_msg, wraplength=400)

            else:
                self.result_label.config(text=LANGUAGES[self.current_language]['no_save_path'])

        except Exception as e:
            self.result_label.config(text=LANGUAGES[self.current_language]['analysis_error'].format(str(e)))

    def switch_language(self, event=None):
        self.current_language = 'en' if self.current_language == 'zh' else 'zh'
        self.root.title(LANGUAGES[self.current_language]['title'])
        self.select_button.config(text=LANGUAGES[self.current_language]['select_button'])
        self.analyze_button.config(text=LANGUAGES[self.current_language]['analyze_button'])
        self.switch_language_label.config(text=LANGUAGES[self.current_language]['switch_language'])
        # 切换语言时更新提示信息
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, LANGUAGES[self.current_language]['file_entry_placeholder'])
        self.file_entry.config(foreground='gray')

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["select_button"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, LANGUAGES[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=LANGUAGES[self.current_language]["analyze_button"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame, text=LANGUAGES[self.current_language]["switch_language"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()

# 为了向后兼容，保留原来的运行方式
def run_app():
    app = EntropyMethodAnalysisApp()
    app.run()

if __name__ == "__main__":
    run_app()