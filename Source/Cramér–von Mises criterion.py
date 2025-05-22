import tkinter as tk
from tkinter import filedialog
import openpyxl
import os
import pandas as pd
from scipy import stats
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
import numpy as np
from docx import Document

# 定义语言字典
languages = {
    "zh": {
        "title": "Cramér–von Mises 检验",
        "select_button_text": "选择文件",
        "file_entry_placeholder": "请输入待分析 Excel 文件的完整路径",
        "analyze_button_text": "分析文件",
        "no_file_selected": "请选择有效的文件路径。",
        "file_not_exists": "文件不存在，请重新选择。",
        "analysis_error": "分析文件时出错: {}",
        "analysis_complete": "分析完成，结果已保存到 {}，正态分布图、QQ 图和 PP 图已保存。",
        "no_save_path_selected": "未选择保存路径，结果未保存。",
        "columns_stats": ["Cramér–von Mises Statistic", "P-value", "", "结果解读"],
        "interpretation_accept": "在 0.05 的显著性水平下，不能拒绝原假设，样本可能来自正态分布。",
        "interpretation_reject": "在 0.05 的显著性水平下，拒绝原假设，样本不太可能来自正态分布。",
        "switch_language_button_text": "切换语言"
    },
    "en": {
        "title": "Cramér–von Mises test ",
        "select_button_text": "Select File",
        "file_entry_placeholder": "Please enter the full path of the Excel file to be analyzed",
        "analyze_button_text": "Analyze File",
        "no_file_selected": "Please select a valid file path.",
        "file_not_exists": "The file does not exist. Please select again.",
        "analysis_error": "An error occurred while analyzing the file: {}",
        "analysis_complete": "Analysis completed. The results have been saved to {}, and the normal distribution images, QQ plots and PP plots have been saved.",
        "no_save_path_selected": "No save path selected. The results were not saved.",
        "columns_stats": ["Cramér–von Mises Statistic", "P-value", "", "Result Interpretation"],
        "interpretation_accept": "At the 0.05 significance level, the null hypothesis cannot be rejected. The sample may come from a normal distribution.",
        "interpretation_reject": "At the 0.05 significance level, the null hypothesis is rejected. The sample is unlikely to come from a normal distribution.",
        "switch_language_button_text": "Switch Language"
    }
}

# 当前语言，默认为中文
current_language = "en"

def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)
        file_entry.config(foreground='black')

def on_entry_click(event):
    if file_entry.get() == languages[current_language]["file_entry_placeholder"]:
        file_entry.delete(0, tk.END)
        file_entry.config(foreground='black')

def on_focusout(event):
    if file_entry.get() == "":
        file_entry.insert(0, languages[current_language]["file_entry_placeholder"])
        file_entry.config(foreground='gray')

def analyze_file():
    file_path = file_entry.get()
    if file_path == languages[current_language]["file_entry_placeholder"]:
        result_label.config(text=languages[current_language]["no_file_selected"])
        return
    if not os.path.exists(file_path):
        result_label.config(text=languages[current_language]["file_not_exists"])
        return
    try:
        # 打开 Excel 文件
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active

        rows = sheet.max_row
        columns = sheet.max_column

        data = []
        columns_stats = languages[current_language]["columns_stats"]

        for col_idx in range(1, columns + 1):
            column_values = []
            for row_idx in range(2, rows + 1):
                cell_value = sheet.cell(row=row_idx, column=col_idx).value
                if isinstance(cell_value, (int, float)):
                    column_values.append(cell_value)

            if column_values:
                col_name = sheet.cell(row=1, column=col_idx).value
                # 进行 Cramér–von Mises 检验，假设检验样本是否来自正态分布
                result = stats.cramervonmises(column_values, 'norm')
                cvm_statistic = result.statistic
                p_value = result.pvalue

                # 根据 P 值进行结果解读
                if p_value > 0.05:
                    interpretation = languages[current_language]["interpretation_accept"]
                else:
                    interpretation = languages[current_language]["interpretation_reject"]

                values = [cvm_statistic, p_value, None, interpretation]
                data.append([col_name] + values)

                # 绘制直方图和拟合的正态分布曲线
                plt.figure()
                n, bins, patches = plt.hist(column_values, bins=30, density=True, alpha=0.7, color='g')
                mu, std = stats.norm.fit(column_values)
                xmin, xmax = plt.xlim()
                x = np.linspace(xmin, xmax, 100)
                p = stats.norm.pdf(x, mu, std)
                plt.plot(x, p, 'k', linewidth=2)
                title = f'{col_name}: mu = {mu:.2f},  std = {std:.2f}'
                plt.title(title)
                plt.xlabel('Value')
                plt.ylabel('Frequency')

                # 保存图片
                image_path = os.path.splitext(file_path)[0] + f'_{col_name}_normal_distribution.png'
                plt.savefig(image_path)
                plt.close()

                # 手动绘制 PP 图
                sorted_data = np.sort(column_values)
                n = len(sorted_data)
                empirical_cdf = np.arange(1, n + 1) / n
                theoretical_cdf = stats.norm.cdf(sorted_data, loc=mu, scale=std)

                plt.figure()
                plt.plot(theoretical_cdf, empirical_cdf, 'o')
                plt.plot([0, 1], [0, 1], 'r--')
                plt.title(f'{col_name} PP Plot')
                plt.xlabel('Theoretical CDF')
                plt.ylabel('Empirical CDF')
                ppplot_path = os.path.splitext(file_path)[0] + f'_{col_name}_ppplot.png'
                plt.savefig(ppplot_path)
                plt.close()

                # 绘制 QQ 图
                plt.figure()
                stats.probplot(column_values, dist="norm", plot=plt)
                plt.title(f'{col_name} QQ Plot')
                qqplot_path = os.path.splitext(file_path)[0] + f'_{col_name}_qqplot.png'
                plt.savefig(qqplot_path)
                plt.close()

        headers = ["Column Name"] + columns_stats
        df = pd.DataFrame(data, columns=headers)

        # 让用户选择保存路径
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            # 创建 Word 文档
            doc = Document()

            # 添加表格标题
            doc.add_heading('Cramér–von Mises 检验结果', level=1)

            # 创建表格
            table = doc.add_table(rows=1, cols=len(headers))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header

            # 填充表格数据
            for index, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # 保存 Word 文档
            doc.save(save_path)

            # 设置 wraplength 属性让文本自动换行
            result_label.config(text=languages[current_language]["analysis_complete"].format(save_path), wraplength=400)
        else:
            result_label.config(text=languages[current_language]["no_save_path_selected"])

    except Exception as e:
        result_label.config(text=languages[current_language]["analysis_error"].format(str(e)))

def switch_language(event):
    global current_language
    if current_language == "zh":
        current_language = "en"
    else:
        current_language = "zh"

    # 更新界面文字
    root.title(languages[current_language]["title"])
    select_button.config(text=languages[current_language]["select_button_text"])
    file_entry.delete(0, tk.END)
    file_entry.insert(0, languages[current_language]["file_entry_placeholder"])
    file_entry.config(foreground='gray')
    analyze_button.config(text=languages[current_language]["analyze_button_text"])
    switch_language_label.config(text=languages[current_language]["switch_language_button_text"])

# 创建主窗口
root = ttk.Window(themename="flatly")
root.title(languages[current_language]["title"])

# 获取屏幕的宽度和高度
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# 设置窗口的宽度和高度
window_width = 500
window_height = 300

# 计算窗口应该放置的位置
x = (screen_width - window_width) // 2
y = (screen_height - window_height) // 2

# 设置窗口的位置和大小
root.geometry(f"{window_width}x{window_height}+{x}+{y}")

# 创建一个框架来包含按钮和输入框
frame = ttk.Frame(root)
frame.pack(expand=True)

# 创建文件选择按钮
select_button = ttk.Button(frame, text=languages[current_language]["select_button_text"], command=select_file, bootstyle=PRIMARY)
select_button.pack(pady=10)

# 创建文件路径输入框
file_entry = ttk.Entry(frame, width=50)
file_entry.insert(0, languages[current_language]["file_entry_placeholder"])
file_entry.config(foreground='gray')
file_entry.bind('<FocusIn>', on_entry_click)
file_entry.bind('<FocusOut>', on_focusout)
file_entry.pack(pady=5)

# 创建分析按钮
analyze_button = ttk.Button(frame, text=languages[current_language]["analyze_button_text"], command=analyze_file, bootstyle=SUCCESS)
analyze_button.pack(pady=10)

# 创建切换语言标签
switch_language_label = ttk.Label(frame, text=languages[current_language]["switch_language_button_text"], foreground="gray", cursor="hand2")
switch_language_label.bind("<Button-1>", switch_language)
switch_language_label.pack(pady=10)

# 创建结果显示标签
result_label = ttk.Label(root, text="", justify=tk.LEFT)
result_label.pack(pady=10)

# 运行主循环
root.mainloop()