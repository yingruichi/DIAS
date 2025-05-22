import tkinter as tk
from tkinter import filedialog
import os
import pandas as pd
import numpy as np
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 设置 matplotlib 支持中文
plt.rcParams['font.family'] = 'SimHei'
plt.rcParams['axes.unicode_minus'] = False

# 定义语言字典
languages = {
    'zh': {
        'title': "TOPSIS 法分析",
        'select_button_text': "选择文件",
        'file_entry_placeholder': "请输入待分析 Excel 文件的完整路径",
        'analyze_button_text': "分析文件",
        'no_file_selected': "请选择有效的文件路径。",
        'file_not_exists': "文件不存在，请重新选择。",
        'analysis_error': "分析文件时出错: {}",
        'analysis_complete': "分析完成，结果已保存到 {}，相关图片已保存。",
        'no_save_path_selected': "未选择保存路径，结果未保存。",
        'switch_language_button_text': "切换语言",
        'explanation': {
            "标准化决策矩阵": "对原始决策矩阵进行标准化处理后的矩阵",
            "加权标准化决策矩阵": "考虑各属性权重后的标准化决策矩阵",
            "正理想解": "各属性的最优值构成的向量",
            "负理想解": "各属性的最劣值构成的向量",
            "各方案到正理想解的距离": "各方案与正理想解的欧几里得距离",
            "各方案到负理想解的距离": "各方案与负理想解的欧几里得距离",
            "各方案的相对贴近度": "反映各方案与正理想解的相对接近程度",
            "方案排序结果": "根据相对贴近度对各方案进行排序的结果"
        },
        'interpretation': {
            "标准化决策矩阵": "消除不同属性量纲的影响",
            "加权标准化决策矩阵": "体现各属性在决策中的重要性",
            "正理想解": "作为衡量各方案优劣的最优参考点",
            "负理想解": "作为衡量各方案优劣的最劣参考点",
            "各方案到正理想解的距离": "距离越小，方案越优",
            "各方案到负理想解的距离": "距离越大，方案越优",
            "各方案的相对贴近度": "值越接近 1，方案越优",
            "方案排序结果": "排名越靠前，方案越优"
        }
    },
    'en': {
        'title': "TOPSIS Method Analysis",
        'select_button_text': "Select File",
        'file_entry_placeholder': "Please enter the full path of the Excel file to be analyzed",
        'analyze_button_text': "Analyze File",
        'no_file_selected': "Please select a valid file path.",
        'file_not_exists': "The file does not exist. Please select again.",
        'analysis_error': "An error occurred while analyzing the file: {}",
        'analysis_complete': "Analysis completed. The results have been saved to {}, and the relevant images have been saved.",
        'no_save_path_selected': "No save path selected. The results were not saved.",
        'switch_language_button_text': "Switch Language",
        'explanation': {
            "标准化决策矩阵": "The matrix after standardizing the original decision matrix",
            "加权标准化决策矩阵": "The standardized decision matrix considering the weights of each attribute",
            "正理想解": "The vector composed of the optimal values of each attribute",
            "负理想解": "The vector composed of the worst values of each attribute",
            "各方案到正理想解的距离": "The Euclidean distance between each alternative and the positive ideal solution",
            "各方案到负理想解的距离": "The Euclidean distance between each alternative and the negative ideal solution",
            "各方案的相对贴近度": "Reflects the relative closeness of each alternative to the positive ideal solution",
            "方案排序结果": "The result of ranking each alternative according to the relative closeness"
        },
        'interpretation': {
            "标准化决策矩阵": "Eliminate the influence of different attribute dimensions",
            "加权标准化决策矩阵": "Reflect the importance of each attribute in the decision-making",
            "正理想解": "As the optimal reference point for measuring the advantages and disadvantages of each alternative",
            "负理想解": "As the worst reference point for measuring the advantages and disadvantages of each alternative",
            "各方案到正理想解的距离": "The smaller the distance, the better the alternative",
            "各方案到负理想解的距离": "The larger the distance, the better the alternative",
            "各方案的相对贴近度": "The closer the value is to 1, the better the alternative",
            "方案排序结果": "The higher the ranking, the better the alternative"
        }
    }
}


class TOPSISMethodAnalysisApp:
    def __init__(self, root=None):
        # 当前语言，默认为英文
        self.current_language = "en"

        # 如果没有提供root，则创建一个新窗口
        if root is None:
            self.root = ttk.Window(themename="flatly")
            self.root.title(languages[self.current_language]["title"])
        else:
            self.root = root
            self.root.title(languages[self.current_language]["title"])

        self.create_ui()

    def select_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
        if file_path:
            self.file_entry.delete(0, tk.END)
            self.file_entry.insert(0, file_path)
            self.file_entry.config(foreground='black')

    def on_entry_click(self, event):
        if self.file_entry.get() == languages[self.current_language]["file_entry_placeholder"]:
            self.file_entry.delete(0, tk.END)
            self.file_entry.config(foreground='black')

    def on_focusout(self, event):
        if self.file_entry.get() == "":
            self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
            self.file_entry.config(foreground='gray')

    def topsis_method(self, decision_matrix, weight_vector):
        """
        实现 TOPSIS 法
        :param decision_matrix: 决策矩阵
        :param weight_vector: 属性权重向量
        :return: 各方案的相对贴近度和方案排序结果
        """
        # 标准化决策矩阵
        standardized_matrix = decision_matrix / np.sqrt(np.sum(decision_matrix ** 2, axis=0))

        # 加权标准化决策矩阵
        weighted_matrix = standardized_matrix * weight_vector

        # 正理想解和负理想解
        positive_ideal_solution = np.max(weighted_matrix, axis=0)
        negative_ideal_solution = np.min(weighted_matrix, axis=0)

        # 各方案到正理想解和负理想解的距离
        distances_to_positive = np.sqrt(np.sum((weighted_matrix - positive_ideal_solution) ** 2, axis=1))
        distances_to_negative = np.sqrt(np.sum((weighted_matrix - negative_ideal_solution) ** 2, axis=1))

        # 各方案的相对贴近度
        relative_closeness = distances_to_negative / (distances_to_positive + distances_to_negative)

        # 方案排序结果
        ranking = np.argsort(-relative_closeness) + 1

        return standardized_matrix, weighted_matrix, positive_ideal_solution, negative_ideal_solution, \
            distances_to_positive, distances_to_negative, relative_closeness, ranking

    def analyze_file(self):
        file_path = self.file_entry.get()
        if file_path == languages[self.current_language]["file_entry_placeholder"]:
            self.result_label.config(text=languages[self.current_language]["no_file_selected"])
            return
        if not os.path.exists(file_path):
            self.result_label.config(text=languages[self.current_language]["file_not_exists"])
            return
        try:
            # 打开 Excel 文件
            df = pd.read_excel(file_path, header=None)
            data = df.values

            # 将数据转换为浮点类型
            data = data.astype(float)

            # 假设第一行为属性权重向量，其余行为决策矩阵
            weight_vector = data[0]
            decision_matrix = data[1:]

            # 进行 TOPSIS 分析
            standardized_matrix, weighted_matrix, positive_ideal_solution, negative_ideal_solution, \
                distances_to_positive, distances_to_negative, relative_closeness, ranking = self.topsis_method(
                decision_matrix,
                weight_vector)

            # 整理数据
            data = [
                ["标准化决策矩阵", standardized_matrix.tolist(), ""],
                ["加权标准化决策矩阵", weighted_matrix.tolist(), ""],
                ["正理想解", positive_ideal_solution.tolist(), ""],
                ["负理想解", negative_ideal_solution.tolist(), ""],
                ["各方案到正理想解的距离", distances_to_positive.tolist(), ""],
                ["各方案到负理想解的距离", distances_to_negative.tolist(), ""],
                ["各方案的相对贴近度", relative_closeness.tolist(), ""],
                ["方案排序结果", ranking.tolist(), ""]
            ]
            headers = ["统计量", "统计量值", "p值"]
            df_result = pd.DataFrame(data, columns=headers)

            # 添加解释说明
            explanations = languages[self.current_language]['explanation']
            interpretations = languages[self.current_language]['interpretation']
            explanation_df = pd.DataFrame([explanations])
            explanation_df = explanation_df.reindex(
                columns=["标准化决策矩阵", "加权标准化决策矩阵", "正理想解", "负理想解", "各方案到正理想解的距离",
                         "各方案到负理想解的距离", "各方案的相对贴近度", "方案排序结果"])
            explanation_df.insert(0, "统计量_解释说明", "解释说明" if self.current_language == 'zh' else "Explanation")

            # 添加分析结果解读
            interpretation_df = pd.DataFrame([interpretations])
            interpretation_df = interpretation_df.reindex(
                columns=["标准化决策矩阵", "加权标准化决策矩阵", "正理想解", "负理想解", "各方案到正理想解的距离",
                         "各方案到负理想解的距离", "各方案的相对贴近度", "方案排序结果"])
            interpretation_df.insert(0, "统计量_结果解读",
                                     "结果解读" if self.current_language == 'zh' else "Interpretation")

            # 让用户选择保存路径
            save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
            if save_path:
                # 创建 Word 文档
                doc = Document()

                # 添加结果表格
                table = doc.add_table(rows=df_result.shape[0] + 1, cols=df_result.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(headers):
                    hdr_cells[col_idx].text = header
                for row_idx, row in df_result.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加解释说明表格
                doc.add_paragraph()
                doc.add_heading("解释说明" if self.current_language == 'zh' else "Explanation", level=2)
                table = doc.add_table(rows=explanation_df.shape[0] + 1, cols=explanation_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(explanation_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in explanation_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 添加分析结果解读表格
                doc.add_paragraph()
                doc.add_heading("结果解读" if self.current_language == 'zh' else "Interpretation", level=2)
                table = doc.add_table(rows=interpretation_df.shape[0] + 1, cols=interpretation_df.shape[1])
                hdr_cells = table.rows[0].cells
                for col_idx, header in enumerate(interpretation_df.columns):
                    hdr_cells[col_idx].text = header
                for row_idx, row in interpretation_df.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)

                # 生成各方案相对贴近度柱状图
                fig, ax = plt.subplots()
                ax.bar(range(len(relative_closeness)), relative_closeness)
                ax.set_title(
                    '各方案相对贴近度柱状图' if self.current_language == 'zh' else 'Bar Chart of Relative Closeness of Each Alternative')
                ax.set_xlabel('方案编号' if self.current_language == 'zh' else 'Alternative Number')
                ax.set_ylabel('相对贴近度' if self.current_language == 'zh' else 'Relative Closeness')
                # 保存图片
                img_path = os.path.splitext(save_path)[0] + '_relative_closeness.png'
                plt.savefig(img_path)
                plt.close()

                # 将图片插入到 Word 文档中
                doc.add_paragraph()
                doc.add_heading(
                    "各方案相对贴近度柱状图" if self.current_language == 'zh' else 'Bar Chart of Relative Closeness of Each Alternative',
                    level=2)
                doc.add_picture(img_path, width=Inches(6))

                # 保存 Word 文档
                doc.save(save_path)

                self.result_label.config(text=languages[self.current_language]["analysis_complete"].format(save_path),
                                         wraplength=400)
            else:
                self.result_label.config(text=languages[self.current_language]["no_save_path_selected"])

        except Exception as e:
            self.result_label.config(text=languages[self.current_language]["analysis_error"].format(str(e)))

    def switch_language(self, event):
        if self.current_language == "zh":
            self.current_language = "en"
        else:
            self.current_language = "zh"

        # 更新界面文字
        self.root.title(languages[self.current_language]["title"])
        self.select_button.config(text=languages[self.current_language]["select_button_text"])
        self.file_entry.delete(0, tk.END)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.analyze_button.config(text=languages[self.current_language]["analyze_button_text"])
        self.switch_language_label.config(text=languages[self.current_language]["switch_language_button_text"])

    def create_ui(self):
        # 获取屏幕的宽度和高度
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # 设置窗口的宽度和高度
        window_width = 500
        window_height = 300

        # 计算窗口应该放置的位置
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # 设置窗口的位置和大小
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # 创建一个框架来包含按钮和输入框
        frame = ttk.Frame(self.root)
        frame.pack(expand=True)

        # 创建文件选择按钮
        self.select_button = ttk.Button(frame, text=languages[self.current_language]["select_button_text"],
                                        command=self.select_file, bootstyle=PRIMARY)
        self.select_button.pack(pady=10)

        # 创建文件路径输入框
        self.file_entry = ttk.Entry(frame, width=50)
        self.file_entry.insert(0, languages[self.current_language]["file_entry_placeholder"])
        self.file_entry.config(foreground='gray')
        self.file_entry.bind('<FocusIn>', self.on_entry_click)
        self.file_entry.bind('<FocusOut>', self.on_focusout)
        self.file_entry.pack(pady=5)

        # 创建分析按钮
        self.analyze_button = ttk.Button(frame, text=languages[self.current_language]["analyze_button_text"],
                                         command=self.analyze_file, bootstyle=SUCCESS)
        self.analyze_button.pack(pady=10)

        # 创建切换语言标签
        self.switch_language_label = ttk.Label(frame,
                                               text=languages[self.current_language]["switch_language_button_text"],
                                               foreground="gray", cursor="hand2")
        self.switch_language_label.bind("<Button-1>", self.switch_language)
        self.switch_language_label.pack(pady=10)

        # 创建结果显示标签
        self.result_label = ttk.Label(self.root, text="", justify=tk.LEFT)
        self.result_label.pack(pady=10)

    def run(self):
        # 运行主循环
        self.root.mainloop()


# 为了向后兼容，保留原来的运行方式
def run_app():
    app = TOPSISMethodAnalysisApp()
    app.run()


if __name__ == "__main__":
    run_app()